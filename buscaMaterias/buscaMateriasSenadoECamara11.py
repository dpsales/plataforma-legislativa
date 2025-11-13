raise SystemExit(
    "Este módulo foi substituído pelo microserviço Django em monitoramento/."
)

import re
import os
import io
import requests
import pandas as pd
from datetime import datetime, timedelta
from docx import Document
import xml.etree.ElementTree as ET
from flask import Flask, send_file

from dash import Dash, html, dcc
from dash.dash_table import DataTable
from dash.dependencies import Input, Output, State
import dash_bootstrap_components as dbc

from requests.adapters import HTTPAdapter
from requests.packages.urllib3.util.retry import Retry

# ---------------------------------------------------------------------
# 1) Sessão requests com retries
# ---------------------------------------------------------------------
session = requests.Session()
retries = Retry(total=5, backoff_factor=1, status_forcelist=[500, 502, 503, 504])
session.mount("https://", HTTPAdapter(max_retries=retries))

# ---------------------------------------------------------------------
# 2) Funções auxiliares (mantidas sem alteração)
# ---------------------------------------------------------------------
def parse_data_hora(data_str, formatos):
    for fmt in formatos:
        try:
            return datetime.strptime(data_str, fmt).strftime("%Y-%m-%d %H:%M:%S")
        except ValueError:
            continue
    return data_str

def substituir_http_por_https(link):
    if isinstance(link, str) and link.startswith("http://legis.senado.leg.br"):
        return link.replace("http://legis.senado.leg.br", "https://legis.senado.leg.br", 1)
    return link

# ---------------------------------------------------------------------
# 3) Senadores em exercício (mantida)
# ---------------------------------------------------------------------
def obter_lista_senadores_em_exercicio():
    url = "https://legis.senado.leg.br/dadosabertos/senador/lista/atual.json"
    try:
        r = session.get(url, timeout=30); r.raise_for_status()
        parl = r.json()["ListaParlamentarEmExercicio"]["Parlamentares"]["Parlamentar"]
        return {
            p["IdentificacaoParlamentar"]["NomeParlamentar"].lower(): (
                p["IdentificacaoParlamentar"].get("SiglaPartidoParlamentar", "").strip(),
                p["IdentificacaoParlamentar"].get("UfParlamentar", "").strip()
            )
            for p in parl
        }
    except:
        return {}

dict_senadores = obter_lista_senadores_em_exercicio()

def ajustar_autor_senado(autor):
    if not autor or autor == "Autor não disponível":
        return autor
    nome = autor.rsplit("-", 1)[-1].strip() if " - " in autor else autor.strip()
    partido, uf = dict_senadores.get(nome.lower(), ("", ""))
    partido = partido if partido and partido != "Sem Partido" else ""
    uf = uf if uf and uf != "Sem UF" else ""
    return f"{nome} ({partido}/{uf})".replace("//", "/").strip(" ()/") if partido or uf else nome

# ---------------------------------------------------------------------
# 4) Funções de busca (Câmara sem alteração; Senado ajustado)
# ---------------------------------------------------------------------
#Câmara

# --- fallback para status da Câmara quando null ---
def obter_status_proposicao_camara(cod):
    url = f"https://dadosabertos.camara.leg.br/api/v2/proposicoes/{cod}"
    r = session.get(url, timeout=30)
    r.raise_for_status()
    d = r.json().get("dados", {})
    return d.get("statusProposicao", {}).get("descricaoSituacao", "Status não disponível")

def obter_ultima_movimentacao_camara(cod):
    url = f"https://dadosabertos.camara.leg.br/api/v2/proposicoes/{cod}/tramitacoes"
    try:
        r = session.get(url, timeout=30); r.raise_for_status()
        dados = r.json()["dados"]
        if dados:
            for d in dados:
                d["dataHora_dt"] = datetime.strptime(d["dataHora"], "%Y-%m-%dT%H:%M")
            dados.sort(key=lambda x: (x["sequencia"], x["dataHora_dt"]), reverse=True)
            ult = dados[0]
            # se descricaoSituacao for None, buscamos no endpoint da proposição
            status = ult.get("descricaoSituacao") \
                     or obter_status_proposicao_camara(cod)
            return {
                "ID": cod,
                "Tipo": "Câmara",
                "DataMovimentacao": ult["dataHora_dt"].strftime("%Y-%m-%d %H:%M:%S"),
                "Status": status,
                "UltimaMovimentacao": ult.get("despacho", "N/A")
            }
    except requests.exceptions.RequestException as e:
        print("[ERRO Câmara]", e)
    return None

def obter_autor_proposicao_camara(cod):
    url = f"https://dadosabertos.camara.leg.br/api/v2/proposicoes/{cod}/autores"
    try:
        r = requests.get(url, timeout=30)
        if r.status_code == 200:
            autores = r.json().get("dados", [])
            autor1 = next((a for a in autores if a.get("ordemAssinatura") == 1), autores[0] if autores else None)
            if not autor1:
                return "Autor não disponível"
            nome = autor1.get("nome", "Autor não disponível")
            if autor1.get("codTipo") == 1:
                dep_uri = autor1.get("uri", "")
                if dep_uri:
                    dep = requests.get(dep_uri, timeout=30).json().get("dados", {}).get("ultimoStatus", {})
                    partido = dep.get("siglaPartido", "").strip(); uf = dep.get("siglaUf", "").strip()
                    partido = partido if partido != "Sem Partido" else ""; uf = uf if uf != "Sem UF" else ""
                    if partido or uf:
                        return f"{nome} ({partido}/{uf})".replace("//", "/").strip(" ()/")
            return nome
    except:
        pass
    return "Autor não disponível"

# --- Senado (novo) ---
def obter_id_processo_senado(codigo_materia):
    url = f"https://legis.senado.leg.br/dadosabertos/processo?codigoMateria={codigo_materia}"
    r = session.get(url, timeout=30); r.raise_for_status()
    resultados = r.json()
    if resultados:
        return resultados[0]["id"]
    return None

def obter_dados_processo_senado(id_processo):
    url = f"https://legis.senado.leg.br/dadosabertos/processo/{id_processo}"
    r = session.get(url, timeout=30); r.raise_for_status()
    d = r.json()

    # Título e ementa
    titulo = f"{d.get('sigla')} {d.get('numero')}/{d.get('ano')}"
    ementa = d.get("conteudo", {}).get("ementa", "Ementa não disponível")

    # Autor
    autor = d.get("documento", {}).get("resumoAutoria", "Autor não disponível")

    # Status atual
    status = ""
    autuacoes = d.get("autuacoes", [])
    if autuacoes:
        situacoes = autuacoes[0].get("situacoes", [])
        if situacoes:
            ultima_sit = max(situacoes, key=lambda s: datetime.strptime(s["inicio"], "%Y-%m-%d"))
            status = ultima_sit.get("descricao", "")

    # Última movimentação e data
    ultima_mov, data_mov = "", ""
    if autuacoes:
        informes = autuacoes[0].get("informesLegislativos", [])
        if informes:
            inf = informes[-1]
            ultima_mov = inf.get("descricao", "")
            data_mov = inf.get("data", "")

    # === Ajuste aqui: monta o link de ficha com o código da matéria ===
    codigo_materia = d.get("codigoMateria")
    ficha_url = f"https://www25.senado.leg.br/web/atividade/materias/-/materia/{codigo_materia}"
    url_inteiro = d.get("documento", {}).get("url", "")

    link_ficha_md    = f"[Ficha]({ficha_url})" if codigo_materia else "Não disponível"
    link_inteiro_md  = f"[Inteiro]({url_inteiro})" if url_inteiro else "Não disponível"

    return {
        "Titulo": titulo,
        "Ementa": ementa,
        "Autor": autor,
        "Status": status,
        "UltimaMovimentacao": ultima_mov,
        "DataMovimentacao": data_mov,
        "LinkFicha": link_ficha_md,
        "LinkInteiroTeor": link_inteiro_md
    }

#Ambas as casas

def obter_titulo_pl(cod, casa):
    if casa == "Senado":
        url = f"https://legis.senado.leg.br/dadosabertos/materia/{cod}"
        try:
            r = requests.get(url, timeout=30); r.raise_for_status()
            root = ET.fromstring(r.content)
            titulo = root.findtext(".//DescricaoIdentificacaoMateria", "Título não disponível")
            ementa = root.findtext(".//EmentaMateria", "Ementa não disponível")
            autor = ajustar_autor_senado(root.findtext(".//Autor", "Autor não disponível"))
            ficha = substituir_http_por_https(f"https://www25.senado.leg.br/web/atividade/materias/-/materia/{cod}")
            inteiro = substituir_http_por_https(obter_inteiro_teor_senado(cod))
            return titulo, ementa, autor, ficha, inteiro
        except requests.exceptions.RequestException:
            pass
    else:
        url = f"https://dadosabertos.camara.leg.br/api/v2/proposicoes/{cod}"
        try:
            r = requests.get(url, timeout=30); r.raise_for_status()
            d = r.json()["dados"]
            titulo = f"{d['siglaTipo']} {d['numero']}/{d['ano']}"
            autor = obter_autor_proposicao_camara(cod)
            ficha = f"https://www.camara.leg.br/proposicoesWeb/fichadetramitacao?idProposicao={cod}"
            return titulo, d.get("ementa", "Ementa não disponível"), autor, ficha, d.get("urlInteiroTeor", "Não disponível")
        except requests.exceptions.RequestException:
            pass
    return ("Título não disponível", "Ementa não disponível", "Autor não disponível", "Não disponível", "Não disponível")

# ---------------------------------------------------------------------
# 5) Leitura do Excel com IDs, Casa e Secretaria
# ---------------------------------------------------------------------
EXCEL_IDS = "proposicoes_ids.xlsx"
df_ids = pd.read_excel(EXCEL_IDS, dtype={"ID": int, "Casa": str, "Secretaria": str})
df_ids["Casa"] = df_ids["Casa"].str.strip().str.upper().map({"CD": "Câmara", "SF": "Senado"}).fillna(df_ids["Casa"])

# ---------------------------------------------------------------------
# 6) Coleta inicial (ajustada)
# ---------------------------------------------------------------------
dados = []
for _, r in df_ids.iterrows():
    cod, casa, sec = int(r["ID"]), r["Casa"], r["Secretaria"]
    if casa == "Senado":
        id_proc = obter_id_processo_senado(cod)
        if not id_proc:
            continue
        mov = obter_dados_processo_senado(id_proc)
        mov.update({"ID": cod, "Tipo": "Senado", "Secretaria": sec})
    else:
        mov = obter_ultima_movimentacao_camara(cod)
        if not mov:
            continue
        titulo, ementa, autor, ficha, inteiro = obter_titulo_pl(cod, casa)
        mov.update({
            "Titulo": titulo, "Ementa": ementa, "Autor": autor, "Secretaria": sec,
            "LinkFicha": f"[Ficha]({ficha})" if ficha != "Não disponível" else ficha,
            "LinkInteiroTeor": f"[Inteiro]({inteiro})" if inteiro != "Não disponível" else inteiro
        })
    dados.append(mov)

df_projetos = pd.DataFrame(dados)
ultima_atualizacao = datetime.now() - timedelta(hours=3)
sec_opts = sorted(df_projetos["Secretaria"].dropna().unique())

# ---------------------------------------------------------------------
# 7) Dash App
# ---------------------------------------------------------------------
server = Flask(__name__)
app = Dash(__name__, server=server,
           external_stylesheets=[dbc.themes.BOOTSTRAP],
           title="Matérias Parlamentares de Interesse")

titulo_style = {
    "fontWeight": "bold", "marginRight": "10px",
    "fontFamily": "Verdana", "fontSize": "14px",
}

app.layout = dbc.Container([
    html.H1("Matérias Parlamentares de Interesse",
            style={"color": "#183EFF", "fontFamily": "Verdana"}),
    html.Hr(),
    dcc.Interval(id="interval-atualizacao", interval=60_000, n_intervals=0),
    html.Div(id="mensagem-atualizacao",
             style={"textAlign": "right",
                    "fontStyle": "italic",
                    "fontSize": "13px"}),

    # ---------------- filtros ----------------
    dbc.Row(dbc.Col([
        html.Span("Filtro por casa:", style=titulo_style),
        dcc.Checklist(
            id="checkbox-casa",
            options=[{"label": " Câmara", "value": "Câmara"},
                     {"label": " Senado", "value": "Senado"}],
            value=["Câmara", "Senado"],
            labelStyle={"display": "inline-block", "margin-right": "15px",
                        "fontFamily": "Verdana"},
            inputStyle={"margin-right": "5px"},
        ),
    ]), className="mb-2"),

    dbc.Row(dbc.Col([
        html.Span("Filtro por unidade:", style=titulo_style),
        dcc.Checklist(
            id="checkbox-secretaria",
            options=[{"label": f" {s}", "value": s} for s in sec_opts],
            value=sec_opts,
            labelStyle={"display": "inline-block", "margin-right": "15px",
                        "fontFamily": "Verdana"},
            inputStyle={"margin-right": "5px"},
        ),
    ]), className="mb-3"),

    # ------------ botões exportação ----------
    dbc.Row([
        dbc.Col(html.A("Exportar para Excel", href="/download_excel",
                       className="btn btn-success me-2"), width="auto"),
        dbc.Col(html.A("Exportar para Word", href="/download_word",
                       className="btn btn-primary me-2"), width="auto"),
    ], className="mb-3"),

    # --------------- CSS extra ---------------
    dcc.Markdown("""
    <style>
        .dash-table-container .dash-header {font-size: 16px !important;}
        .dash-table-container .dash-sort   {font-size: 20px !important; margin-left: 10px !important;}
        .dash-table-container              {overflow-x: hidden !important; max-width: 100% !important;}
        .dash-table-container .dash-spreadsheet-container {max-width: 100% !important;}
    </style>
    """, dangerously_allow_html=True),

    # --------------- DataTable ---------------
    DataTable(
        id="tabela-projetos",
        data=df_projetos.to_dict("records"),
        columns=[
            {"name": "Casa",          "id": "Tipo"},
            {"name": "Secretaria",    "id": "Secretaria"},
            {"name": "Proposição",    "id": "Titulo"},
            {"name": "Autor",         "id": "Autor"},
            {"name": "Ementa",        "id": "Ementa"},
            {"name": "Status",        "id": "Status"},
            {"name": "Última movimentação",     "id": "UltimaMovimentacao"},
            {"name": "Data",          "id": "DataMovimentacao"},
            {"name": "Ficha",         "id": "LinkFicha",       "presentation": "markdown"},
            {"name": "Inteiro Teor",  "id": "LinkInteiroTeor", "presentation": "markdown"},
        ],
        fixed_rows={"headers": True},
        style_table={"overflowY": "auto", "overflowX": "visible", "maxHeight": "80vh"},
        style_cell_conditional=[
            {"if": {"column_id": "Tipo"},             "width": "80px",  "maxWidth": "90px"},
            {"if": {"column_id": "Secretaria"},       "width": "100px", "maxWidth": "100px"},
            {"if": {"column_id": "Titulo"},           "width": "100px", "maxWidth": "100px"},
            {"if": {"column_id": "Ementa"},           "width": "400px", "maxWidth": "400px"},
            {"if": {"column_id": "Autor"},            "width": "140px", "maxWidth": "140px"},
            {"if": {"column_id": "DataMovimentacao"}, "width": "120px"},
            {"if": {"column_id": "Status"},           "width": "180px", "maxWidth": "180px"},
            {"if": {"column_id": "UltimaMovimentacao"},  "width": "140px", "maxWidth": "170px"},
            {"if": {"column_id": "LinkFicha"},        "width": "90px"},
            {"if": {"column_id": "LinkInteiroTeor"},  "width": "90px"},
        ],
        style_cell={"textAlign": "left", "fontFamily": "Verdana",
                    "fontSize": "14px", "whiteSpace": "normal", "height": "auto"},
        style_header={"backgroundColor": "#183EFF", "fontWeight": "bold", "color": "#FFFFFF"},
        sort_action="native", filter_action="native", page_action="native", page_size=100,
        row_selectable="multi",
    ),

    html.Button("Gerar Mensagem WhatsApp", id="gerar-msg-whatsapp",
                n_clicks=0, className="btn btn-secondary", style={"marginTop": "10px"}),
    dcc.Textarea(id="whatsapp-message",
                 style={"width": "100%", "height": "150px", "marginTop": "10px"},
                 placeholder="A mensagem gerada aparecerá aqui..."),
],
    fluid=True   # <-- fecha a chamada do Container
)                # <-- parêntese que faltava

# ---------------------------------------------------------------------
# 8) Callbacks
# ---------------------------------------------------------------------
@app.callback(Output("mensagem-atualizacao", "children"), Input("interval-atualizacao", "n_intervals"))
def msg_atualizacao(_):
    data = ultima_atualizacao.strftime("%d/%m/%Y")
    hora = ultima_atualizacao.strftime("%H:%M")
    return [
        f"Dados atualizados em {data} às {hora}",
        html.Br(),
        "Dados obtidos via API da Câmara e do Senado"
    ]

@app.callback(Output("tabela-projetos", "data"),
              Input("checkbox-casa", "value"), Input("checkbox-secretaria", "value"))
def filtrar(casas, secs):
    df = df_projetos
    if casas: df = df[df["Tipo"].isin(casas)]
    if secs:  df = df[df["Secretaria"].isin(secs)]
    return df.to_dict("records")

@app.callback(Output("whatsapp-message", "value"),
              Input("gerar-msg-whatsapp", "n_clicks"),
              State("tabela-projetos", "selected_rows"), State("tabela-projetos", "data"),
              prevent_initial_call=True)
def gerar_msg(_, rows, data):
    if not rows: return ""
    msgs = []
    for i in rows:
        l = data[i]
        link = re.search(r"\[Ficha\]\(([^)]+)\)", l["LinkFicha"] or "")
        link_url = link.group(1) if link else ""
        try:
            data_fmt = datetime.strptime(l["DataMovimentacao"], "%Y-%m-%d %H:%M:%S").strftime("%d/%m/%Y %H:%M")
        except:
            data_fmt = l["DataMovimentacao"]
        msg = (f"📌 *Proposição*: {l['Titulo']}\n"
               f"👤 *Autor*: {l['Autor']}\n"
               f"📝 *Ementa*: {l['Ementa']}\n"
               f"🏢 *Secretaria*: {l['Secretaria']}\n"
               f"📅 *Data de atualização*: {data_fmt}\n"
               f"🔎 *Status*: {l['Status']}\n"
               f"✏️ *Última movimentação*: {l['UltimaMovimentacao']}\n")
        if link_url: msg += f"🔗 *Ficha*: {link_url}\n"
        msg += "-"*40 + "\n"
        msgs.append(msg)
    return "📢 *Atualização Legislativa* 📢\n\n" + "".join(msgs)

# ---------------------------------------------------------------------
# 9) Rotas de exportação
# ---------------------------------------------------------------------
@server.route("/download_excel")
def dl_excel():
    try:
        out = io.BytesIO()
        with pd.ExcelWriter(out, engine="xlsxwriter") as w:
            df = df_projetos.copy()
            df["DataMovimentacao"] = pd.to_datetime(df["DataMovimentacao"]).dt.strftime("%d/%m/%Y %H:%M")
            url = lambda md: re.search(r"\(([^)]+)\)", md).group(1) if isinstance(md, str) and md.startswith("[") else md
            df["LinkFicha"] = df["LinkFicha"].apply(url); df["LinkInteiroTeor"] = df["LinkInteiroTeor"].apply(url)
            df.to_excel(w, index=False, sheet_name="Projetos")
        out.seek(0)
        return send_file(out, download_name="projetos_legislativos.xlsx", as_attachment=True,
                         mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
    except Exception as e:
        return f"Erro: {e}", 500

@server.route("/download_word")
def dl_word():
    df = df_projetos.copy()
    df["DataMovimentacao_dt"] = pd.to_datetime(df["DataMovimentacao"], errors="coerce")
    df.sort_values("DataMovimentacao_dt", ascending=False, inplace=True)
    doc = Document(); doc.add_heading("Relatório de Projetos Legislativos", level=1)
    for _, r in df.iterrows():
        try: data_fmt = datetime.strptime(r["DataMovimentacao"], "%Y-%m-%d %H:%M:%S").strftime("%d/%m/%Y")
        except: data_fmt = r["DataMovimentacao"]
        p = doc.add_paragraph()
        p.add_run(f"{r['Titulo']} ").bold = True
        p.add_run(f"({r['Tipo']} - {r['Secretaria']}) ").bold = True
        p.add_run(f"- {r['Ementa']} - {data_fmt} - {r['Status']}")
    out = io.BytesIO(); doc.save(out); out.seek(0)
    return send_file(out, download_name="projetos_legislativos.docx", as_attachment=True)

# ---------------------------------------------------------------------
# 10) Run
# ---------------------------------------------------------------------
if __name__ == "__main__":
    port = int(os.environ.get("PORT", 8080))
    app.run(debug=False, host="0.0.0.0", port=port)
