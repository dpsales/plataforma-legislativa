from __future__ import annotations

import io
import json
from typing import Iterable

import pandas as pd
from django.contrib import messages
from django.db.models import Q
from django.http import HttpRequest, HttpResponse, JsonResponse
from django.shortcuts import get_object_or_404, redirect, render
from django.urls import reverse
from django.utils import timezone
from django.views.decorators.http import require_POST
from docx import Document

from .forms import TrackedPropositionForm, UploadJsonForm
from .models import TrackedDocument, TrackedProposition
from .services.sync import sync_document
from .utils import build_document_payload, load_document_from_json

ALLOWED_CONFIG_PROFILES = {"admin", "normal"}


def _request_profile(request: HttpRequest) -> str:
    for source in (request.GET, request.POST):
        profile = source.get("profile")
        if profile:
            return profile.strip().lower()
    header = request.META.get("HTTP_X_USER_PROFILE")
    if header:
        return header.strip().lower()
    return ""


def _prefixed_path(request: HttpRequest, path: str) -> str:
    prefix = request.META.get("HTTP_X_FORWARDED_PREFIX") or request.META.get("SCRIPT_NAME") or ""
    prefix = prefix.rstrip("/")
    if path.startswith("http://") or path.startswith("https://"):
        return path
    relative = path.lstrip("/")
    if prefix:
        base = prefix or "/"
        return f"{base}/{relative}" if relative else base
    return f"/{relative}" if relative else "/"


def _append_profile(path: str, profile: str) -> str:
    if not profile:
        return path
    separator = "&" if "?" in path else "?"
    return f"{path}{separator}profile={profile}"


def _parse_multi(values: str) -> list[str]:
    if not values:
        return []
    return [item.strip() for item in values.split(",") if item.strip()]


def _active_document() -> TrackedDocument | None:
    return TrackedDocument.objects.filter(slug=TrackedDocument.DEFAULT_SLUG).first()


def index(request: HttpRequest) -> HttpResponse:
    profile = _request_profile(request)
    can_configure = profile in ALLOWED_CONFIG_PROFILES

    document = _active_document()
    queryset = document.propositions.all() if document else TrackedProposition.objects.none()

    casas = []
    for casa_value, casa_label in TrackedProposition.CASA_CHOICES:
        if queryset.filter(casa=casa_value).exists():
            casas.append({"value": casa_value, "label": casa_label})

    secretarias = sorted(
        [
            value
            for value in queryset.order_by().values_list("secretaria", flat=True).distinct()
            if value
        ]
    )
    prioridades = sorted(
        {
            valor
            for valor in queryset.order_by().values_list("prioridade", flat=True)
            if valor is not None
        }
    )

    last_update = None
    if queryset.exists():
        last_update = queryset.order_by("-updated_at").values_list("updated_at", flat=True).first()

    api_url = _append_profile(_prefixed_path(request, reverse("monitoramento:propositions-api")), profile)
    export_xlsx_url = _append_profile(_prefixed_path(request, reverse("monitoramento:export-xlsx")), profile)
    export_docx_url = _append_profile(_prefixed_path(request, reverse("monitoramento:export-docx")), profile)
    configure_url = _append_profile(_prefixed_path(request, reverse("monitoramento:configure")), profile) if can_configure else ""

    context = {
        "document": document,
        "document_updated_at": timezone.localtime(document.updated_at) if document else None,
        "document_reference": document.reference_label if document else "",
        "document_description": document.description if document else "",
        "has_dataset": queryset.exists(),
        "casas": casas,
        "secretarias": secretarias,
        "prioridades": prioridades,
        "api_url": api_url,
        "export_xlsx_url": export_xlsx_url,
        "export_docx_url": export_docx_url,
        "configure_url": configure_url,
        "static_prefix": _prefixed_path(request, "static").rstrip("/"),
        "profile": profile,
        "can_configure": can_configure,
    "last_update": timezone.localtime(last_update) if last_update else None,
    }
    return render(request, "monitoramento/index.html", context)


def _filter_queryset(
    queryset,
    casas: Iterable[str],
    secretarias: Iterable[str],
    prioridades: Iterable[str],
    busca: str,
):
    if casas:
        queryset = queryset.filter(casa__in=list(casas))
    if secretarias:
        queryset = queryset.filter(secretaria__in=list(secretarias))
    if prioridades:
        valores = []
        for item in prioridades:
            try:
                valores.append(int(item))
            except (TypeError, ValueError):
                continue
        if valores:
            queryset = queryset.filter(prioridade__in=valores)
    if busca:
        queryset = queryset.filter(
            Q(titulo__icontains=busca)
            | Q(ementa__icontains=busca)
            | Q(assunto__icontains=busca)
            | Q(justificativa__icontains=busca)
        )
    return queryset


def propositions_api(request: HttpRequest) -> JsonResponse:
    document = _active_document()
    if not document:
        return JsonResponse({"results": []})

    casas = _parse_multi(request.GET.get("casas", ""))
    secretarias = _parse_multi(request.GET.get("secretarias", ""))
    prioridades = _parse_multi(request.GET.get("prioridades", ""))
    busca = request.GET.get("busca", "").strip()

    queryset = document.propositions.all()
    queryset = _filter_queryset(queryset, casas, secretarias, prioridades, busca)

    dados = []
    for item in queryset.order_by("-data_movimentacao", "-updated_at"):
        dados.append(
            {
                "id": item.proposition_id,
                "casa": item.get_casa_display(),
                "secretaria": item.secretaria,
                "titulo": item.titulo or item.assunto or "Título não disponível",
                "autor": item.autor,
                "ementa": item.ementa,
                "status": item.status,
                "ultima_movimentacao": item.ultima_movimentacao,
                "data_movimentacao": item.data_movimentacao.strftime("%d/%m/%Y %H:%M") if item.data_movimentacao else "",
                "link_ficha": item.link_ficha,
                "link_inteiro_teor": item.link_inteiro_teor,
                "prioridade": item.prioridade,
            }
        )

    return JsonResponse({"results": dados})


def _query_for_export(request: HttpRequest):
    document = _active_document()
    if not document:
        return TrackedProposition.objects.none()
    casas = _parse_multi(request.GET.get("casas", ""))
    secretarias = _parse_multi(request.GET.get("secretarias", ""))
    prioridades = _parse_multi(request.GET.get("prioridades", ""))
    busca = request.GET.get("busca", "").strip()
    queryset = document.propositions.all()
    return _filter_queryset(queryset, casas, secretarias, prioridades, busca)


def export_xlsx(request: HttpRequest) -> HttpResponse:
    queryset = _query_for_export(request)
    dados = []
    for item in queryset.order_by("-data_movimentacao", "-updated_at"):
        dados.append(
            {
                "Casa": item.get_casa_display(),
                "Secretaria": item.secretaria,
                "Proposição": item.titulo or item.assunto or "Título não disponível",
                "Autor": item.autor,
                "Ementa": item.ementa,
                "Status": item.status,
                "Última movimentação": item.ultima_movimentacao,
                "Data": item.data_movimentacao.strftime("%d/%m/%Y %H:%M") if item.data_movimentacao else "",
                "Ficha": item.link_ficha,
                "Inteiro Teor": item.link_inteiro_teor,
                "Prioridade": item.prioridade,
            }
        )

    df = pd.DataFrame(dados)
    buffer = io.BytesIO()
    with pd.ExcelWriter(buffer, engine="xlsxwriter") as writer:
        df.to_excel(writer, index=False, sheet_name="Proposicoes")
    buffer.seek(0)
    filename = timezone.now().strftime("proposicoes_monitoradas_%Y%m%d_%H%M%S.xlsx")
    response = HttpResponse(
        buffer.read(),
        content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )
    response["Content-Disposition"] = f"attachment; filename={filename}"
    return response


def export_docx(request: HttpRequest) -> HttpResponse:
    queryset = _query_for_export(request)
    documento = Document()
    documento.add_heading("Relatório de Proposições Monitoradas", level=1)

    itens = list(queryset.order_by("-data_movimentacao", "-updated_at"))
    for item in itens:
        paragrafo = documento.add_paragraph()
        titulo = item.titulo or item.assunto or f"ID {item.proposition_id}"
        paragrafo.add_run(titulo + " ").bold = True
        paragrafo.add_run(f"({item.get_casa_display()} - {item.secretaria or 'Sem secretaria'}) ")
        data_str = item.data_movimentacao.strftime("%d/%m/%Y %H:%M") if item.data_movimentacao else ""
        paragrafo.add_run(f"- {item.status} - {data_str}")
        if item.ementa:
            documento.add_paragraph(item.ementa)
        documento.add_paragraph(f"Autor: {item.autor or 'Autor não disponível'}")
        if item.link_ficha:
            documento.add_paragraph(f"Ficha: {item.link_ficha}")
        if item.link_inteiro_teor:
            documento.add_paragraph(f"Inteiro Teor: {item.link_inteiro_teor}")

    buffer = io.BytesIO()
    documento.save(buffer)
    buffer.seek(0)
    filename = timezone.now().strftime("proposicoes_monitoradas_%Y%m%d_%H%M%S.docx")
    response = HttpResponse(
        buffer.read(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    response["Content-Disposition"] = f"attachment; filename={filename}"
    return response


def configure_document(request: HttpRequest) -> HttpResponse:
    profile = _request_profile(request)
    if profile not in ALLOWED_CONFIG_PROFILES:
        messages.error(request, "Somente perfis administrador ou operador podem configurar o monitoramento.")
        index_url = _append_profile(_prefixed_path(request, reverse("monitoramento:index")), profile)
        return redirect(index_url)

    document = _active_document()
    if not document:
        document = TrackedDocument.objects.create(name="Proposições monitoradas")

    upload_form = UploadJsonForm()
    add_form = TrackedPropositionForm()

    if request.method == "POST":
        action = request.POST.get("action")
        if action == "upload":
            upload_form = UploadJsonForm(request.POST, request.FILES)
            if upload_form.is_valid():
                arquivo = upload_form.cleaned_data["arquivo"]
                try:
                    documento_importado = load_document_from_json(arquivo.read(), profile)
                except (ValueError, json.JSONDecodeError) as exc:
                    messages.error(request, f"Falha ao importar o JSON: {exc}")
                else:
                    sync_document(documento_importado)
                    messages.success(request, "Documento JSON importado com sucesso.")
                    return redirect(_append_profile(_prefixed_path(request, reverse("monitoramento:configure")), profile))
        elif action == "add":
            add_form = TrackedPropositionForm(request.POST)
            if add_form.is_valid():
                proposicao = add_form.save(commit=False)
                proposicao.document = document
                proposicao.save()
                document.last_updated_profile = profile
                document.raw_payload = build_document_payload(document)
                document.save(update_fields=["last_updated_profile", "raw_payload", "updated_at"])
                sync_document(document)
                messages.success(request, "Proposição adicionada à lista de monitoramento.")
                return redirect(_append_profile(_prefixed_path(request, reverse("monitoramento:configure")), profile))
        else:
            messages.error(request, "Ação não reconhecida.")

    proposicoes = document.propositions.order_by("casa", "proposition_id")

    context = {
        "document": document,
        "upload_form": upload_form,
        "add_form": add_form,
        "propositions": proposicoes,
        "profile": profile,
        "index_url": _append_profile(_prefixed_path(request, reverse("monitoramento:index")), profile),
        "download_url": _append_profile(_prefixed_path(request, reverse("monitoramento:download-document")), profile),
    }
    return render(request, "monitoramento/configure.html", context)


def download_document(request: HttpRequest) -> HttpResponse:
    profile = _request_profile(request)
    if profile not in ALLOWED_CONFIG_PROFILES:
        return redirect(_append_profile(_prefixed_path(request, reverse("monitoramento:index")), profile))

    document = _active_document()
    if not document:
        return HttpResponse("{}", content_type="application/json")

    payload = build_document_payload(document)
    filename = f"documento_monitoramento_{timezone.now().strftime('%Y%m%d_%H%M%S')}.json"
    data = json.dumps(payload, ensure_ascii=False, indent=2).encode("utf-8")
    response = HttpResponse(data, content_type="application/json")
    response["Content-Disposition"] = f"attachment; filename={filename}"
    return response


@require_POST
def delete_tracked(request: HttpRequest, pk: int) -> HttpResponse:
    profile = _request_profile(request)
    if profile not in ALLOWED_CONFIG_PROFILES:
        messages.error(request, "Sem permissão para remover proposições.")
        return redirect(_append_profile(_prefixed_path(request, reverse("monitoramento:index")), profile))

    proposition = get_object_or_404(TrackedProposition, pk=pk)
    document = proposition.document
    proposition.delete()
    document.last_updated_profile = profile
    document.raw_payload = build_document_payload(document)
    document.save(update_fields=["last_updated_profile", "raw_payload", "updated_at"])
    messages.success(request, "Proposição removida com sucesso.")
    return redirect(_append_profile(_prefixed_path(request, reverse("monitoramento:configure")), profile))
