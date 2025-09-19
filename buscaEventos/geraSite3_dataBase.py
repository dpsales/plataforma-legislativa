# ----------------------------------------------------------
# Relatório de Proposições Legislativas
# Lendo dados da base SQLite agendaSemana.db
# ----------------------------------------------------------
# Standard Library Imports
import os
import sqlite3
from pathlib import Path
from datetime import datetime

# Third-party Library Imports
import pandas as pd
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ----------------------------------------------------------
# 1) Conexão e leitura do banco de dados
# ----------------------------------------------------------
SCRIPT_DIR = Path(__file__).resolve().parent          # .../buscaEventos
DB_FILE    = (SCRIPT_DIR / "codUnificado" / "data" / "agendaSemana.db").resolve()

def get_conn() -> sqlite3.Connection:
    """Devolve uma conexão SQLite em modo autocommit."""
    return sqlite3.connect(DB_FILE, isolation_level=None)

def _parse_data_coluna(col: pd.Series) -> pd.Series:
    """
    Converte strings de data em datetime:
    1. tenta ISO  (YYYY-MM-DD)
    2. tenta dia/mês/ano (DD/MM/YYYY)
    """
    # Primeiro – ISO, que é o caso do banco
    parsed = pd.to_datetime(col, format="%Y-%m-%d", errors="coerce")
    # Depois – formatos com barra, assumindo dia/mês/ano
    mask = parsed.isna()
    if mask.any():
        parsed.loc[mask] = pd.to_datetime(col[mask], dayfirst=True, errors="coerce")
    return parsed

def carregar_dados() -> pd.DataFrame:
    """
    Lê a tabela 'eventos' e devolve DataFrame com
    a mesma estrutura usada no código original.
    """
    query = """
        SELECT
            proposicao           AS "PL",
            linkInteiroTeor      AS "Link PL",
            autorPartidoUf       AS "Autor/Partido",
            ementa               AS "Ementa",
            tipoImpactoFiscal    AS "Motivo da Seleção",
            impactoFiscal        AS "Incluir no REL?(X)",
            casa                 AS "Casa",
            plenarioOuComissao   AS "Plenário ou Comissao?",
            dataEvento           AS "Data",
            horaEvento           AS "Hora",
            nomeComissaoPlenario AS "Comissão/Plenário",
            linkComissaoPlenario AS "Link Comissão/Plenário"
        FROM eventos;
    """
    with get_conn() as con:
        df = pd.read_sql_query(query, con)

    # Normaliza campo de inclusão no relatório — 'S' ↔ 'X'
    df["Incluir no REL?(X)"] = (
        df["Incluir no REL?(X)"]
          .fillna("")
          .str.strip()
          .str.upper()
    )

    # Filtra apenas registros marcados com 'S'
    df = df[df["Incluir no REL?(X)"] == "S"].copy()

    # Converte coluna de data de forma robusta
    df["Data"] = _parse_data_coluna(df["Data"])

    return df

# ----------------------------------------------------------
# 2) Funções utilitárias de formatação (inalteradas)
# ----------------------------------------------------------
def add_hyperlink(paragraph, text, url):
    """
    Adiciona hiperlink azul e sublinhado a um parágrafo do Word.
    Se a URL for vazia/NaN, adiciona apenas o texto.
    """
    if pd.isna(url) or url == "":
        paragraph.add_run(text)
        return

    url = str(url)
    hyperlink = OxmlElement("w:hyperlink")
    r_id = paragraph._parent.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0000FF")
    rPr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rPr.append(underline)

    run.append(rPr)
    text_elem = OxmlElement("w:t")
    text_elem.text = text
    run.append(text_elem)
    hyperlink.append(run)
    paragraph._element.append(hyperlink)

def add_pl_info(document: Document, row: pd.Series):
    """
    Adiciona PL + Autor + Ementa + Impacto fiscal no documento.
    """
    # PL (com link)
    p = document.add_paragraph()
    add_hyperlink(p, row["PL"], row["Link PL"])
    p.style.font.name = "Arial"

    # Autor
    if pd.notna(row["Autor/Partido"]):
        author_paragraph = document.add_paragraph()
        run = author_paragraph.add_run("Autor: ")
        run.bold = True
        run.font.name = "Arial"
        author_paragraph.add_run(str(row["Autor/Partido"])).font.name = "Arial"

    # Ementa
    ementa_paragraph = document.add_paragraph()
    run = ementa_paragraph.add_run("Ementa: ")
    run.bold = True
    run.font.name = "Arial"
    ementa_paragraph.add_run(str(row["Ementa"])).font.name = "Arial"

    # Impacto Fiscal
    impacto = str(row["Motivo da Seleção"]) if pd.notna(row["Motivo da Seleção"]) else ""
    impacto_paragraph = document.add_paragraph()
    run = impacto_paragraph.add_run("Impacto fiscal: ")
    run.bold = True
    run.font.name = "Arial"
    impacto_paragraph.add_run(impacto).font.name = "Arial"

# ----------------------------------------------------------
# 3) Geração do Relatório Word
# ----------------------------------------------------------
def main() -> None:
    df = carregar_dados()

    dias_semana = {
        "Monday": "Segunda-feira",
        "Tuesday": "Terça-feira",
        "Wednesday": "Quarta-feira",
        "Thursday": "Quinta-feira",
        "Friday": "Sexta-feira",
        "Saturday": "Sábado",
        "Sunday": "Domingo",
    }
    df["Dia da Semana"] = (
        df["Data"].dt.strftime("%A").map(dias_semana)
    )

    df = df.sort_values(by=["Casa", "Plenário ou Comissao?", "Data", "Hora"])

    document = Document()
    heading = document.add_heading("Relatório para o Site", level=1)
    heading.style.font.name = "Arial"

    def add_section(title: str, data_frame: pd.DataFrame) -> None:
        if data_frame.empty:
            return

        sec_heading = document.add_heading(title, level=2)
        sec_heading.style.font.name = "Arial"

        for _, sub_group in data_frame.groupby("Plenário ou Comissao?"):
            for (dt, meeting), rows in sub_group.groupby(["Data", "Comissão/Plenário"]):
                dia_sem = dias_semana[dt.strftime("%A")]
                hora    = rows.iloc[0]["Hora"]

                date_heading = document.add_heading(
                    f"{dt.strftime('%d/%m/%Y')} - {dia_sem} - {hora}", level=3
                )
                date_heading.style.font.name = "Arial"

                p = document.add_paragraph()
                add_hyperlink(p, meeting, rows.iloc[0]["Link Comissão/Plenário"])
                p.style.font.name = "Arial"

                for _, row in rows.iterrows():
                    add_pl_info(document, row)

    add_section("Câmara dos Deputados", df[df["Casa"] == "CD"])
    add_section("Senado Federal",     df[df["Casa"] == "SF"])

    start_date = df["Data"].min().strftime("%d-%m-%Y")
    end_date   = df["Data"].max().strftime("%d-%m-%Y")
    output_file = f"Relatório para o Site - {start_date} a {end_date}.docx"
    document.save(output_file)
    print(f"\nRelatório gerado com sucesso: {output_file}")

# ----------------------------------------------------------
# 4) Execução
# ----------------------------------------------------------
if __name__ == "__main__":
    main()
