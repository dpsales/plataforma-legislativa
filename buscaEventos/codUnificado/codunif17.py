#aqui vamos começar a integração simultânea com o Banco de Dados
#começamos a integração do backupPl no banco também
# Standard Library Imports
import os
import io
import re
import time
from datetime import datetime, timedelta # Necessário para o DatePickerRange
import dash

# Third-party Library Imports
import requests
import pandas as pd

# Dash and Related Libraries
import dash_bootstrap_components as dbc
from dash import Dash, dcc, html, Input, Output, State, callback_context
import dash_ag_grid as dag # Importe dash_ag_grid como dag

# Dash Exceptions
from dash.exceptions import PreventUpdate

# Bibliotecas para gerar o relatório em Word
import docx
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.shared import RGBColor
from docx.oxml import OxmlElement
from docx.enum.section import WD_ORIENT
from docx.shared import Inches, Cm

# --- persistência -------------------------------------------------
import sqlite3
#parte nova
import ssl
import certifi
import requests
from requests.adapters import HTTPAdapter
from urllib3.util.retry import Retry
import logging

def _build_http_session():
    s = requests.Session()
    s.verify = certifi.where()  # usa o bundle do certifi
    retries = Retry(
        total=3,
        backoff_factor=0.8,
        status_forcelist=[429, 500, 502, 503, 504],
        allowed_methods={"GET"},
        raise_on_status=False,
    )
    adapter = HTTPAdapter(max_retries=retries)
    s.mount("https://", adapter)
    s.mount("http://", adapter)
    s.headers.update({"User-Agent": "agenda-semanal/1.0 (+contato@exemplo.gov.br)"})
    return s

HTTP = _build_http_session()
# --- Fallback TEMPORÁRIO para SSL do Senado --------------------------
from urllib.parse import urlparse

_SENADO_HOSTS_SEM_SSL = {"legis.senado.leg.br"}  # adicione aqui outros hosts, se precisar

def _http_get(url: str, **kwargs):
    """
    Wrapper para GET: mantém verificação TLS para tudo,
    mas DESLIGA (verify=False) SOMENTE para hosts do Senado
    com problema de cadeia (uso temporário).
    """
    parsed = urlparse(url)
    host = parsed.hostname or ""
    if host in _SENADO_HOSTS_SEM_SSL:
        # Loga claramente que está sem verificação para este host
        logging.warning(f"[TLS-FALLBACK] verify=False habilitado para {host} -> {url}")
        kwargs.setdefault("verify", False)
        return HTTP.get(url, **kwargs)
    # Demais domínios seguem fluxo normal (com verificação)
    return HTTP.get(url, **kwargs)
# --------------------------------------------------------------------
# logs úteis na inicialização
logging.info(f"OpenSSL: {ssl.OPENSSL_VERSION}")
logging.info(f"CA bundle (certifi): {certifi.where()}")


BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DATA_DIR = os.path.join(BASE_DIR, "data")
os.makedirs(DATA_DIR, exist_ok=True)
DB_FILE = os.path.join(DATA_DIR, "agendaSemana.db")

def get_conn():
    return sqlite3.connect(DB_FILE, isolation_level=None)

# Funções de banco e de geração da planilha
def init_db():
    with get_conn() as con:
        con.execute("""
        CREATE TABLE IF NOT EXISTS eventos (
            evento_id TEXT NOT NULL,
            nomeComissaoPlenario TEXT,
            dataEvento TEXT,
            horaEvento TEXT,
            linkComissaoPlenario TEXT,
            proposicao TEXT NOT NULL,
            ementa TEXT,
            autorPartidoUf TEXT,
            linkInteiroTeor TEXT,
            casa TEXT,
            temPL TEXT,
            plenarioOuComissao TEXT,
            impactoFiscal TEXT,
            marcarParaRelatorio TEXT,
            tipoImpactoFiscal TEXT,
            PRIMARY KEY (evento_id, proposicao)
        );
        """)
        con.execute("""
        CREATE TABLE IF NOT EXISTS backup_pl (
            proposicao TEXT PRIMARY KEY,
            impactoFiscal TEXT,
            tipoImpactoFiscal TEXT,
            linkInteiroTeor TEXT,
            dataGeracaoPlanilha TEXT
        );
        """)

init_db()

with get_conn() as con:
    cur = con.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='eventos'")
    assert cur.fetchone(), "⚠️  Tabela 'eventos' não foi criada!"

# -----------------------------------------------------------
# FUNÇÃO PARA GERAR PLANILHA (BUSCA DADOS E ATUALIZA DB)
# -----------------------------------------------------------
def gerar_planilha(data_inicio: str, data_fim: str) -> pd.DataFrame:
    cols_esperadas = [
            "evento_id", "nomeComissaoPlenario", "dataEvento", "horaEvento",
            "linkComissaoPlenario","proposicao","ementa","linkInteiroTeor",
            "casa","temPL","plenarioOuComissao","impactoFiscal",
            "marcarParaRelatorio","tipoImpactoFiscal","autorPartidoUf",
            "dataGeracaoPlanilha","textosAssociados","buscaPalavrasChave"
        ]
    print(f"[LOG] Iniciando geração da planilha de {data_inicio} a {data_fim}")
    start_time = time.time()

    ev_cam       = buscar_eventos_camara(data_inicio, data_fim)
    ev_sen       = buscar_eventos_senado_json(data_inicio, data_fim)
    
    regs_cam     = formatar_dados_camara(ev_cam, data_inicio, data_fim)
    regs_sen_com = formatar_dados_comissoes_senado(
        [e for e in ev_sen if 'Materias' not in e], 
        data_inicio, data_fim
    )
    regs_sen_pl  = formatar_dados_plenario_senado(
        [e for e in ev_sen if 'Materias' in e], 
        data_inicio, data_fim
    )
    
    todos_regs = regs_cam + regs_sen_com + regs_sen_pl

    if not todos_regs:
        print("[LOG] Sem registros novos encontrados para o período.")
        with get_conn() as con:
            con.execute("DELETE FROM eventos") 
            print("[LOG] Banco de dados 'eventos' limpo pois não foram encontrados novos registros.")
        return pd.DataFrame(columns=cols_esperadas) 

    df_novos_eventos = pd.DataFrame(todos_regs)
    df_com_backup = importar_dados_do_backup(df_novos_eventos)

    for col in cols_esperadas:
        if col not in df_com_backup.columns:
            df_com_backup[col] = None 

    df_db = df_com_backup.copy()
    df_db["dataEvento"] = (
        pd.to_datetime(df_db["dataEvento"], format="%d/%m/%Y", errors="coerce")
          .dt.strftime("%Y-%m-%d")
    )
    df_db = df_db.where(pd.notnull(df_db), None) 

    with get_conn() as con:
        print("[LOG] Limpando dados antigos da tabela 'eventos' no SQLite.")
        con.execute("DELETE FROM eventos")
        
        print(f"[LOG] Tentando inserir/atualizar {len(df_db)} registros no SQLite.")
        for _, row in df_db.iterrows():
            insert_values = (
                row.get("evento_id"), row.get("proposicao"),
                row.get("nomeComissaoPlenario"), row.get("dataEvento"), row.get("horaEvento"),
                row.get("linkComissaoPlenario"), row.get("ementa"), row.get("autorPartidoUf"),
                row.get("linkInteiroTeor"), row.get("casa"), row.get("temPL"),
                row.get("plenarioOuComissao"), row.get("impactoFiscal"), row.get("marcarParaRelatorio"), 
                row.get("tipoImpactoFiscal")
            )
            update_values = (
                row.get("nomeComissaoPlenario"), row.get("dataEvento"), row.get("horaEvento"),
                row.get("linkComissaoPlenario"), row.get("ementa"), row.get("autorPartidoUf"),
                row.get("linkInteiroTeor"), row.get("casa"), row.get("temPL"),
                row.get("plenarioOuComissao"), row.get("impactoFiscal"), row.get("marcarParaRelatorio"),
                row.get("tipoImpactoFiscal"),
                row.get("evento_id"), row.get("proposicao") 
            )
            try:
                con.execute("""
                    INSERT INTO eventos (
                        evento_id, proposicao,
                        nomeComissaoPlenario, dataEvento, horaEvento,
                        linkComissaoPlenario, ementa, autorPartidoUf, linkInteiroTeor,
                        casa, temPL, plenarioOuComissao, impactoFiscal, marcarParaRelatorio, tipoImpactoFiscal
                    ) VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)
                """, insert_values)
            except sqlite3.IntegrityError: 
                con.execute("""
                    UPDATE eventos SET
                        nomeComissaoPlenario  = ?, dataEvento = ?, horaEvento = ?,
                        linkComissaoPlenario  = ?, ementa = ?, autorPartidoUf = ?,
                        linkInteiroTeor       = ?, casa = ?, temPL = ?,
                        plenarioOuComissao    = ?, impactoFiscal = ?, marcarParaRelatorio = ?,
                        tipoImpactoFiscal     = ?
                    WHERE evento_id = ? AND proposicao = ?
                """, update_values)
            except Exception as e_insert:
                print(f"[LOG ERROR DB] Erro ao inserir/atualizar linha: {row.to_dict()}. Erro: {e_insert}")

    total_db, = con.execute("SELECT COUNT(*) FROM eventos").fetchone()
    print(f"[LOG] 🚀 {total_db} registros no SQLite após atualização.")

    df_disp = carregar_dados_formatados_do_sqlite() 

    print(f"[LOG] Geração da planilha e atualização do DB concluída em {(time.time()-start_time)/60:.2f} minutos.")
    return df_disp


# =======================================================================
# Funções de busca e formatação de dados (Câmara e Senado)
# =======================================================================
TIPOS_EVENTOS_CAMARA = [
    "Sessão Deliberativa", "Reunião Deliberativa", "Audiência Pública",
    "Audiência Pública e Deliberação", "Reunião de Instalação e Eleição",
    "Reunião de Instalação","Reunião de Eleição","Reunião","Tomada de Depoimento e Deliberação"
]
KEYWORDS = ["Ministério do Planejamento","MPO","Secretaria de Orçamento","Simone Tebet"]
COMISSOES_IMPORTANTES = [
    "Comissão de Constituição, Justiça e Cidadania",
    "Plenário do Senado Federal",
    "Comissão de Assuntos Econômicos",
    "Plenário da Câmara dos Deputados",
    "Comissão de Finanças e Tributação",
    "Comissão de Constituição e Justiça e de Cidadania",
    "Comissão Mista de Planos, Orçamentos Públicos e Fiscalização",
    "Plenário" 
]

BACKUP_TABLE = "backup_pl" 

TIPOS_PERMITIDOS = {'PDL', 'PEC', 'PL', 'PLP', 'PRC', 'MPV', 'PLN', 'SUG', 'PLS'} 

def buscar_eventos_camara(data_inicio, data_fim):
    print(f"[LOG] Buscando eventos da Câmara de {data_inicio} a {data_fim}")
    url = "https://dadosabertos.camara.leg.br/api/v2/eventos"
    params = {
        'dataInicio': data_inicio,
        'dataFim': data_fim,
        'itens': 100,
        'ordem': 'asc',
        'ordenarPor': 'dataHoraInicio'
    }
    try:
        resp = HTTP.get(url, params=params, timeout=30)
        resp.raise_for_status()
        dados = resp.json().get('dados', [])
        eventos_filtrados = [e for e in dados if e.get('descricaoTipo') in TIPOS_EVENTOS_CAMARA]
        print(f"[LOG] Encontrados {len(dados)} eventos na Câmara, {len(eventos_filtrados)} após filtro de tipo.")
        return eventos_filtrados
    except requests.exceptions.RequestException as e:
        print(f"[LOG] Erro ao buscar Câmara: {e}")
        return []

def obter_pauta_camara(evento_id, max_attempts=3, delay=2):
    url = f"https://dadosabertos.camara.leg.br/api/v2/eventos/{evento_id}/pauta"
    for attempt in range(1, max_attempts + 1):
        try:
            resp = HTTP.get(url, timeout=30)
            resp.raise_for_status()
            dados = resp.json().get('dados', [])
            return dados
        except requests.exceptions.RequestException as e:
            print(f"[LOG] Erro pauta Câmara {evento_id} (tentativa {attempt}): {e}")
            if attempt < max_attempts:
                time.sleep(delay)
    return []

def obter_ementa_proposicao(proposicao_id, max_attempts=3, delay=2):
    url = f"https://dadosabertos.camara.leg.br/api/v2/proposicoes/{proposicao_id}"
    attempts = 0
    while attempts < max_attempts:
        try:
            response = HTTP.get(url, timeout=20)
            response.raise_for_status()
            if response.content:
                dados = response.json().get('dados', {})
                ementa = dados.get('ementa', '')
                url_inteiro_teor = dados.get('urlInteiroTeor', '')
                sigla_tipo_prop = dados.get('siglaTipo', '')
                id_para_autores = proposicao_id

                if sigla_tipo_prop not in ['REQ', 'RIC', 'INC', 'RQS'] and dados.get('proposicaoRelacionada_'):
                    proposta_rel = dados.get('proposicaoRelacionada_')
                    if isinstance(proposta_rel, dict) and proposta_rel.get('id'):
                        related_id = proposta_rel.get('id')
                        url_relacionada = f"https://dadosabertos.camara.leg.br/api/v2/proposicoes/{related_id}"
                        response_relacionada = HTTP.get(url_relacionada, timeout=10)
                        response_relacionada.raise_for_status()
                        if response_relacionada.content:
                            dados_relacionada = response_relacionada.json().get('dados', {})
                            ementa = dados_relacionada.get('ementa', ementa)
                            url_inteiro_teor = dados_relacionada.get('urlInteiroTeor', url_inteiro_teor)
                            id_para_autores = related_id

                autores_str = ''
                if id_para_autores:
                    autores_url = f"https://dadosabertos.camara.leg.br/api/v2/proposicoes/{id_para_autores}/autores"
                    autores_response = HTTP.get(autores_url, timeout=10)
                    autores_response.raise_for_status()
                    if autores_response.content:
                        autores_dados = autores_response.json().get('dados', [])
                        autores_dados = sorted(autores_dados, key=lambda a: a.get('ordemAssinatura', float('inf')))
                        autores_list = []
                        for autor in autores_dados:
                            autor_nome = autor.get('nome', '')
                            uri_autor = autor.get('uri', '')
                            tipo_autor = autor.get('tipo', '')
                            partido, uf = '', ''
                            if tipo_autor == 'Deputado(a)':
                                partido, uf = obter_partido_uf_deputado(uri_autor)
                            elif tipo_autor == 'Senador(a)':
                                partido, uf = obter_partido_uf_senador(uri_autor)
                            autor_partido_uf = autor_nome
                            if partido and uf:
                                autor_partido_uf = f"{autor_nome} ({partido}/{uf})"
                            elif partido:
                                autor_partido_uf = f"{autor_nome} ({partido})"
                            autores_list.append(autor_partido_uf)

                        if autores_list:
                            autores_str = (autores_list[0] + " e outros") if len(autores_list) > 1 else autores_list[0]
                        else:
                            autores_str = ''

                return ementa, url_inteiro_teor, autores_str
            return '', '', ''
        except requests.exceptions.HTTPError as e:
            status_code = e.response.status_code
            attempts += 1
            if status_code in [403, 404]:
                break
            if attempts < max_attempts:
                time.sleep(delay)
        except (requests.exceptions.RequestException, ValueError):
            attempts += 1
            if attempts < max_attempts:
                time.sleep(delay)
    return '', '', ''

def obter_partido_uf_parlamentar(uri_parlamentar, tipo_parlamentar="parlamentar"):
    attempts = 0
    max_attempts = 3
    delay = 1
    if not uri_parlamentar:
        return '', ''
    while attempts < max_attempts:
        try:
            response = HTTP.get(uri_parlamentar, timeout=10)
            response.raise_for_status()
            if response.content:
                dados = response.json().get('dados', {}).get('ultimoStatus', {})
                return dados.get('siglaPartido', ''), dados.get('siglaUf', '')
            return '', ''
        except requests.exceptions.HTTPError as e:
            status_code = e.response.status_code
            attempts += 1
            if status_code in [403, 404]:
                break
            if status_code == 429:
                time.sleep(delay * 5)
            else:
                time.sleep(delay)
        except (requests.exceptions.RequestException, ValueError):
            attempts += 1
            if attempts < max_attempts:
                time.sleep(delay)
    return '', ''

def obter_partido_uf_deputado(uri_deputado):
    return obter_partido_uf_parlamentar(uri_deputado, "deputado")

def obter_partido_uf_senador(uri_senador):
    return obter_partido_uf_parlamentar(uri_senador, "senador")


def buscar_eventos_senado_json(data_inicio, data_fim):
    print(f"[LOG] Buscando eventos do Senado de {data_inicio} a {data_fim}")
    try:
        dt0 = datetime.strptime(data_inicio, "%Y-%m-%d")
        dt1 = datetime.strptime(data_fim, "%Y-%m-%d")
    except Exception as e:
        print(f"[LOG] Erro datas Senado: {e}")
        return []

    eventos = []
    for i in range((dt1 - dt0).days + 1):
        dia = dt0 + timedelta(days=i)
        s = dia.strftime("%Y%m%d")

        urls = {
            "comissoes": f"https://legis.senado.leg.br/dadosabertos/agendareuniao/{s}.json",
            "plenario": f"https://legis.senado.leg.br/dadosabertos/plenario/agenda/dia/{s}.json",
        }

        for tipo, url in urls.items():
            try:
                resp = _http_get(url, timeout=20)
                resp.raise_for_status()
                j = resp.json()

                if tipo == "comissoes":
                    reunioes_obj = j.get("AgendaReuniao", {}).get("reunioes")
                    if reunioes_obj is None:
                        reunioes = []
                    elif isinstance(reunioes_obj, dict):
                        reunioes = reunioes_obj.get("reuniao", []) or []
                    else:
                        reunioes = reunioes_obj
                    print(f"[LOG] {len(reunioes):2d} comissões em {s}")
                    eventos.extend(reunioes)

                else:  # plenario
                    sessoes = (
                        j.get("AgendaPlenario", {})
                         .get("Sessoes", {})
                         .get("Sessao", [])
                    )
                    print(f"[LOG] {len(sessoes):2d} plenário  em {s}")
                    eventos.extend(sessoes)

            except Exception as e:
                print(f"[LOG] Erro Senado {tipo} {s}: {e}")

    print(f"[LOG] Total Senado: {len(eventos)} eventos")
    return eventos

def _as_list(x):
    if x is None:
        return []
    return x if isinstance(x, list) else [x]

def obter_pauta_senado_json(evento):
    """
    Extrai a pauta tanto de Plenário (Materias) quanto de Comissões (partes/itens),
    aceitando dict ou list em todos os níveis.
    """
    if not isinstance(evento, dict):
        print(f"Evento inválido (não é um dict). Ignorando: {evento}")
        return []

    try:
        # Caso Plenário
        if 'Materias' in evento:
            materias_obj = evento.get('Materias')
            if isinstance(materias_obj, dict):
                materias = materias_obj.get('Materia', [])
                return materias if isinstance(materias, list) else _as_list(materias)
            # fallback (já vindo como lista)
            return materias_obj if isinstance(materias_obj, list) else []

        # Caso Comissões
        pauta = []
        partes = _as_list(evento.get('partes'))
        for parte in partes:
            if not isinstance(parte, dict):
                continue
            itens = _as_list(parte.get('itens'))
            for item in itens:
                # alguns dumps trazem item vazio/None
                if isinstance(item, dict) and item:
                    pauta.append(item)

        # Log leve de diagnóstico
        cod = evento.get('codigo') or evento.get('CodigoSessao')
        print(f"[LOG] itens extraídos para reunião/sessão {cod}: {len(pauta)}")
        return pauta

    except Exception as e:
        print(f"Erro ao extrair a pauta do evento ID {evento.get('CodigoSessao') or evento.get('codigo')}: {e}")
        return []

def obter_ementa_proposicao_senado(evento):
    p = evento.get('proposicao',{})
    e = p.get('ementa','Sem ementa disponível')
    u = p.get('urlInteiroTeor','N/A')
    autores = [a.get('nome','Desconhecido') for a in p.get('autores',[])]
    if autores:
        autor_str = autores[0] + " e outros" if len(autores) > 1 else autores[0]
    else:
        autor_str = ''
    return e, u, autor_str  

def formatar_dados_comissoes_senado(eventos_comissoes, data_inicio: str, data_fim: str):
    registros = []
    for ev in eventos_comissoes:
        d0 = ev.get('dataInicio', '')[:10]
        if (data_inicio <= d0 <= data_fim):
            orgao = ev.get('colegiadoCriador', {}).get('nome', 'N/A')
            try:
                dt_str = datetime.strptime(d0, '%Y-%m-%d').strftime('%d/%m/%Y')
            except:
                dt_str = d0
            hr = ev.get('dataInicio', '')[11:16]
            descr = ev.get('tipo', {}).get('descricao', 'Desconhecido')
            sit = ev.get('situacao', 'Desconhecida')

            pauta = obter_pauta_senado_json(ev)
            for item in pauta:
                if not isinstance(item, dict):
                    continue

                doma = item.get('doma', {})
                ident = doma.get('identificacao', 'Sem identificação')
                tipo_item = ident.split()[0] if ident else ""
                if tipo_item not in TIPOS_PERMITIDOS:
                    continue
                ementa = doma.get('ementa', 'Sem ementa disponível')

                # tratamento robusto de 'textos', que pode ser lista ou dict
                textos = doma.get('textos')
                if isinstance(textos, list) and textos:
                    primeiro_texto = textos[0]
                elif isinstance(textos, dict):
                    primeiro_texto = textos
                else:
                    primeiro_texto = {}
                url_inteiro_teor = primeiro_texto.get('urlDownload', 'N/A')

                autoria = doma.get('autoria', 'N/A')

                tem_pl = 'S' if any(ident.startswith(pref) for pref in ['PL ', 'PLP ', 'PLS ', 'PEC ']) else 'N'
                imp = 'S' if any(nome in f"{orgao} - {descr} - {sit}" for nome in COMISSOES_IMPORTANTES) else 'N'

                reg = {
                    "evento_id": ev.get("codigo"),
                    'nomeComissaoPlenario': f"{orgao} - {descr} - {sit}",
                    'dataEvento': dt_str,
                    'horaEvento': hr,
                    'linkComissaoPlenario': f"https://legis.senado.leg.br/comissoes/reuniao?reuniao={ev.get('codigo')}",
                    'proposicao': ident,
                    'ementa': ementa,
                    'linkInteiroTeor': url_inteiro_teor,
                    'casa': 'SF',
                    'temPL': tem_pl,
                    'plenarioOuComissao': 'CM',
                    'impactoFiscal': '',
                    'marcarParaRelatorio': imp,
                    'tipoImpactoFiscal': '',
                    'autorPartidoUf': autoria,
                    'dataGeracaoPlanilha': datetime.now().strftime('%d/%m/%Y'),
                    'textosAssociados': url_inteiro_teor,
                    'buscaPalavrasChave': 'ND'
                }
                registros.append(reg)

    print(f"[LOG] formatar_dados_comissoes_senado -> {len(registros)} registros")
    return registros

def formatar_dados_plenario_senado(eventos_plenario, data_inicio: str, data_fim: str):
    registros = []
    for sess in eventos_plenario:
        d0 = sess.get('Data', '')
        if not (data_inicio <= d0 <= data_fim):
            continue

        # data e hora
        try:
            dt_str = datetime.strptime(d0, '%Y-%m-%d').strftime('%d/%m/%Y')
        except:
            dt_str = d0
        hr    = sess.get('Hora', '')
        tipo_s = sess.get('TipoSessao', 'Desconhecido')
        sit_s  = sess.get('SituacaoSessao', 'Desconhecida')
        casa   = sess.get('Casa', 'SF')

        # percorre cada matéria na pauta
        for mat in sess.get('Materias', {}).get('Materia', []):
            cod  = mat.get('CodigoMateria', 'N/A')
            # identificação: sigla + número/ano
            sigla = mat.get('SiglaMateria', '').strip()
            numero = mat.get('NumeroMateria', '')
            ano    = mat.get('AnoMateria', '')
            ident = f"{sigla} {numero}/{ano}".strip()

            # 🔹 filtro de TIPOS_PERMITIDOS
            if sigla not in TIPOS_PERMITIDOS:
                continue

            autoria = mat.get('NomeAutor', 'N/A')
            imp     = 'S' if any(
                          nome in f"Plenário - {tipo_s} - {sit_s}"
                          for nome in COMISSOES_IMPORTANTES
                      ) else 'N'

            reg = {
                'evento_id': sess.get("CodigoSessao"),
                'nomeComissaoPlenario': f"Plenário - {tipo_s} - {sit_s}",
                'dataEvento': dt_str,
                'horaEvento': hr,
                'linkComissaoPlenario': (
                    f"https://www25.senado.leg.br/web/atividade/"
                    f"sessao-plenaria/-/pauta/{sess.get('CodigoSessao')}"
                ),
                # usa ident aqui para garantir consistência
                'proposicao': f"{ident} - {autoria}",
                'ementa': mat.get('Ementa', 'Sem ementa disponível'),
                'linkInteiroTeor': (
                    f"https://www25.senado.leg.br/web/atividade/"
                    f"materias/-/materia/{cod}"
                    if cod != 'N/A' else 'N/A'
                ),
                'casa': casa,
                'temPL': 'S' if cod != 'N/A' else 'N',
                'plenarioOuComissao': 'PLEN',
                'impactoFiscal': '',
                'marcarParaRelatorio': imp,
                'tipoImpactoFiscal': '',
                'autorPartidoUf': autoria,
                'dataGeracaoPlanilha': datetime.now().strftime('%d/%m/%Y'),
                'textosAssociados': "",
                'buscaPalavrasChave': 'ND'
            }
            registros.append(reg)

    print(f"[LOG] formatar_dados_plenario_senado -> {len(registros)} registros")
    return registros

def formatar_dados_camara(eventos_camara, data_inicio_str: str, data_fim_str: str):
    registros = []
    for ev_camara in eventos_camara: 
        data_hora_inicio_api = ev_camara.get("dataHoraInicio") 
        if not data_hora_inicio_api: continue
        try:
            dt_obj = datetime.fromisoformat(data_hora_inicio_api)
            data_evento_fmt_db = dt_obj.strftime('%Y-%m-%d')
            if not (data_inicio_str <= data_evento_fmt_db <= data_fim_str): continue
            dt_str_display = dt_obj.strftime("%d/%m/%Y")
            hr_str_display = dt_obj.strftime("%H:%M")
        except ValueError: continue
        orgaos_info = ev_camara.get("orgaos", [])
        nome_orgao_cam = orgaos_info[0].get("nome", "Órgão Desconhecido") if orgaos_info else "Órgão Desconhecido"
        sigla_orgao_cam = orgaos_info[0].get("sigla", "") if orgaos_info else ""
        nome_display_orgao = f"{nome_orgao_cam} ({sigla_orgao_cam})" if sigla_orgao_cam else nome_orgao_cam
        tipo_evento_desc = ev_camara.get('descricaoTipo', 'Evento Desconhecido')
        situacao_evento_desc = ev_camara.get('situacao', 'Situação Desconhecida')
        id_evento_api = ev_camara.get("id")
        link_evento_oficial = f"https://www.camara.leg.br/evento-legislativo/{id_evento_api}" if id_evento_api else "N/A"
        nome_evento_completo = f"{nome_display_orgao} - {tipo_evento_desc} - {situacao_evento_desc}"
        pauta_itens = obter_pauta_camara(id_evento_api)
        if not pauta_itens: continue
        for item_pauta in pauta_itens:
            if not isinstance(item_pauta, dict): continue
            prop_principal_info = item_pauta.get("proposicao_", {})
            id_prop_principal_api = prop_principal_info.get("id")
            sigla_prop_principal_api = prop_principal_info.get("siglaTipo", "")
            num_prop_principal_api = prop_principal_info.get("numero", "")
            ano_prop_principal_api = prop_principal_info.get("ano", "")
            titulo_item_pauta = item_pauta.get("titulo", "") 
            id_para_detalhes_prop, sigla_para_detalhes_prop, num_para_detalhes_prop, ano_para_detalhes_prop = \
                id_prop_principal_api, sigla_prop_principal_api, num_prop_principal_api, ano_prop_principal_api
            prop_relacionada_info = item_pauta.get("proposicaoRelacionada_", {})
            if isinstance(prop_relacionada_info, dict) and prop_relacionada_info and sigla_prop_principal_api not in ("REQ", "RIC", "INC", "RQS"):
                id_prop_rel_api = prop_relacionada_info.get("id")
                if id_prop_rel_api: 
                    id_para_detalhes_prop = id_prop_rel_api
                    sigla_para_detalhes_prop = prop_relacionada_info.get("siglaTipo", sigla_prop_principal_api)
                    num_para_detalhes_prop = prop_relacionada_info.get("numero", num_prop_principal_api)
                    ano_para_detalhes_prop = prop_relacionada_info.get("ano", ano_prop_principal_api)
            ident_prop_display = f"{sigla_para_detalhes_prop} {num_para_detalhes_prop}/{ano_para_detalhes_prop}".strip()
            if not sigla_para_detalhes_prop or not num_para_detalhes_prop or not ano_para_detalhes_prop : 
                 ident_prop_display = titulo_item_pauta if titulo_item_pauta else "Proposição N/A"
            if sigla_para_detalhes_prop not in TIPOS_PERMITIDOS and sigla_para_detalhes_prop: continue
            ementa_prop, link_prop, autores_prop = "", "", ""
            if id_para_detalhes_prop: 
                ementa_prop, link_prop, autores_prop = obter_ementa_proposicao(id_para_detalhes_prop)
            else: ementa_prop = prop_principal_info.get("ementa", "Sem ementa disponível")
            busca_palavra_chave = "ND"
            if ementa_prop:
                for kw in KEYWORDS:
                    if kw.lower() in ementa_prop.lower():
                        busca_palavra_chave = "Ementa: " + kw; break
            tem_pl_flag = "S" if sigla_para_detalhes_prop in ["PL", "PLP", "PLS", "PEC", "MPV", "PLN"] else "N"
            marcar_rel_flag = "S" if any(com_imp in nome_evento_completo for com_imp in COMISSOES_IMPORTANTES) else "N"
            plen_ou_comissao_flag = "PLEN" if "Plenário" in nome_display_orgao else "CM"
            registros.append({
                "evento_id": f"CD_EVT_{id_evento_api}" if id_evento_api else f"CD_EVT_DATA_{data_evento_fmt_db}",
                "nomeComissaoPlenario": nome_evento_completo, "dataEvento": dt_str_display, "horaEvento": hr_str_display,
                "linkComissaoPlenario": link_evento_oficial, "proposicao": ident_prop_display, 
                "ementa": ementa_prop if ementa_prop else "Sem ementa disponível",
                "linkInteiroTeor": link_prop if link_prop else "N/A", "casa": "CD", "temPL": tem_pl_flag,
                "plenarioOuComissao": plen_ou_comissao_flag, "impactoFiscal": None, "marcarParaRelatorio": marcar_rel_flag,
                "tipoImpactoFiscal": None, "autorPartidoUf": autores_prop if autores_prop else "N/A",
                "dataGeracaoPlanilha": datetime.now().strftime("%d/%m/%Y"),
                "textosAssociados": link_prop if link_prop else "", "buscaPalavrasChave": busca_palavra_chave })
    if registros: print(f"[LOG] formatar_dados (Câmara) -> {len(registros)} registros.")
    return registros

def importar_dados_do_backup(df_novos_eventos: pd.DataFrame) -> pd.DataFrame:
    # garante a coluna proposicao
    if "proposicao" not in df_novos_eventos.columns:
        df_novos_eventos["proposicao"] = ""
    
    # carregando backup do banco
    with get_conn() as con:
        df_backup = pd.read_sql(f"SELECT proposicao, impactoFiscal, tipoImpactoFiscal FROM {BACKUP_TABLE}", con)
    print(f"[LOG] Backup da tabela '{BACKUP_TABLE}' carregado com {len(df_backup)} registros.")
    
    # prepara merge
    def extrair_codigo_pl(pl_str):
        if pd.isna(pl_str): return None
        m = re.match(r"^([A-Z]+\s*\d+/\d+)", str(pl_str).strip())
        return m.group(1) if m else str(pl_str).strip()
    
    df_novos_eventos["PL_Code_Merge"] = df_novos_eventos["proposicao"].apply(extrair_codigo_pl)
    df_backup["PL_Code_Merge"] = df_backup["proposicao"].apply(extrair_codigo_pl)
    df_backup = df_backup.drop_duplicates(subset=["PL_Code_Merge"], keep="first")
    
    cols_para_merge = ["PL_Code_Merge", "impactoFiscal", "tipoImpactoFiscal"]
    df_merged = df_novos_eventos.merge(
        df_backup[cols_para_merge],
        on="PL_Code_Merge",
        how="left",
        suffixes=("", "_bk")
    )
    
    # preenche com backup onde houver
    for col in ["impactoFiscal", "tipoImpactoFiscal"]:
        bk = col + "_bk"
        if bk in df_merged.columns:
            # só aplica se o backup não estiver em branco
            mask = (
                df_merged[bk].notna() 
                & df_merged[bk].astype(str).str.strip().astype(bool)
            )
            df_merged.loc[mask, col] = df_merged.loc[mask, bk]
            df_merged.drop(columns=[bk], inplace=True)
        else:
            df_merged[col] = df_merged.get(col, None)
    
    # remove coluna auxiliar
    df_merged.drop(columns=["PL_Code_Merge"], inplace=True, errors="ignore")
    print(f"[LOG] importar_dados_do_backup -> {len(df_merged)} registros após merge com backup.")
    
    # agora, atualiza o backup_pl no SQLite (upsert)
    with get_conn() as con:
        for _, row in df_merged[
            ["proposicao", "impactoFiscal", "tipoImpactoFiscal", "linkInteiroTeor", "dataGeracaoPlanilha"]
        ].drop_duplicates(subset=["proposicao"]).iterrows():
            con.execute(f"""
            INSERT INTO {BACKUP_TABLE}
                (proposicao, impactoFiscal, tipoImpactoFiscal, linkInteiroTeor, dataGeracaoPlanilha)
            VALUES (?,?,?,?,?)
            ON CONFLICT(proposicao) DO UPDATE SET
                impactoFiscal=excluded.impactoFiscal,
                tipoImpactoFiscal=excluded.tipoImpactoFiscal,
                linkInteiroTeor=excluded.linkInteiroTeor,
                dataGeracaoPlanilha=excluded.dataGeracaoPlanilha;
            """, (
                row["proposicao"], row["impactoFiscal"],
                row["tipoImpactoFiscal"], row["linkInteiroTeor"],
                row["dataGeracaoPlanilha"]
            ))
    print(f"[LOG] Tabela '{BACKUP_TABLE}' sincronizada com {len(df_merged)} PLs.")
    
    return df_merged

def carregar_dados_formatados_do_sqlite() -> pd.DataFrame: 
    with get_conn() as con:
        df = pd.read_sql("SELECT * FROM eventos", con)
    if df.empty:
        print("[LOG] carregar_dados_formatados_do_sqlite: Banco de dados vazio.")
        return pd.DataFrame() 
    for col in ['impactoFiscal', 'tipoImpactoFiscal', 'marcarParaRelatorio']:
        if col not in df.columns: df[col] = '' 
        else: df[col] = df[col].fillna('') 
    df["proposicao_raw"] = df["proposicao"] 
    df["dataEvento"] = pd.to_datetime(df["dataEvento"], errors='coerce').dt.strftime("%d/%m/%Y").fillna("N/A")
    df["horaEvento"] = df["horaEvento"].astype(str).fillna("N/A")
    df["dataHora"] = df.apply(lambda r: f"{r['dataEvento']} {r['horaEvento']}" if r['dataEvento'] != "N/A" and r['horaEvento'] != "N/A" else "N/A", axis=1)
    def _format_link_for_display(txt, url):
        txt_str = str(txt) if pd.notna(txt) else "N/A"
        url_str = str(url) if pd.notna(url) and str(url).strip() != "N/A" and str(url).startswith("http") else ""
        return f"[{txt_str}]({url_str})" if url_str else txt_str
    df["nomeComissaoPlenario"] = df.apply(lambda r: _format_link_for_display(r.get("nomeComissaoPlenario"), r.get("linkComissaoPlenario")), axis=1)
    df["proposicao"] = df.apply(lambda r: _format_link_for_display(r.get("proposicao_raw"), r.get("linkInteiroTeor")), axis=1)
    print(f"[LOG] carregar_dados_formatados_do_sqlite: {len(df)} registros formatados.")
    return df

# -------------------------------------------------
# Dash App com ag-Grid
# -------------------------------------------------
impacto_editor = {"cellEditor": "agSelectCellEditor", "cellEditorParams": {"values": ["S", "N", ""]}}
tipo_editor = {
    "cellEditor": "agTextCellEditor"   # texto livre
}
#editor com parâmetros 
#tipo_editor = {"cellEditor": "agSelectCellEditor", "cellEditorParams": {"values": ["Aumento de Despesa", "Redução de Receita", "Sem Impacto", "Análise Pendente", ""]}} 
marcar_relatorio_editor = {"cellEditor": "agSelectCellEditor", "cellEditorParams": {"values": ["S", "N", ""]}}

date_filter_comparator_js = """
function(filterLocalDateAtMidnight, cellValue) {
    if (cellValue == null || cellValue.trim() === '' || cellValue.trim() === 'N/A') { return 0; }
    const datePartMatch = cellValue.match(/^(\\d{2}\\/\\d{2}\\/\\d{4})/);
    if (!datePartMatch) { return 0; }
    const dateParts = datePartMatch[0].split("/");
    const cellDate = new Date(Number(dateParts[2]), Number(dateParts[1]) - 1, Number(dateParts[0]));
    if (filterLocalDateAtMidnight.getTime() === cellDate.getTime()) { return 0; }
    if (cellDate < filterLocalDateAtMidnight) { return -1; }
    if (cellDate > filterLocalDateAtMidnight) { return 1; }
    return 0; 
}
"""
columnDefs = [
    {"headerName": "Casa", "field": "casa", "width": 80, "editable": False, "filter": True, "floatingFilter": True},
    {"headerName": "Comissão/Plenário", "field": "nomeComissaoPlenario", "cellRenderer": "markdownLinkRenderer", "width": 250, "wrapText": True, "autoHeight": True, "editable": False, "cellStyle": {"white-space": "normal", "lineHeight": "1.2"}, "filter": "agTextColumnFilter", "floatingFilter": True},
    {"headerName": "Data/Hora", "field": "dataHora", "width": 150, "editable": False, "filter": "agDateColumnFilter", "floatingFilter": True, "filterParams": {"comparator": date_filter_comparator_js }},
    {"headerName": "Proposição", "field": "proposicao", "cellRenderer": "markdownLinkRenderer", "width": 200, "autoHeight": True, "wrapText": True, "editable": False, "cellStyle": {"white-space": "normal", "lineHeight": "1.2"}, "filter": "agTextColumnFilter", "floatingFilter": True},
    {"headerName": "Ementa", "field": "ementa", "width": 350, "wrapText": True, "autoHeight": True, "editable": False, "cellStyle": {"white-space": "normal", "lineHeight": "1.2"}, "filter": "agTextColumnFilter", "floatingFilter": True},
    {"headerName": "Autor/Partido", "field": "autorPartidoUf", "width": 220, "wrapText": True, "autoHeight": True, "editable": False, "cellStyle": {"white-space": "normal", "lineHeight": "1.2"}, "filter": "agTextColumnFilter", "floatingFilter": True},
    {"headerName": "Impacto Fiscal", "field": "impactoFiscal", **impacto_editor, "width": 160, "editable": True, "filter": True, "floatingFilter": True},
    {"headerName": "Tipo Impacto", "field": "tipoImpactoFiscal", **tipo_editor, "width": 200, "editable": True, "filter": True, "floatingFilter": True, "wrapText": True, "autoHeight": True,"cellStyle": {
            "white-space": "normal",
            "lineHeight": "1.2"
        }},
    {"field": "proposicao_raw", "hide": True},
    {"field": "evento_id", "hide": True}
]

def make_layout():
    return dbc.Container(
        [
            dcc.Location(id="url", refresh=False),
            html.H1("Agenda da Semana", style={"color": "#003366", "textAlign": "left", "marginBottom": "20px"}), 
            html.Hr(),
            dbc.Row(
                dbc.Col(
                    [
                        html.Label("Selecione o intervalo de datas:",
                                   style={"fontWeight": "bold"}),
                        dcc.DatePickerRange(
                            id="date-picker-range",
                            start_date=datetime.today().strftime("%Y-%m-%d"),
                            end_date=(datetime.today() + timedelta(days=7))
                                     .strftime("%Y-%m-%d"),
                            display_format="DD/MM/YYYY",
                            className="mb-2",
                            start_date_placeholder_text="Data Início",
                            end_date_placeholder_text="Data Fim",
                        ),

                        # ❶ Botão que abre o modal
                        dbc.Button(
                            "Buscar Novos Dados e Atualizar Base",
                            id="btn-gerar-planilha",
                            color="primary",
                            className="ms-2",
                            n_clicks=0,
                        ),

                        # ❷ Modal de confirmação
                        dbc.Modal(
                            [
                                dbc.ModalHeader(
                                    dbc.ModalTitle("Confirmação da Operação")
                                ),
                                dbc.ModalBody(
                                    "A operação a seguir irá apagar a base "
                                    "atual e inserir uma nova. "
                                    "As proposições já analisadas estão salvas.\n\n"
                                    "Deseja continuar?"
                                ),
                                dbc.ModalFooter(
                                    [
                                        dbc.Button(
                                            "Sim",
                                            id="btn-modal-sim",
                                            color="danger",
                                            className="me-2",
                                            n_clicks=0,
                                        ),
                                        dbc.Button(
                                            "Não",
                                            id="btn-modal-nao",
                                            color="secondary",
                                            n_clicks=0,
                                        ),
                                    ]
                                ),
                            ],
                            id="modal-confirm",
                            is_open=False,
                            backdrop="static",   # impede clique fora fechar
                            keyboard=False,      # bloqueia ESC
                            centered=True,
                        ),
                    ],
                    width=12,
                    lg=8,
                    md=10,
                    className="mb-3",
                ),
                justify="start",
            ),
            dbc.Toast(
                [dbc.Spinner(size="sm"), " Atualizando dados…"],
                id="toast-loading", header="Processando",
                icon="primary",  # cor do cabeçalho
                dismissable=False, duration=None, is_open=False,
                style={
                    "position": "fixed", "top": 70, "right": 10,
                    "width": 320, "zIndex": 1100
                },
            ),
            html.Div([
                html.Label(
                    "Filtrar por Casa:",
                    style={"fontWeight": "bold", "marginBottom": "8px"}  # espaço abaixo do título
                ),
                dcc.Checklist(
                    id="checklist_casa",
                    options=[
                        {"label": " Câmara (CD)", "value": "CD"},
                        {"label": " Senado (SF)", "value": "SF"}
                    ],
                    value=["CD", "SF"],
                    inline=True,
                    className="mb-4",                  # aumenta a distância para o próximo bloco
                    style={"display": "flex", "gap": "20px"},
                    inputStyle={"marginRight": "4px"}  # reduz o espaço entre checkbox e texto
                ),
                html.Label(
                    "Filtrar Plenário/Comissão:",
                    style={"fontWeight": "bold", "marginBottom": "8px", "marginTop": "16px"}  
                    # margem superior extra para reforçar o espaçamento entre blocos, se necessário
                ),
                dcc.Checklist(
                    id="checklist_plen_com",
                    options=[
                        {"label": " Comissão (CM)", "value": "CM"},
                        {"label": " Plenário (PLEN)", "value": "PLEN"}
                    ],
                    value=["CM", "PLEN"],
                    inline=True,
                    className="mb-3",                  # distância padrão abaixo deste bloco
                    style={"display": "flex", "gap": "20px"},
                    inputStyle={"marginRight": "4px"}  # reduz o espaço entre checkbox e texto
                )
            ],
            className="mb-3 bg-light p-3 border rounded"),
            dbc.Row([
                dbc.Col(dbc.Button("Exportar para Excel (XLSX)", id="btn-exportar", color="success", className="me-2 mb-2 mb-md-0"), width="auto"), 
                dbc.Col(dcc.Download(id="download-xlsx"), width="auto"),
                dbc.Col(dbc.Button("Exportar para Word (DOCX)", id="btn-exportar-word", color="info", className="mb-2 mb-md-0"), width="auto"),
                dbc.Col(dcc.Download(id="download-docx"), width="auto")], className="mb-3", justify="start"),
            dcc.Store(id="full_data", storage_type="session"), 
            dcc.Loading(id="loading-wrapper", type="default", children=[
                html.Div(id="div-mensagens", className="mb-2 alert alert-info", role="alert"), 
                dag.AgGrid(id="tabela_dados", rowData=[], columnDefs=columnDefs,
                           defaultColDef={"resizable": True, "sortable": True, "filter": True, "floatingFilter": True, "editable": False, "minWidth": 100,
                                          "cellStyle": {"fontSize": "14px", "lineHeight": "24px"},},
                           dashGridOptions={
                                "suppressFieldDotNotation": True,
                                "domLayout": "autoHeight",
                                "undoRedoCellEditing": True,
                                "rowSelection": "multiple",
                                "suppressRowClickSelection": True,
                                "frameworkComponents": {
                                    "markdownLinkRenderer": "markdownLinkRenderer"
                                },
                                "enableClipboard": True,
                                "enableRangeSelection": True
                            },
                           style={"height": "600px", "width": "100%"}, className="ag-theme-alpine", dangerously_allow_code=True, enableEnterpriseModules=False)]),
            html.Div(id="div-salvar-edicao-feedback", className="mt-3")
        ], fluid=True, className="p-4" )

app = Dash(__name__, external_stylesheets=[dbc.themes.BOOTSTRAP])
app.title = "Agenda Legislativa Semanal" 
server = app.server 
app.layout = make_layout()

@app.callback(Output("full_data", "data"), Input("url", "pathname"), prevent_initial_call=False)
def inicializar_dados(_):
    print("[LOG] inicializar_dados: Carregando dados iniciais do SQLite.")
    df_dados = carregar_dados_formatados_do_sqlite()
    print(f"[LOG] inicializar_dados: {len(df_dados)} registros carregados e formatados.")
    return df_dados.to_dict("records")

### INÍCIO BLOCO CALLBACK_MODAL
@app.callback(
    Output("modal-confirm", "is_open"),
    Input("btn-gerar-planilha", "n_clicks"),
    Input("btn-modal-sim", "n_clicks"),
    Input("btn-modal-nao", "n_clicks"),
    State("modal-confirm", "is_open"),
    prevent_initial_call=True,
)

def toggle_modal(n_clicks_open, n_clicks_yes, n_clicks_no, is_open):
    """
    Abre o modal quando o usuário clica no botão principal
    e fecha quando ele clica em 'Sim' ou 'Não'.
    """
    ctx = callback_context
    if not ctx.triggered:
        raise PreventUpdate

    trigger = ctx.triggered[0]["prop_id"].split(".")[0]

    # Clique no botão principal → abrir
    if trigger == "btn-gerar-planilha" and n_clicks_open:
        return True

    # Clique em 'Sim' ou 'Não' → fechar
    if trigger in ("btn-modal-sim", "btn-modal-nao"):
        return False

    return is_open

@app.callback(
    Output("toast-loading", "is_open"),
    Input("btn-modal-sim", "n_clicks"),
    prevent_initial_call=True,
)
def show_toast_loading(n):
    if n:
        return True
    raise PreventUpdate

@app.callback(
    [
        Output("full_data", "data", allow_duplicate=True),
        Output("div-salvar-edicao-feedback", "children",
               allow_duplicate=True),
        Output("toast-loading", "is_open", allow_duplicate=True),
    ],
    Input("btn-modal-sim", "n_clicks"),
    State("date-picker-range", "start_date"),
    State("date-picker-range", "end_date"),
    prevent_initial_call=True,
)
def callback_gerar_planilha_e_atualizar_base(
    n_clicks_yes, start_date, end_date
):
    if not n_clicks_yes:
        raise PreventUpdate

    if not start_date or not end_date:
        return dash.no_update, dbc.Alert(
            "Por favor, selecione as datas de início e fim.",
            color="warning", dismissable=True, duration=4000
        ), False # fecha toast se houve erro

    print(
        f"[LOG] callback_gerar_planilha: Confirmação recebida "
        f"para {start_date} a {end_date}."
    )

    df_disp = gerar_planilha(start_date[:10], end_date[:10])
    novos_dados_records = df_disp.to_dict("records")

    msg_texto = (
        f"Busca concluída! {len(novos_dados_records)} "
        "registros processados. Base atualizada."
    )

    return (
        novos_dados_records,                       # 1ª saída
        dbc.Alert(                                 # 2ª saída
            msg_texto,
            color="success",
            dismissable=True,
            duration=5000,
        ),
        False                                       # 3ª saída → fecha o toast
    )

@app.callback(
    [Output("tabela_dados", "rowData"), Output("div-mensagens", "children")],
    [Input("full_data", "data"), Input("checklist_casa", "value"), Input("checklist_plen_com", "value")] )
def callback_aplicar_filtros_e_preencher_grade(dados_completos_records, casas_selecionadas, plen_com_selecionados):
    if dados_completos_records is None: return [], "Aguardando dados..." 
    if not dados_completos_records: return [], "Nenhum dado carregado."
    df = pd.DataFrame(dados_completos_records)
    if df.empty: return [], "Nenhum dado para exibir."
    df_filtrado = df.copy() 
    if casas_selecionadas: df_filtrado = df_filtrado[df_filtrado["casa"].isin(casas_selecionadas)]
    else: df_filtrado = pd.DataFrame(columns=df.columns) 
    if not df_filtrado.empty and plen_com_selecionados:
        df_filtrado = df_filtrado[df_filtrado["plenarioOuComissao"].isin(plen_com_selecionados)]
    elif not plen_com_selecionados and not df_filtrado.empty : 
        df_filtrado = pd.DataFrame(columns=df_filtrado.columns)
    num_total, num_filtrado = len(df), len(df_filtrado)
    mensagens = [html.Strong(f"Registros na base (período): {num_total}"), html.Br(), html.Span(f"Exibidos após filtros: {num_filtrado}")]
    return df_filtrado.to_dict("records"), mensagens

@app.callback(Output("download-xlsx", "data"), Input("btn-exportar", "n_clicks"), State("tabela_dados", "rowData"), prevent_initial_call=True)
def exportar_xlsx_callback(n_clicks, dados_grade): 
    if not n_clicks or not dados_grade: raise PreventUpdate
    df_export = pd.DataFrame(dados_grade)
    if df_export.empty: return None 
    df_export = df_export.drop(columns=[col for col in ['proposicao_raw', 'evento_id'] if col in df_export.columns], errors='ignore')
    def fmt_link_excel(md_link):
        if pd.isna(md_link) or not isinstance(md_link, str): return md_link 
        m = re.match(r'\[([^\]]+)\]\((https?://[^\)]+)\)', md_link)
        if m: text, url = m.groups(); return f'=HYPERLINK("{url}", "{text[:250].replace("\"", "\"\"")}")' 
        return md_link 
    for col_link in ['nomeComissaoPlenario', 'proposicao']:
        if col_link in df_export.columns: df_export[col_link] = df_export[col_link].apply(fmt_link_excel)
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine="xlsxwriter") as writer:
        df_export.to_excel(writer, sheet_name="AgendaLegislativa", index=False)
        ws = writer.sheets["AgendaLegislativa"]
        wrap_fmt = writer.book.add_format({'text_wrap': True, 'valign': 'top'})
        for idx, col_nome in enumerate(df_export.columns):
            max_len_col = 0
            if not df_export[col_nome].empty:
                 series_len = df_export[col_nome].astype(str).map(len)
                 if not series_len.empty:
                    max_len_val = series_len.max() 
                    if pd.isna(max_len_val): max_len_col = 0
                    else: max_len_col = int(max_len_val) 
                 else: 
                    max_len_col = 0
            
            header_len_col = len(str(col_nome))
            final_col_width = max(max_len_col, header_len_col) + 2
            final_col_width = min(max(final_col_width, 12), 60) 
            
            cols_com_wrap = ['ementa', 'nomeComissaoPlenario', 'proposicao', 'autorPartidoUf', 'tipoImpactoFiscal']
            ws.set_column(idx, idx, final_col_width, wrap_fmt if col_nome in cols_com_wrap else None)
    return dcc.send_bytes(output.getvalue(), "AgendaLegislativa_Exportada.xlsx")


# ==============================
# Callback para Exportar DOCX
# ==============================
dias_semana = {
    'Monday': 'Segunda-feira', 'Tuesday': 'Terça-feira',
    'Wednesday': 'Quarta-feira', 'Thursday': 'Quinta-feira',
    'Friday': 'Sexta-feira', 'Saturday': 'Sábado', 'Sunday': 'Domingo'
}
mapa_casa = {"CD": "Câmara dos Deputados", "SF": "Senado Federal", "CN": "Congresso Nacional"}

COLOR_SENADO = "4472C4"
COLOR_CAMARA = "00B050"
LIGHT_COLOR_SENADO = "DDEBF7"
LIGHT_COLOR_CAMARA = "E2EFDA"

# ------------------------------------------------------------------
def set_paragraph_background(paragraph, fill_color):
    """Colorir fundo do parágrafo e deixar fonte branca."""
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill_color.upper())
    pPr.append(shd)
    for run in paragraph.runs:
        run.font.color.rgb = RGBColor(255, 255, 255)

def set_paragraph_light_background(paragraph, fill_color):
    """Fundo suave (para cabeçalhos de reuniões)."""
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill_color.upper())
    pPr.append(shd)

# --- remove NBSP e espaços excedentes ------------------------------
def _clean_str(val):
    if pd.isna(val):
        return val
    return str(val).replace('\u00A0', ' ').strip()

# ==============================
# Callback para Exportar DOCX
# ==============================
@app.callback(
    Output("download-docx", "data"),
    Input("btn-exportar-word", "n_clicks"),
    State("tabela_dados", "rowData"),
    prevent_initial_call=True
)
def exportar_docx(n_clicks, current_rows):
    if not n_clicks:
        raise dash.exceptions.PreventUpdate

    # ------------------------------------------------------------
    # 1) Converte em DF e filtra: impactoFiscal == 'S'
    # ------------------------------------------------------------
    df = pd.DataFrame(current_rows)

    if df.empty or "impactoFiscal" not in df.columns:
        # nada para exportar
        buf = io.BytesIO()
        Document().save(buf)
        buf.seek(0)
        return dcc.send_bytes(buf.getvalue(), "Relatorio_Vazio.docx")

    impacto = (
        df["impactoFiscal"]
        .fillna("").astype(str).str.strip().str.upper()
    )
    df = df[impacto == "S"].copy()

    if df.empty:
        buf = io.BytesIO()
        Document().save(buf)
        buf.seek(0)
        return dcc.send_bytes(buf.getvalue(), "Relatorio_Vazio.docx")

    # ------------------------------------------------------------
    # 2) Colunas auxiliares que o relatório precisa
    # ------------------------------------------------------------
    # 2.1 Data e Hora separadas
    df["Data"] = (
        df["dataEvento"]
        .apply(_clean_str)
        .pipe(pd.to_datetime, dayfirst=True, errors="coerce")
    )
    df = df.dropna(subset=["Data"])
    df["Hora"] = df["horaEvento"].apply(_clean_str)

    # 2.2 Texto sem markdown para Comissão/Plenário e PL
    def _strip_md(md):
        if pd.isna(md) or not isinstance(md, str):
            return str(md)
        m = re.match(r"\[([^\]]+)\]\([^)]*\)", md)
        return m.group(1) if m else md

    df["ComissaoPlenTexto"] = df["nomeComissaoPlenario"].apply(_strip_md)
    df["PL_Texto"] = df["proposicao_raw"].apply(_strip_md)

    # ------------------------------------------------------------
    # 3) Cria o documento
    # ------------------------------------------------------------
    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width, section.page_height = Inches(8.27), Inches(11.69)
    for m in (section.left_margin, section.right_margin,
              section.top_margin, section.bottom_margin):
        m = Cm(2)

    styles = document.styles
    styles["Normal"].font.name = styles["Heading 1"].font.name = "Calibri"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    for h in ("Heading 1", "Heading 2", "Heading 3"):
        styles[h]._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")

    document.add_heading("Agenda da Semana", level=1)

    # ------------------------------------------------------------
    # 4) Grupos: casa + tipo (PLEN/CM)
    # ------------------------------------------------------------
    groups = [
        ("SF", "PLEN", "Plenário - Senado Federal"),
        ("SF", "CM",   "Comissões - Senado Federal"),
        ("CD", "PLEN", "Plenário - Câmara dos Deputados"),
        ("CD", "CM",   "Comissões - Câmara dos Deputados"),
    ]

    for casa_code, flag_tipo, titulo in groups:
        subset = df[
            (df["casa"] == casa_code) &
            (df["plenarioOuComissao"] == flag_tipo)
        ]
        if subset.empty:
            continue

        # Barra colorida do título
        p_tit = document.add_heading(titulo, level=2)
        cor_barra = COLOR_SENADO if casa_code == "SF" else COLOR_CAMARA
        set_paragraph_background(p_tit, cor_barra)
        document.add_paragraph("")

        # Reuniões agrupadas por Comissão/Plenário + Data + Hora
        for (nome, dia, hora), bloco in subset.groupby(
                ["ComissaoPlenTexto", "Data", "Hora"], sort=False):

            dia_str  = dia.strftime("%d/%m/%Y")
            hora_str = str(hora) if pd.notna(hora) else ""
            header   = f"{nome} | {dia_str} | {hora_str}"

            p_meet = document.add_paragraph()
            run_meet = p_meet.add_run(header)
            run_meet.bold = True
            p_meet.style.font.size = Pt(11)
            cor_clara = LIGHT_COLOR_SENADO if casa_code == "SF" else LIGHT_COLOR_CAMARA
            set_paragraph_light_background(p_meet, cor_clara)

            # Cada proposição
            for _, row in bloco.iterrows():
                p_reg = document.add_paragraph()
                run_pl = p_reg.add_run(str(row["PL_Texto"]))
                run_pl.bold = True
                p_reg.add_run(f" ({row['autorPartidoUf']})").font.size = Pt(11)
                p_reg.add_run("\n" + str(row["ementa"]))
                p_reg.style.font.size = Pt(11)

                motivo = row.get("tipoImpactoFiscal")
                if pd.notna(motivo) and str(motivo).strip():
                    p_mot = document.add_paragraph("Impacto fiscal: " + str(motivo))
                    p_mot.style.font.size = Pt(11)

            document.add_paragraph("")

    # ------------------------------------------------------------
    # 5) Retorna DOCX
    # ------------------------------------------------------------
    buf = io.BytesIO()
    document.save(buf)
    buf.seek(0)
    return dcc.send_bytes(buf.getvalue(), "Agenda_Semana.docx")


@app.callback(Output("div-salvar-edicao-feedback", "children"), Input("tabela_dados", "cellValueChanged"), prevent_initial_call=True)
def salvar_edicao_celula(evento_edicao): 
    if not evento_edicao: raise PreventUpdate
    if isinstance(evento_edicao, list): evento_edicao = evento_edicao[0] if evento_edicao else None
    if not isinstance(evento_edicao, dict) or "colId" not in evento_edicao: raise PreventUpdate
    col_mod = evento_edicao["colId"]
    if col_mod not in ("impactoFiscal", "tipoImpactoFiscal", "marcarParaRelatorio"): raise PreventUpdate 
    dados_linha = evento_edicao.get("data", {})
    if not dados_linha: raise PreventUpdate
    ev_id, prop_raw = dados_linha.get("evento_id"), dados_linha.get("proposicao_raw") 
    if not ev_id or not prop_raw:
        return dbc.Alert("Erro: IDs não encontrados para salvar.", color="danger", duration=5000)
    val_imp = dados_linha.get("impactoFiscal", ""); val_tipo = dados_linha.get("tipoImpactoFiscal", ""); val_marcar = dados_linha.get("marcarParaRelatorio", "")
    try:
        with get_conn() as con:
            cur = con.execute("UPDATE eventos SET impactoFiscal=?, tipoImpactoFiscal=?, marcarParaRelatorio=? WHERE evento_id=? AND proposicao=?", 
                              (val_imp, val_tipo, val_marcar, ev_id, prop_raw))
            con.execute(
                """
                INSERT INTO backup_pl (proposicao, impactoFiscal, tipoImpactoFiscal)
                VALUES (?, ?, ?)
                ON CONFLICT(proposicao) DO UPDATE SET
                    impactoFiscal    = excluded.impactoFiscal,
                    tipoImpactoFiscal= excluded.tipoImpactoFiscal
                """,
                (prop_raw, val_imp, val_tipo)
            )
            if cur.rowcount > 0: msg = f"Salvo: Prop '{prop_raw}' (Ev: {ev_id}) atualizada."
            else: msg = f"Aviso: Nenhum registro atualizado para Ev: {ev_id}, Prop: '{prop_raw}'."
            return dbc.Alert(msg, color="success" if cur.rowcount > 0 else "warning", duration=4000)
    except sqlite3.Error as e: return dbc.Alert(f"Erro DB: {e}", color="danger", duration=5000)
    except Exception as e: return dbc.Alert(f"Erro: {e}", color="danger", duration=5000)

if __name__ == "__main__":
    app.run(debug=False, host="0.0.0.0", port=8081)
