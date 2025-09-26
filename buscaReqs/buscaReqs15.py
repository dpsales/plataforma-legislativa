#!/usr/bin/env python
# -*- coding: utf-8 -*-

# ----------------------------------------------------------
# IMPORTS
# ----------------------------------------------------------
import os
import io                      # já existia
import pandas as pd
import requests
import datetime
from datetime import datetime as dt
import logging
from apscheduler.schedulers.background import BackgroundScheduler
import time
import re
from flask import Flask, jsonify
import dash
import threading
import pytz
from requests.adapters import HTTPAdapter, Retry
import http.client
from json import JSONDecodeError
from requests.exceptions import ChunkedEncodingError
from PyPDF2 import PdfReader

# >>> Biblioteca do Google Cloud Storage <<<
from google.cloud import storage

from concurrent.futures import ThreadPoolExecutor, as_completed
MAX_WORKERS = 12

# ---- Banco de Dados (novo) ----
from sqlalchemy import create_engine, text
from sqlalchemy.exc import SQLAlchemyError

# ----------------------------------------------------------
# FLASK SERVER / LOG / VARIÁVEIS GLOBAIS
# ----------------------------------------------------------
server = Flask(__name__)

logging.basicConfig(
    level=logging.INFO,
    format='[%(asctime)s] %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)
logger.setLevel(logging.DEBUG)

# Flags / Modo
LOCAL_MODE = str(os.getenv("LOCAL_MODE", "true")).lower() == "true"   # True = teste local | False = produção (GCS)
USE_DB     = str(os.getenv("USE_DB", "false")).lower() == "true"      # Liga/desliga uso de banco

# GCS (opcional)
BUCKET_NAME       = os.getenv("BUCKET_NAME", "data")
GCS_OBJECT_NAME   = os.getenv("GCS_OBJECT_NAME", "resultados_filtrados_expressa.xlsx")
LOCAL_OUTPUT_FILE = os.getenv("LOCAL_OUTPUT_FILE", "data/resultados_filtrados_expressa.xlsx")

# DB URL (Postgres ou MySQL)
DB_URL = os.getenv("DB_URL", "")  # Ex.: postgresql+psycopg2://usuario:senha@host:5432/db

scheduler_global = None

def is_db_empty() -> bool:
    """Verifica se a tabela de proposicoes está vazia."""
    if not USE_DB:
        return True # Se não usa DB, considera "vazio" para fins de arquivo
    engine = get_engine()
    if engine is None:
        return True
    try:
        with engine.begin() as conn:
            # Query mais eficiente para verificar se existe algum registro
            result = conn.execute(text("SELECT 1 FROM proposicoes LIMIT 1")).scalar()
            is_empty = result is None
            logger.info(f"Verificando se DB está vazio: {'Sim' if is_empty else 'Não'}")
            return is_empty
    except SQLAlchemyError as e:
        logger.error(f"Erro ao verificar se DB está vazio: {e}")
        # Em caso de erro, é mais seguro assumir que está vazio e tentar buscar dados.
        return True

# ----------------------------------------------------------
# TERMOS DE INTERESSE E CONSTANTES CAM/SF
# ----------------------------------------------------------
TERMOS = [
    "Tebet",
    "Ministério do Planejamento",
    "Ministério do Planejamento e Orçamento",
    "Ministra do Planejamento",
    "(MPO)",
    "Ministra de Estado do Planejamento e Orçamento",
    "Instituto Brasileiro de Geografia e Estatística",
    "Pochmann",
    "Luciana Mendes Santos Servo",
    "Instituto de Pesquisa Econômica Aplicada",
    "Clayton Luiz Montes",
    "Secretaria de Orçamento Federal",
    "(SOF)",
    "Wesley Matheus de Oliveira",
    "Secretaria de Monitoramento e Avaliação de Políticas Públicas e Assuntos Econômicos",
    "Virgínia de Ângelis Oliveira de Paula",
    "Secretaria Nacional de Planejamento",
    "SEPLAN",
    "Lei Complementar nº 101",
    "Plano Plurianual",
    "Lei de Diretrizes Orçamentárias",
    "Projeto de Lei de Diretrizes Orçamentárias",
    "Lei Orçamentária Anual",
    "Projeto de Lei Orçamentária Anual",
    "(PLOA 2026)"
]

TIPOS_BUSCA_CAMARA = ["RIC", "INC", "REQ"]
URL_CAMARA         = "https://dadosabertos.camara.leg.br/api/v2/proposicoes"

# ==============================================================================
# CAMADA DE BANCO – helpers
# ==============================================================================
_engine = None

def get_engine():
    global _engine
    if not USE_DB:
        return None
    if _engine is None:
        if not DB_URL:
            logger.error("USE_DB=true mas DB_URL não foi definido.")
            return None
        _engine = create_engine(DB_URL, pool_pre_ping=True, pool_recycle=1800)
    return _engine

def ensure_tables():
    """
    Cria a tabela 'proposicoes' se não existir.
    Colunas espelham o DataFrame consolidado.
    """
    if not USE_DB:
        return
    engine = get_engine()
    if engine is None:
        return
    ddl = """
    CREATE TABLE IF NOT EXISTS proposicoes (
        CodigoMateria TEXT,
        Titulo TEXT,
        Autor TEXT,
        DataApresentacao TEXT,
        Ementa TEXT,
        SituacaoAtual TEXT,
        DataUltimaTramitacao TEXT,
        DescricaoUltimaTramitacao TEXT,
        LinkFicha TEXT,
        LinkInteiroTeor TEXT,
        TermosEncontrados TEXT,
        Local TEXT,
        "Prazo para Resposta" TEXT,
        Casa TEXT
    );
    """
    with engine.begin() as conn:
        conn.execute(text(ddl))

def df_to_db(df: pd.DataFrame):
    """
    Grava o DF na tabela (estratégia simples: truncate + bulk insert).
    Se preferir 'upsert', dá para trocar por MERGE/ON CONFLICT.
    """
    if not USE_DB:
        return
    engine = get_engine()
    if engine is None:
        return
    # Normaliza colunas para garantir o mesmo esquema
    expected_cols = [
        "CodigoMateria","Titulo","Autor","DataApresentacao","Ementa",
        "SituacaoAtual","DataUltimaTramitacao","DescricaoUltimaTramitacao",
        "LinkFicha","LinkInteiroTeor","TermosEncontrados","Local",
        "Prazo para Resposta","Casa"
    ]
    for c in expected_cols:
        if c not in df.columns:
            df[c] = None
    df = df[expected_cols].copy()

    # TRUNCATE + INSERT (simples e rápido para esse caso)
    try:
        with engine.begin() as conn:
            conn.execute(text("DELETE FROM proposicoes"))
        df.to_sql("proposicoes", engine, if_exists="append", index=False, method="multi", chunksize=1000)
        logger.info("Dados gravados no banco com sucesso.")
    except SQLAlchemyError as e:
        logger.error(f"Erro ao gravar no banco: {e}")

def load_df_from_db() -> pd.DataFrame:
    if not USE_DB:
        return pd.DataFrame()
    engine = get_engine()
    if engine is None:
        return pd.DataFrame()
    try:
        with engine.begin() as conn:
            df = pd.read_sql(text("SELECT * FROM proposicoes"), conn)
        logger.info("Dados carregados do banco com sucesso.")
        return df
    except SQLAlchemyError as e:
        logger.error(f"Erro ao carregar do banco: {e}")
        return pd.DataFrame()

# ==============================================================================
# FUNÇÕES AUXILIARES – SALVAR/CARREGAR XLSX
# ==============================================================================
def _forcar_https(url: str | None) -> str:
    if not url:
        return ""
    if url.startswith("http://www.camara.leg.br") or url.startswith("http://www25.senado.leg.br"):
        return url.replace("http://", "https://", 1)
    return url

def salvar_df(df: pd.DataFrame):
    # 1) Banco (se habilitado)
    if USE_DB:
        ensure_tables()
        df_to_db(df)

    # 2) XLSX (mantido para compatibilidade)
    try:
        if LOCAL_MODE:
            df.to_excel(LOCAL_OUTPUT_FILE, index=False)
            logger.info(f"Dados salvos localmente em {LOCAL_OUTPUT_FILE}")
        else:
            from io import BytesIO
            output = BytesIO()
            df.to_excel(output, index=False, engine='openpyxl')
            output.seek(0)
            try:
                client = storage.Client()
                bucket = client.bucket(BUCKET_NAME)
                blob = bucket.blob(GCS_OBJECT_NAME)
                blob.upload_from_file(
                    output,
                    content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
                )
                logger.info(f"Arquivo XLSX enviado ao GCS: gs://{BUCKET_NAME}/{GCS_OBJECT_NAME}")
            except Exception as e:
                logger.error(f"Erro ao salvar DataFrame no GCS: {e}")
    except Exception as e:
        logger.error(f"Erro ao salvar XLSX: {e}")

def carregar_df() -> pd.DataFrame:
    # Lógica exclusiva para o banco de dados
    if USE_DB:
        # Se USE_DB for verdadeiro, SÓ tentamos o banco. Não há fallback.
        return load_df_from_db()

    # O código abaixo só executa se USE_DB for falso
    logger.info("USE_DB=false, carregando de arquivo XLSX.")
    try:
        if LOCAL_MODE:
            if not os.path.exists(LOCAL_OUTPUT_FILE):
                logger.warning(f"Arquivo {LOCAL_OUTPUT_FILE} não encontrado.")
                return pd.DataFrame()
            df = pd.read_excel(LOCAL_OUTPUT_FILE, engine='openpyxl')
            logger.info("Arquivo carregado localmente com sucesso.")
            return df
        else:
            # Lógica GCS
            client = storage.Client()
            bucket = client.bucket(BUCKET_NAME)
            blob = bucket.blob(GCS_OBJECT_NAME)
            if not blob.exists():
                logger.warning(f"Objeto gs://{BUCKET_NAME}/{GCS_OBJECT_NAME} não existe.")
                return pd.DataFrame()
            from io import BytesIO
            data = BytesIO()
            blob.download_to_file(data)
            data.seek(0)
            df = pd.read_excel(data, engine='openpyxl')
            logger.info("Carregado XLSX do GCS com sucesso.")
            return df
    except Exception as e:
        logger.error(f"Erro ao carregar o arquivo XLSX: {e}")
        return pd.DataFrame()
    
def pegar_data_ultima_atualizacao() -> str:
    try:
        # Se estiver usando o banco, a "última atualização" é agora.
        if USE_DB:
            sp_tz = pytz.timezone("America/Sao_Paulo")
            dt_sp = dt.now(sp_tz)
            return dt_sp.strftime("Dados do Banco em %d/%m/%Y às %H:%M")

        # O código abaixo só executa se USE_DB for falso
        if LOCAL_MODE:
            if not os.path.exists(LOCAL_OUTPUT_FILE):
                return "Nenhum arquivo local gerado ainda."
            timestamp = os.path.getmtime(LOCAL_OUTPUT_FILE)
            dt_obj = dt.fromtimestamp(timestamp)
        else:
            # Lógica GCS
            client = storage.Client()
            bucket = client.bucket(BUCKET_NAME)
            blob = bucket.blob(GCS_OBJECT_NAME)
            if not blob.exists():
                return "Nenhum arquivo no GCS gerado ainda."
            blob.reload()
            if not blob.updated:
                return "Arquivo sem data de atualização."
            dt_obj = blob.updated
            
        sp_tz = pytz.timezone("America/Sao_Paulo")
        dt_sp = dt_obj.astimezone(sp_tz)
        return dt_sp.strftime("Arquivo atualizado em %d/%m/%Y às %H:%M")
    except Exception as e:
        logger.error(f"Erro ao obter data de atualização: {e}")
        return "Falha ao obter data de atualização."
    
# ==============================================================================
# FUNÇÕES AUXILIARES GERAIS (parse, safe_request, ordenar, ...)
# ==============================================================================
def parse_data(strdata):
    formatos = [
        "%Y-%m-%dT%H:%M:%S",
        "%Y-%m-%dT%H:%M",
        "%Y-%m-%d %H:%M:%S",
        "%Y-%m-%d %H:%M",
        "%Y-%m-%d",
        "%d/%m/%Y %H:%M:%S",
        "%d/%m/%Y"
    ]
    for fmt in formatos:
        try:
            return dt.strptime(strdata, fmt)
        except:
            pass
    return None

def contem_termos(texto, termos):
    if not texto:
        return False
    texto_lower = texto.lower()
    return any(termo.lower() in texto_lower for termo in termos)

def safe_request(
    url, *, params=None, timeout=60, max_tentativas=3,
    stream=False, headers=None, **kwargs
):
    default_headers = {
        "User-Agent": (
            "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
            "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/124.0 Safari/537.36"
        ),
        "Accept": "*/*",
        "Accept-Encoding": "identity",
    }
    if headers:
        default_headers.update(headers)

    sess = requests.Session()
    retry_cfg = Retry(
        connect=max_tentativas,
        read=max_tentativas,
        total=max_tentativas,
        backoff_factor=1.5,
        status_forcelist=[500, 502, 503, 504],
        allowed_methods=["GET", "HEAD"],
        raise_on_status=False,
    )
    sess.mount("https://", HTTPAdapter(max_retries=retry_cfg))
    sess.mount("http://",  HTTPAdapter(max_retries=retry_cfg))

    for tent in range(max_tentativas):
        try:
            # --- CORREÇÃO APLICADA AQUI ---
            # Passamos os parâmetros corretamente para a chamada 'get'
            r = sess.get(
                url,
                params=params,
                timeout=timeout,
                stream=stream,
                headers=default_headers,
                **kwargs
            )
            # --------------------------------

            if not stream:
                try:
                    _ = r.content
                except (http.client.IncompleteRead, ChunkedEncodingError) as e:
                    raise e
            
            # Checa se o status code indica sucesso antes de retornar
            r.raise_for_status()
            return r
            
        except (http.client.IncompleteRead, ChunkedEncodingError) as e:
            logger.warning(f"safe_request - corpo incompleto em {url} – tent {tent+1}/{max_tentativas}")
            time.sleep(1.5 * (tent + 1))
            continue
        except requests.exceptions.RequestException as e:
            logger.warning(f"safe_request falhou em {url} – {e}")
            time.sleep(1.5 * (tent + 1))
            continue
            
    logger.error(f"safe_request esgotou tentativas em {url}")
    return None

def ordenar_por_ultima_tramitacao(df: pd.DataFrame) -> pd.DataFrame:
    if 'DataUltimaTramitacao' in df.columns:
        df['DataUltimaTramitacao_dt'] = pd.to_datetime(
            df['DataUltimaTramitacao'], dayfirst=True, errors='coerce'
        )
        df.sort_values(by='DataUltimaTramitacao_dt', ascending=False, inplace=True)
        df.drop(columns=['DataUltimaTramitacao_dt'], inplace=True)
    return df

# =====================================================================
# === BUSCA TODOS OS TERMOS EM TODAS AS PÁGINAS DO PDF ================
# =====================================================================
def pdf_tem_termos(url_pdf: str,
                   termos: list[str] = TERMOS) -> list[str]:
    """
    Baixa o PDF (url_pdf) e devolve a lista de TERMOS encontrados
    em *todas* as páginas.  Lista vazia  ⇒  nenhum termo.
    """
    if not url_pdf:
        return []

    resp = safe_request(url_pdf, timeout=120, stream=True)
    if not (resp and resp.status_code == 200 and resp.content):
        return []

    try:
        pdf = PdfReader(io.BytesIO(resp.content))
    except Exception as e:
        logger.error(f"Erro ao abrir PDF {url_pdf}: {e}")
        return []

    termos_lower = [t.lower() for t in termos]
    encontrados: set[str] = set()

    for page in pdf.pages:
        try:
            texto = (page.extract_text() or "").lower()
        except Exception:
            texto = ""

        for termo, termo_l in zip(termos, termos_lower):
            if termo_l in texto:
                encontrados.add(termo)

    return list(encontrados)

# ==============================================================================
# FUNÇÕES DE COLETA DE DADOS – CÂMARA
# ==============================================================================
def obter_local_ultima_tramitacao(tramitacao):
    sigla = tramitacao.get("siglaOrgao", "")
    uri_orgao = tramitacao.get("uriOrgao", "")
    nome_pub = ""
    if uri_orgao:
        r = safe_request(uri_orgao, timeout=30)
        if r and r.status_code == 200:
            dados_orgao = r.json().get("dados", {})
            nome_pub = dados_orgao.get("nomePublicacao", "")
    if sigla and nome_pub:
        return f"{sigla} - {nome_pub}"
    elif sigla:
        return sigla
    return ""

def obter_ultima_tramitacao_cam(det_id: int):
    """
    Retorna os dados da tramitação mais recente
    (ordenada por dataHora e, em empate, pelo maior sequencial).
    """
    tram_url = f"{URL_CAMARA}/{det_id}/tramitacoes"
    tram_r   = safe_request(tram_url, timeout=30)
    tram     = tram_r.json().get("dados", []) if (tram_r and tram_r.status_code == 200) else []

    if not tram:
        return "", "", "", "", ""

    for t in tram:
        t["dataHora_dt"]  = parse_data(t.get("dataHora", ""))
        try:
            t["sequencia_int"] = int(t.get("sequencia", 0))
        except ValueError:
            t["sequencia_int"] = 0

    tram.sort(
        key=lambda x: (x["dataHora_dt"], x["sequencia_int"]),
        reverse=True
    )
    ultima = tram[0]

    data_ultima = (ultima["dataHora_dt"].strftime("%d/%m/%Y %H:%M:%S")
                   if ultima["dataHora_dt"] else "")
    desc_ultima = ultima.get("descricaoTramitacao", "") or ""
    desp_ultima = ultima.get("despacho", "") or ""
    link_inteiro = ultima.get("url", "") or ""
    local        = obter_local_ultima_tramitacao(ultima)

    return data_ultima, desc_ultima, link_inteiro, local, desp_ultima

def obter_autor_proposicao_camara(codigo_proposicao):
    logger.info(f"[DEBUG] Obtendo autor da proposição da Câmara ID {codigo_proposicao}")
    url_autores = f"https://dadosabertos.camara.leg.br/api/v2/proposicoes/{codigo_proposicao}/autores"
    try:
        response = safe_request(url_autores, timeout=30)
        if response and response.status_code == 200 and response.content:
            dados = response.json().get("dados", [])
            autor_com_ordem_1 = next((item for item in dados if item.get("ordemAssinatura")==1), None)
            autor_alvo = autor_com_ordem_1 if autor_com_ordem_1 else (dados[0] if dados else None)
            if autor_alvo:
                nome = autor_alvo.get("nome", "Autor não disponível")
                deputado_uri = autor_alvo.get("uri")
                partido = "Sem Partido"
                uf = "Sem UF"
                if deputado_uri:
                    logger.info(f"[DEBUG] Obtendo dados adicionais do deputado via URI: {deputado_uri}")
                    deputado_response = safe_request(deputado_uri, timeout=30)
                    if deputado_response and deputado_response.status_code==200 and deputado_response.content:
                        deputado_dados = deputado_response.json().get("dados", {})
                        ultimo_status = deputado_dados.get("ultimoStatus", {})
                        partido = (ultimo_status.get("siglaPartido") or "Sem Partido").strip()
                        uf = (ultimo_status.get("siglaUf") or "Sem UF").strip()
                        logger.info(f"[DEBUG] Dados do deputado obtidos: Partido={partido}, UF={uf}")
                autor_str = f"{nome} ({partido}/{uf})"
                logger.info(f"[DEBUG] Autor (ordem=1) formatado: {autor_str}")
                return autor_str
            else:
                logger.warning("[AVISO] Nenhum autor encontrado na lista.")
        else:
            logger.error(f"[ERRO] Falha ao obter autores da proposição ID {codigo_proposicao}. Status: {response.status_code if response else 'N/A'}")
    except Exception as e:
        logger.error(f"[ERRO] Exceção ao obter autores da proposição ID {codigo_proposicao}: {e}")
    return "Autor não disponível"

def obter_prazo_resposta_cam(det_id: int) -> str:
    url = f"{URL_CAMARA}/{det_id}/tramitacoes"
    r = safe_request(url, timeout=30)
    if not (r and r.status_code==200):
        return ""
    tram = r.json().get("dados", [])
    tram.sort(
        key=lambda t: parse_data(t.get("dataHora","")) or dt.min,
        reverse=True
    )
    for t in tram:
        texto = t.get("despacho") or t.get("descricaoTramitacao","")
        if "prazo para resposta" in texto.lower():
            m = re.search(r"\(de\s*([^\)]+)\)", texto, flags=re.I)
            if m:
                return m.group(1).strip()
    return ""

def _baixar_processar_camara(uri_prop: str, termos=TERMOS):
    det_r = safe_request(uri_prop, timeout=30)
    if not (det_r and det_r.status_code == 200):
        return None
    det = det_r.json().get("dados", {})

    ementa  = det.get("ementa", "") or ""
    url_pdf = _forcar_https(det.get("urlInteiroTeor"))

    if not contem_termos(ementa, termos):
        logger.debug(f"[Câmara] {det.get('siglaTipo')} {det.get('numero')}/{det.get('ano')} – sem termo, descartado")
        return None

    (data_ultima,
     desc_ultima,
     link_inteiro,
     local,
     desp_ultima) = obter_ultima_tramitacao_cam(det["id"])

    autor_str = obter_autor_proposicao_camara(det["id"])
    ficha_link = (f"https://www.camara.leg.br/proposicoesWeb/"
                  f"fichadetramitacao?idProposicao={det['id']}")

    termos_encontrados = [t for t in termos if t.lower() in ementa.lower()]
    onde_achou = "Ementa"

    dt_ap = parse_data(det.get("dataApresentacao", ""))
    data_apresentacao = (dt_ap.strftime("%d/%m/%Y %H:%M:%S")
                         if dt_ap else det.get("dataApresentacao", ""))

    casa_final = ("Congresso Nacional"
                  if "CMO - Comissão Mista de Orçamento" in local else "Câmara")

    prazo_resposta = obter_prazo_resposta_cam(det["id"])

    return {
        "CodigoMateria": det.get("id", ""),
        "Titulo": f"[{det['siglaTipo']} {det['numero']}/{det['ano']}]({url_pdf})",
        "Autor": autor_str,
        "DataApresentacao": data_apresentacao,
        "Ementa": ementa,
        "SituacaoAtual": desc_ultima,
        "DataUltimaTramitacao": data_ultima,
        "DescricaoUltimaTramitacao": desp_ultima,
        "LinkFicha": ficha_link,
        "LinkInteiroTeor": url_pdf,
        "TermosEncontrados": f"{', '.join(termos_encontrados)} ({onde_achou})",
        "Local": local,
        "Prazo para Resposta": prazo_resposta,
        "Casa": casa_final,
    }

def coletar_camara_datas(data_inicio, data_fim):
    resultados = []
    itens_por_pagina = 100
    pagina_atual = 1

    with ThreadPoolExecutor(MAX_WORKERS) as executor:
        logger.debug("→ coletar_camara_datas INICIO")
        while True:
            params = {
                "siglaTipo": ",".join(TIPOS_BUSCA_CAMARA),
                "itens": itens_por_pagina,
                "pagina": pagina_atual,
                "dataApresentacaoInicio": data_inicio,
                "dataApresentacaoFim": data_fim,
            }
            resp = safe_request(URL_CAMARA, params=params, timeout=30)
            if not (resp and resp.status_code==200):
                break

            proposicoes = resp.json().get("dados", [])
            if not proposicoes:
                break
            logger.debug(f"[Câmara] página {pagina_atual} – {len(proposicoes)} itens")
            futuros = {
                executor.submit(_baixar_processar_camara, p["uri"]): p["id"]
                for p in proposicoes
            }

            for fut in as_completed(futuros):
                try:
                    dado = fut.result()
                    if dado:
                        resultados.append(dado)
                except Exception as e:
                    logger.error("Thread falhou para id %s – %s", futuros[fut], e)

            if len(proposicoes) < itens_por_pagina:
                break
            pagina_atual += 1
    logger.debug(f"[Câmara] total coletado: {len(resultados)}")
    return resultados

# ==============================================================================
# FUNÇÕES SENADO
# ==============================================================================
def _processar_item_senado(item: dict) -> dict | None:
    ementa = item.get("ementa", "") or ""
    termos_ementa = [t for t in TERMOS if t.lower() in ementa.lower()]

    url_pdf       = _forcar_https(item.get("urlDocumento"))
    termos_pdf    = pdf_tem_termos(url_pdf)

    termos_encontrados = sorted(set(termos_ementa + termos_pdf))
    if not termos_encontrados:
        return None

    if termos_ementa and termos_pdf:
        onde_achou = "Ementa + Inteiro Teor"
    elif termos_ementa:
        onde_achou = "Ementa"
    else:
        onde_achou = "Inteiro Teor"

    detail_url   = f"http://legis.senado.gov.br/dadosabertos/processo/{item.get('id')}"
    detail_resp  = safe_request(detail_url, timeout=60)
    data_ultima_str = ""
    ultima_desc     = ""
    if detail_resp and detail_resp.status_code == 200:
        try:
            detail_data = detail_resp.json()
            autuacoes = detail_data.get("autuacoes", [])
            if isinstance(autuacoes, dict):
                autuacoes = [autuacoes]

            latest_dt = None
            for aut in autuacoes:
                situacoes = aut.get("situacoes", [])
                if isinstance(situacoes, dict):
                    situacoes = [situacoes]
                for sit in situacoes:
                    inicio = sit.get("inicio", "")
                    if inicio:
                        if " " not in inicio:
                            inicio += " 00:00:00"
                        dt_inicial = parse_data(inicio)
                        if dt_inicial and (latest_dt is None or dt_inicial > latest_dt):
                            latest_dt  = dt_inicial
                            ultima_desc = sit.get("descricao", "")
            if latest_dt:
                data_ultima_str = latest_dt.strftime("%d/%m/%Y %H:%M:%S")
        except Exception as e:
            logger.error(f"Erro ao processar detalhes {detail_url}: {e}")

    data_ap_raw = item.get("dataApresentacao", "")
    if data_ap_raw and " " not in data_ap_raw:
        data_ap_raw += " 00:00:00"
    dt_ap = parse_data(data_ap_raw)
    data_apresentacao = dt_ap.strftime("%d/%m/%Y %H:%M:%S") if dt_ap else data_ap_raw

    local       = item.get("siglaEnteIdentificador", "")
    casa_final  = ("Congresso Nacional"
                   if "CMO - Comissão Mista de Orçamento" in local else "Senado")
    url_pdf     = _forcar_https(item.get("urlDocumento"))

    return {
        "CodigoMateria": item.get("codigoMateria", ""),
        "Titulo": f"[{item.get('identificacao', '')}]({url_pdf})",
        "Autor": item.get("autoria", ""),
        "DataApresentacao": data_apresentacao,
        "Ementa": ementa,
        "SituacaoAtual": ultima_desc,
        "DataUltimaTramitacao": data_ultima_str,
        "DescricaoUltimaTramitacao": "",
        "LinkFicha": (
            f"https://www25.senado.leg.br/web/atividade/materias/-/materia/"
            f"{item.get('codigoMateria', '')}"
        ),
        "LinkInteiroTeor": url_pdf,
        "TermosEncontrados": f"{', '.join(termos_encontrados)} ({onde_achou})",
        "Local": local,
        "Prazo para Resposta": "",
        "Casa": casa_final,
    }

def coletar_senado_novo():
    resultados = []
    seen       = set()
    headers    = {"Accept": "application/json"}
    siglas     = ["req", "rqs"]

    with ThreadPoolExecutor(MAX_WORKERS) as executor:
        futuros = []
        for sigla in siglas:
            url = f"http://legis.senado.gov.br/dadosabertos/processo?sigla={sigla}"
            r = safe_request(url, headers=headers, timeout=120)
            if not (r and r.status_code == 200):
                logger.error(f"Falha ao acessar {url}. Status: {r.status_code if r else 'N/A'}")
                continue

            try:
                data = r.json()
            except Exception as e:
                logger.error(f"Erro ao decodificar JSON de {url}: {e}")
                continue

            for item in data:
                key = item.get("codigoMateria") or item.get("identificacao", "")
                if key in seen:
                    continue
                seen.add(key)
                futuros.append(executor.submit(_processar_item_senado, item))

        for fut in as_completed(futuros):
            try:
                res = fut.result()
                if res:
                    resultados.append(res)
            except Exception as e:
                logger.error(f"Thread Senado falhou: {e}")

    return resultados

def coletar_senado_unificado():
    return coletar_senado_novo()

# ==============================================================================
# === BUSCA PRINCIPAL
# ==============================================================================
def buscar_dados() -> pd.DataFrame:
    logger.info("Iniciando coleta de dados unificada (Câmara e Senado).")

    hoje = dt.now()
    inicio_periodo = (hoje - datetime.timedelta(days=90)).strftime("%Y-%m-%d")
    fim_periodo    = hoje.strftime("%Y-%m-%d")

    logger.info("Coletando dados da Câmara...")
    resultados_camara = coletar_camara_datas(inicio_periodo, fim_periodo)
    logger.info(f"Encontradas {len(resultados_camara)} proposições na Câmara.")

    logger.info("Coletando dados do Senado...")
    resultados_senado = coletar_senado_unificado()
    logger.info(f"Encontradas {len(resultados_senado)} proposições no Senado.")

    logger.info("Unificando resultados...")
    df_camara = pd.DataFrame(resultados_camara)
    df_senado = pd.DataFrame(resultados_senado)
    df_final = pd.concat([df_camara, df_senado], ignore_index=True)

    logger.info(f"Tamanho do DataFrame após concatenação: {df_final.shape}")
    if df_final.empty:
        logger.warning("Nenhum resultado após coleta unificada.")
        return pd.DataFrame()

    df_final = ordenar_por_ultima_tramitacao(df_final)
    logger.info(f"Tamanho final após ordenar: {df_final.shape}")
    return df_final

def busca_inicial():
    logger.info("Executando busca inicial (gerando XLSX/DB).")
    df = buscar_dados()
    if not df.empty:
        salvar_df(df)
        logger.info("Busca inicial OK e destino(s) salvo(s).")
    else:
        logger.warning("Busca inicial não retornou dados.")

def busca_atualizacao():
    logger.info("Executando atualização (via agendador/endpoint).")
    df = buscar_dados()
    if not df.empty:
        salvar_df(df)
        logger.info("Atualização concluída (XLSX/DB).")
    else:
        logger.warning("Nenhum dado retornado na atualização.")

def iniciar_agendador():
    global scheduler_global
    if scheduler_global is None:
        logger.info("Iniciando o agendador de atualizações...")
        scheduler = BackgroundScheduler()
        scheduler.add_job(busca_atualizacao, 'interval', minutes=30, id='atualizacao_periodica')
        scheduler.start()
        logger.info("Agendador iniciado com sucesso!")
        scheduler_global = scheduler
    else:
        logger.warning("Agendador já estava em execução.")

def obter_opcoes_situacao():
    df = carregar_df()
    if df.empty or "SituacaoAtual" not in df.columns:
        return []
    return [{"label": s, "value": s} for s in sorted(df["SituacaoAtual"].dropna().unique())]

# ============================================
# DASH – DEFINIÇÃO DO LAYOUT
# ============================================
from dash import Dash, html, dcc, dash_table
import dash_bootstrap_components as dbc
from dash.dependencies import Input, Output, State

ultima_mensagem_log = "Carregando dados..."

app = Dash(
    __name__,
    external_stylesheets=[dbc.themes.BOOTSTRAP],
    requests_pathname_prefix="/busca-reqs/",
    routes_pathname_prefix="/busca-reqs/",
    server=server
)
app.title = "Reqs de Interesse"

app.layout = dbc.Container([
    html.H1("Requerimentos de Interesse do MPO - Câmara e Senado", style={"color": "#183EFF", "fontFamily": "Verdana"}),
    html.Hr(),
    html.P(id="mensagem-log", children=ultima_mensagem_log, style={
        "whiteSpace": "pre-wrap",
        "font-style": "italic",
        "color": "#6c757d",
        "text-align": "right",
        "font-size": "0.9em"
    }),
    dbc.Button("Exportar XLSX", id="btn-exportar", color="primary", className="mb-3"),
    dbc.Button("Gerar Relatório Word", id="btn-gerar-word", color="success", className="mb-3", style={"marginLeft": "10px"}),
    dcc.Download(id="download-xlsx"),
    dcc.Download(id="download-docx"),
    # Filtros
    html.Div([
        html.Label("Filtrar por Situação Atual:", style={"fontFamily": "Verdana"}),
        dcc.Dropdown(
            id="filtros-situacao",
            options=obter_opcoes_situacao(),
            value=[],
            multi=True,
            placeholder="Selecione uma ou mais situações...",
            style={"fontFamily": "Verdana"}
        )
    ], className="mb-4"),
    html.Div([
        html.Label("Filtrar por Tipo de Proposição:", style={"fontFamily": "Verdana"}),
        dcc.Checklist(
            id="filtros-tipoprop",
            options=[
                {"label": "RIC", "value": "RIC"},
                {"label": "RQS", "value": "RQS"},
                {"label": "REQ", "value": "REQ"},
                {"label": "INC", "value": "INC"},
            ],
            value=[],
            labelStyle={"display": "inline-block", "margin-right": "15px", "fontFamily": "Verdana"},
            inputStyle={"margin-right": "5px"}
        )
    ], className="mb-4"),
    html.Div([
        html.Label("Filtrar por Ano de Apresentação:", style={"fontFamily": "Verdana"}),
        dcc.Checklist(
            id="filtros-ano-apresentacao",
            options=[
                {"label": "2023", "value": "2023"},
                {"label": "2024", "value": "2024"},
                {"label": "2025", "value": "2025"}
            ],
            value=[],
            labelStyle={"display": "inline-block", "margin-right": "15px", "fontFamily": "Verdana"},
            inputStyle={"margin-right": "5px"}
        )
    ], className="mb-4"),
    html.Div([
        html.Label("Filtrar por Unidade/Autoridade:", style={"fontFamily": "Verdana"}),
        dcc.Checklist(
            id="filtros-termos",
            options=[
                {"label": " Tebet/Ministra", "value": "grupo1"},
                {"label": " IBGE", "value": "grupo2"},
                {"label": " IPEA", "value": "grupo3"},
                {"label": " SOF", "value": "grupo4"},
                {"label": " SMA", "value": "grupo5"},
                {"label": " SEPLAN", "value": "grupo6"}
            ],
            value=[],
            labelStyle={"display": "inline-block", "margin-right": "15px", "fontFamily": "Verdana"},
            inputStyle={"margin-right": "5px"}
        )
    ], className="mb-4"),
    html.Div([
        html.Label("Filtrar por Assunto:", style={"fontFamily": "Verdana"}),
        dcc.Checklist(
            id="filtros-assunto",
            options=[
                {"label": "LC 101/2000", "value": "LC"},
                {"label": "PPA",          "value": "PPA"},
                {"label": "LDO",          "value": "LDO"},
                {"label": "LOA",          "value": "LOA"},
            ],
            value=[],
            labelStyle={"display": "inline-block", "margin-right": "15px", "fontFamily": "Verdana"},
            inputStyle={"margin-right": "5px"}
        )
    ], className="mb-4"),
    html.Div([
        html.Label("Filtrar por Casa:", style={"fontFamily": "Verdana"}),
        dcc.Checklist(
            id="filtros-casa",
            options=[
                {"label": "Câmara", "value": "Câmara"},
                {"label": "Senado", "value": "Senado"},
                {"label": "Congresso Nacional", "value": "Congresso Nacional"}
            ],
            value=[],
            labelStyle={"display": "inline-block", "margin-right": "15px", "fontFamily": "Verdana"},
            inputStyle={"margin-right": "5px"}
        )
    ], className="mb-4"),
    html.Br(),
    dcc.Interval(id="intervalo-atualizacao", interval=1800000, n_intervals=0),
    dcc.Loading(
        id="loading-tabela-resultados",
        type="default",
        children=[
            dash_table.DataTable(
                id="tabela-resultados",
                columns=[
                    {"name": "Proposição", "id": "Titulo", "presentation": "markdown"},
                    {"name": "Autor", "id": "Autor"},
                    {"name": "Data Apresentação", "id": "DataApresentacao"},
                    {"name": "Ementa", "id": "Ementa"},
                    {"name": "Situação Atual", "id": "SituacaoAtual"},
                    {"name": "Data Última Tramitação", "id": "DataUltimaTramitacao"},
                    {"name": "Descrição Última Tramitação", "id": "DescricaoUltimaTramitacao"},
                    {"name": "Ficha", "id": "LinkFicha", "presentation": "markdown"},
                    {"name": "Termos Encontrados", "id": "TermosEncontrados"},
                    {"name": "Local da Última Tramitação", "id": "Local"},
                    {"name": "Prazo para Resposta", "id": "Prazo para Resposta"},
                    {"name": "Casa", "id": "Casa", "hidden": True}
                ],
                data=[],
                filter_action="native",
                sort_action="native",
                sort_mode="multi",
                fixed_rows={'headers': True},
                page_action="none",
                style_table={"maxHeight": "1500px", "overflowY": "auto"},
                style_cell={
                    "whiteSpace": "pre-line",
                    "height": "auto",
                    "textAlign": "left",
                    "padding": "5px",
                    "fontFamily": "Verdana",
                    "fontSize": "14px",
                    "color": "#000000",
                    "userSelect": "text"
                },
                style_cell_conditional=[
                    {"if": {"column_id": "Titulo"}, "width": "100px", "maxWidth": "100px", "whiteSpace": "normal"},
                    {"if": {"column_id": "DataApresentacao"}, "width": "130px", "maxWidth": "130px", "whiteSpace": "normal"},
                    {"if": {"column_id": "Ementa"}, "width": "300px", "maxWidth": "300px", "whiteSpace": "normal"},
                    {"if": {"column_id": "SituacaoAtual"}, "width": "90px", "maxWidth": "90px", "whiteSpace": "normal"},
                    {"if": {"column_id": "DataUltimaTramitacao"}, "width": "110px", "maxWidth": "110px", "whiteSpace": "normal"},
                    {"if": {"column_id": "DescricaoUltimaTramitacao"}, "width": "350px", "maxWidth": "350px", "whiteSpace": "normal"},
                    {"if": {"column_id": "LinkFicha"}, "width": "75px", "maxWidth": "75px", "whiteSpace": "normal"},
                    {"if": {"column_id": "LinkInteiroTeor"}, "width": "75px", "maxWidth": "75px", "whiteSpace": "normal"},
                    {"if": {"column_id": "TermosEncontrados"}, "width": "120px", "maxWidth": "120px", "whiteSpace": "normal"},
                    {"if": {"column_id": "Autor"}, "width": "160px", "maxWidth": "160px", "whiteSpace": "normal"},
                    {"if": {"column_id": "Local"}, "width": "100px", "maxWidth": "100px", "whiteSpace": "normal"},
                    {"if": {"column_id": "Prazo para Resposta"}, "width": "100px", "maxWidth": "100px", "whiteSpace": "normal"},
                    {"if": {"column_id": "Casa"}, "width": "70px", "maxWidth": "70px", "whiteSpace": "normal"}
                ],
                style_header={"backgroundColor": "#183EFF", "fontWeight": "bold", "color": "#FFFFFF", "fontFamily": "Verdana"},
                css=[
                    {"selector": ".dash-table-container .dash-cell.column-Titulo", "rule": "pointer-events: auto !important;"},
                    {"selector": ".dash-table-container .dash-cell.column-Titulo a",
                     "rule": "pointer-events: auto !important; text-decoration: underline;"},
                ],
                style_data_conditional=[
                    {"if": {"row_index": "even"}, "backgroundColor": "#F2F2F2"},
                    {"if": {"column_id": "Titulo"}, "pointerEvents": "auto", "textDecoration": "underline"},
                ],
                row_selectable="multi",
                markdown_options={"link_target": "_blank"},
            )
        ]
    ),
    html.Br(),
    html.Div([
        dbc.Button("Gerar Mensagem WhatsApp", id="gerar-mensagem-btn", color="success", className="mb-3"),
        dcc.Clipboard(
            target_id="mensagem-whatsapp",
            title="Copiar Mensagem",
            style={"cursor": "pointer", "fontSize": 20, "margin-left": "10px"}
        )
    ], style={"display": "flex", "alignItems": "center"}),
    dcc.Textarea(id="mensagem-whatsapp", style={"width": "100%", "height": "150px", "margin-top": "10px"},
                 placeholder="A mensagem gerada aparecerá aqui...")
], fluid=True, style={"backgroundColor": "#FFFFFF", "color": "#000000", "padding": "20px"})

# ============================================
# CALLBACKS
# ============================================
@app.callback(
    Output("mensagem-log", "children"),
    Input("tabela-resultados", "data")
)
def atualizar_mensagem_exibida(_):
    return pegar_data_ultima_atualizacao()

@app.callback(
    Output("tabela-resultados", "data"),
    [
        Input("intervalo-atualizacao", "n_intervals"),
        Input("filtros-termos", "value"),
        Input("filtros-situacao", "value"),
        Input("filtros-tipoprop", "value"),
        Input("filtros-ano-apresentacao", "value"),
        Input("filtros-casa", "value"),
        Input("filtros-assunto", "value"),
    ]
)
def atualizar_tabelas(n_intervals, filtros_termos, filtros_situacao,
                      filtros_tipos, filtros_anos, filtros_casa, filtros_assunto):
    df_latest = carregar_df()
    if df_latest.empty:
        return []
    
    if "LinkInteiroTeor" in df_latest.columns and "Titulo" in df_latest.columns:
        def rebuild(row):
            m = re.match(r"\[([^\]]+)\]", row.get("Titulo") or "")
            label = m.group(1) if m else (row.get("Titulo") or "")
            url   = row.get("LinkInteiroTeor", "") or ""
            return f"[{label}]({url})" if url else label
        df_latest["Titulo"] = df_latest.apply(rebuild, axis=1)

    if "LinkFicha" in df_latest.columns:
        df_latest["LinkFicha"] = df_latest["LinkFicha"].apply(
            lambda x: f"[Ver Ficha]({x})" if pd.notnull(x) and x else ""
        )

    df_filtrado = df_latest.copy()

    if filtros_assunto:
        mask = pd.Series(False, index=df_filtrado.index)
        if "LC"  in filtros_assunto:
            mask |= df_filtrado["TermosEncontrados"].str.contains(
                "Lei Complementar nº 101", case=False, na=False)
        if "PPA" in filtros_assunto:
            mask |= df_filtrado["TermosEncontrados"].str.contains(
                "Plano Plurianual", case=False, na=False)
        if "LDO" in filtros_assunto:
            mask |= (
                df_filtrado["TermosEncontrados"].str.contains(
                    "Lei de Diretrizes Orçamentárias", case=False, na=False
                ) |
                df_filtrado["TermosEncontrados"].str.contains(
                    "Projeto de Lei de Diretrizes Orçamentárias", case=False, na=False
                )
            )
        if "LOA" in filtros_assunto:
            mask |= (
                df_filtrado["TermosEncontrados"].str.contains(
                    "Lei Orçamentária Anual", case=False, na=False
                ) |
                df_filtrado["TermosEncontrados"].str.contains(
                    "Projeto de Lei Orçamentária Anual", case=False, na=False
                )
            )
        df_filtrado = df_filtrado[mask]

    grupo1_terms = ["Tebet", "Ministério do Planejamento", "Ministério do Planejamento e Orçamento",
                    "Ministra do Planejamento", "Ministra de Estado do Planejamento e Orçamento", "(MPO)"]
    grupo2_terms = ["IBGE", "Instituto Brasileiro de Geografia e Estatística", "Pochmann"]
    grupo3_terms = ["Luciana Mendes Santos Servo", "IPEA", "Instituto de Pesquisa Econômica Aplicada"]
    grupo4_terms = ["Clayton Luiz Montes", "Secretaria de Orçamento Federal", "SOF"]
    grupo5_terms = ["Sergio Pinheiro Firpo", "Secretaria de Monitoramento e Avaliação de Políticas Públicas e Assuntos Econômicos"]
    grupo6_terms = ["Virginia de Angelis", "Secretaria Nacional de Planejamento", "SEPLAN"]
    grupo7_terms = ["Lei Complementar nº 101"]

    termos_filtragem = []
    if not filtros_termos:
        filtros_termos = []
    if "grupo1" in filtros_termos: termos_filtragem.extend(grupo1_terms)
    if "grupo2" in filtros_termos: termos_filtragem.extend(grupo2_terms)
    if "grupo3" in filtros_termos: termos_filtragem.extend(grupo3_terms)
    if "grupo4" in filtros_termos: termos_filtragem.extend(grupo4_terms)
    if "grupo5" in filtros_termos: termos_filtragem.extend(grupo5_terms)
    if "grupo6" in filtros_termos: termos_filtragem.extend(grupo6_terms)
    if "grupo7" in filtros_termos: termos_filtragem.extend(grupo7_terms)

    if not df_filtrado.empty:
        if termos_filtragem:
            df_filtrado = df_filtrado[
                df_filtrado["TermosEncontrados"].apply(
                    lambda x: any(t.lower() in x.lower() for t in termos_filtragem)
                    if isinstance(x, str) else False
                )
            ]
        if filtros_situacao:
            df_filtrado = df_filtrado[df_filtrado["SituacaoAtual"].isin(filtros_situacao)]
        if filtros_tipos:
            siglas = df_filtrado["Titulo"].str.extract(r"\[?(RIC|RQS|REQ|INC)\b", expand=False)
            df_filtrado = df_filtrado[siglas.isin(filtros_tipos)]
        if filtros_anos:
            df_filtrado["DataApresentacao_dt"] = pd.to_datetime(
                df_filtrado["DataApresentacao"],
                dayfirst=True,
                errors="coerce"
            )
            anos_selecionados = [int(a) for a in filtros_anos if str(a).isdigit()]
            df_filtrado = df_filtrado[df_filtrado["DataApresentacao_dt"].dt.year.isin(anos_selecionados)]
            df_filtrado.drop(columns=["DataApresentacao_dt"], inplace=True, errors="ignore")
        if filtros_casa and "Casa" in df_filtrado.columns:
            df_filtrado = df_filtrado[df_filtrado["Casa"].isin(filtros_casa)]
        if not df_filtrado.empty and "DataUltimaTramitacao" in df_filtrado.columns:
            df_filtrado["DataUltimaTramitacao"] = pd.to_datetime(
                df_filtrado["DataUltimaTramitacao"], errors="coerce", dayfirst=True
            )
            df_filtrado.sort_values(by="DataUltimaTramitacao", ascending=False, inplace=True)
            # ← converte de volta para string dd/mm/aaaa HH:MM:SS
            df_filtrado["DataUltimaTramitacao"] = df_filtrado["DataUltimaTramitacao"].dt.strftime("%d/%m/%Y %H:%M:%S").fillna("")
        return df_filtrado.to_dict("records")
    else:
        return []

@app.callback(
    Output("filtros-situacao", "options"),
    Input("intervalo-atualizacao", "n_intervals"),
)
def atualizar_opcoes_situacao(n_intervals):
    df = carregar_df()
    if df.empty or "SituacaoAtual" not in df.columns:
        return []
    opcs = sorted(df["SituacaoAtual"].dropna().unique())
    return [{"label": s, "value": s} for s in opcs]

@app.callback(
    Output("download-xlsx", "data"),
    Input("btn-exportar", "n_clicks"),
    State("tabela-resultados", "data"),
    prevent_initial_call=True
)
def exportar_xlsx(n_clicks, current_data):
    df_export = pd.DataFrame(current_data)
    from io import BytesIO
    import xlsxwriter
    output = BytesIO()
    with pd.ExcelWriter(output, engine="xlsxwriter") as writer:
        df_export.to_excel(writer, index=False)
    output.seek(0)
    return dash.dcc.send_bytes(output.getvalue(), "resultados_exportados.xlsx")

@app.callback(
    Output("mensagem-whatsapp", "value"),
    Input("gerar-mensagem-btn", "n_clicks"),
    State("tabela-resultados", "selected_rows"),
    State("tabela-resultados", "data")
)
def gerar_mensagem_whatsapp(n_clicks, selected_rows, data):
    if n_clicks is None or n_clicks <= 0 or not selected_rows:
        return ""
    linhas_msg = []
    for i in selected_rows:
        linha = data[i]
        titulo = linha.get("Titulo", "")
        ementa = linha.get("Ementa", "")
        autor = linha.get("Autor", "")
        data_ultima = linha.get("DataUltimaTramitacao", "")
        local = linha.get("Local", "")
        link_inteiro = linha.get("LinkInteiroTeor", "")
        link_ficha = linha.get("LinkFicha", "")
        situacao_atual = linha.get("SituacaoAtual", "")
        descricao = linha.get("DescricaoUltimaTramitacao", "")
        if data_ultima:
            dt_parsed = parse_data(data_ultima)
            if dt_parsed:
                data_ultima = dt_parsed.strftime("%d/%m/%Y %H:%M:%S")
        msg = (
            f"📌 *Proposição*: {titulo} - {autor}\n"
            f"📝 *Ementa*: {ementa}\n"
            f"🗒️ *Situação Atual*: {situacao_atual}\n"
            f"📃 *Descrição*: {descricao}\n"
            f"📍 *Local*: {local}\n"
            f"⏰ *Última Tramitação*: {data_ultima}\n"
            f"🔗 *Inteiro Teor*: {link_inteiro}\n"
            f"🔎 *Ficha*: {link_ficha}\n"
            f"----------------------------------------\n"
        )
        linhas_msg.append(msg)
    return "📢 *Atualização Legislativa* 📢\n\n" + "".join(linhas_msg)

@app.callback(
    Output("download-docx", "data"),
    Input("btn-gerar-word", "n_clicks"),
    State("tabela-resultados", "data"),
    prevent_initial_call=True
)
def gerar_relatorio_word(n_clicks, current_data):
    if n_clicks is None or n_clicks == 0:
        return None
    df_export = pd.DataFrame(current_data)
    if df_export.empty:
        return None
    from io import BytesIO
    from docx import Document
    from docx.shared import Pt
    from docx.oxml.ns import qn
    document = Document()
    styles = document.styles
    normal_style = styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
    titulo = document.add_heading("Relatório de Proposições Filtradas", level=1)
    titulo.style.font.name = 'Calibri'
    titulo.style.font.size = Pt(18)
    document.add_paragraph("")
    for _, row in df_export.iterrows():
        p_prop = document.add_paragraph()
        run_prop = p_prop.add_run(f"Proposição: {row.get('Titulo', '')}")
        run_prop.bold = True
        run_prop.font.size = Pt(14)
        run_prop.font.name = 'Calibri'
        run_prop._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')

        ementa_par = document.add_paragraph()
        ementa_run = ementa_par.add_run(f"Ementa: {row.get('Ementa', '')}")
        ementa_run.font.size = Pt(12)
        ementa_run.font.name = 'Calibri'
        ementa_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')

        situacao_par = document.add_paragraph()
        situacao_run = situacao_par.add_run(f"Situação Atual: {row.get('SituacaoAtual', '')}")
        situacao_run.font.size = Pt(12)
        situacao_run.font.name = 'Calibri'
        situacao_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')

        local_par = document.add_paragraph()
        local_run = local_par.add_run(f"Local: {row.get('Local', '')}")
        local_run.font.size = Pt(12)
        local_run.font.name = 'Calibri'
        local_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')

        ultima_par = document.add_paragraph()
        ultima_run = ultima_par.add_run(f"Última Tramitação: {row.get('DataUltimaTramitacao', '')}")
        ultima_run.font.size = Pt(12)
        ultima_run.font.name = 'Calibri'
        ultima_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')

        document.add_paragraph("")

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return dash.dcc.send_bytes(buffer.getvalue(), "Relatorio_Proposicoes.docx")

# ============================================
# ENDPOINTS DE SAÚDE (Huawei CCE readiness/liveness)
# ============================================
@server.route("/healthz", methods=["GET"])
def healthz():
    return jsonify(status="ok"), 200

@server.route("/readyz", methods=["GET"])
def readyz():
    # Opcionalmente testar DB
    if USE_DB and get_engine() is None:
        return jsonify(status="db-not-ready"), 503
    return jsonify(status="ready"), 200

# ============================================
# NOVO ENDPOINT PARA ATUALIZAÇÃO VIA CLOUD SCHEDULER
# ============================================

@server.route('/atualizar', methods=['POST'])
def atualizar_manual():
    try:
        threading.Thread(target=busca_atualizacao, daemon=True).start()
        return ("Atualização iniciada em background.", 202)
    except Exception as e:
        logger.error(f"Erro ao agendar atualização: {e}")
        return (f"Erro: {e}", 500)


# ============================================
# INICIALIZAÇÃO
# ============================================

def run_initial_search_in_background():
    """
    Função para ser executada em uma thread separada.
    Verifica se os dados iniciais são necessários e os busca.
    """
    time.sleep(5) 
    logger.info("Thread de busca inicial iniciada em segundo plano.")

    needs_initial_search = False
    if USE_DB:
        # Se usamos DB, a condição é se a tabela está vazia
        if is_db_empty():
            logger.info("Banco de dados está vazio. Disparando busca inicial.")
            needs_initial_search = True
    elif LOCAL_MODE:
        # Se não usamos DB e estamos em modo local, checa o arquivo
        if not os.path.exists(LOCAL_OUTPUT_FILE):
            logger.info("Arquivo local não encontrado. Disparando busca inicial.")
            needs_initial_search = True
    else:
        # Se não usamos DB e não estamos em modo local, checa GCS
        try:
            client = storage.Client()
            bucket = client.bucket(BUCKET_NAME)
            blob = bucket.blob(GCS_OBJECT_NAME)
            if not blob.exists():
                logger.info("Arquivo no GCS não encontrado. Disparando busca inicial.")
                needs_initial_search = True
        except Exception as e:
            logger.error(f"Erro ao verificar GCS: {e}. Disparando busca inicial por segurança.")
            needs_initial_search = True

    if needs_initial_search:
        busca_inicial()
    else:
        logger.info("Dados já existem. A busca inicial em segundo plano não é necessária.")

# O Gunicorn não executa o bloco __name__ == "__main__",
# então movemos a lógica de inicialização para o escopo global do módulo.

# 1. Garante que as tabelas do banco de dados existam (se USE_DB=true)
ensure_tables()

# 2. Inicia a busca de dados inicial em uma thread separada para não bloquear o servidor.
#    Isso garante que o servidor suba imediatamente e responda aos health checks.
initial_search_thread = threading.Thread(target=run_initial_search_in_background)
initial_search_thread.daemon = True  # Permite que a aplicação finalize mesmo se a thread estiver rodando
initial_search_thread.start()

# O comando 'app.run()' é removido. O Gunicorn cuidará de rodar o 'server'.
# if __name__ == "__main__":
#     port = int(os.getenv("PORT", "8080"))
#     app.run(host="0.0.0.0", port=port, debug=False)
