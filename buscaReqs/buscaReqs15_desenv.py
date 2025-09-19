# ----------------------------------------------------------
# IMPORTS
# ----------------------------------------------------------
import os
import io                      # já existia
import pandas as pd
import requests
import datetime
from datetime import datetime as dt
import logging
from apscheduler.schedulers.background import BackgroundScheduler
import time
import re
from flask import Flask
import dash
import threading
import pytz
from requests.adapters import HTTPAdapter, Retry
import http.client
from json import JSONDecodeError
from requests.exceptions import ChunkedEncodingError
from PyPDF2 import PdfReader

# >>> Biblioteca do Google Cloud Storage <<<
from google.cloud import storage

from concurrent.futures import ThreadPoolExecutor, as_completed
MAX_WORKERS = 12

# ----------------------------------------------------------
# FLASK SERVER / LOG / VARIÁVEIS GLOBAIS
# ----------------------------------------------------------
server = Flask(__name__)

logging.basicConfig(
    level=logging.INFO,
    format='[%(asctime)s] %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)
logger.setLevel(logging.DEBUG)

LOCAL_MODE = True   # True = teste local | False = produção (GCS)

BUCKET_NAME      = "aspar-429519_cloudbuild"
GCS_OBJECT_NAME  = "reqs/resultados_filtrados_expressa.xlsx"
LOCAL_OUTPUT_FILE = "resultados_filtrados_expressa.xlsx"

scheduler_global = None

# ----------------------------------------------------------
# TERMOS DE INTERESSE E CONSTANTES CAM/SF
# ----------------------------------------------------------
TERMOS = [
    "Tebet",
    "Ministério do Planejamento",
    "Ministério do Planejamento e Orçamento",
    "Ministra do Planejamento",
    "(MPO)",
    "Ministra de Estado do Planejamento e Orçamento",
    "Instituto Brasileiro de Geografia e Estatística",
    "Pochmann",
    "Luciana Mendes Santos Servo",
    "Instituto de Pesquisa Econômica Aplicada",
    "Clayton Luiz Montes",
    "Secretaria de Orçamento Federal",
    "(SOF)",
    "Sergio Pinheiro Firpo",
    "Secretaria de Monitoramento e Avaliação de Políticas Públicas e Assuntos Econômicos",
    "Virginia de Angelis",
    "Secretaria Nacional de Planejamento",
    "SEPLAN",
    "Lei Complementar nº 101",
    "Plano Plurianual",
    "Lei de Diretrizes Orçamentárias",
    "Projeto de Lei de Diretrizes Orçamentárias",
    "Lei Orçamentária Anual",
    "Projeto de Lei Orçamentária Anual",
    "(PLOA 2026)"
]

TIPOS_BUSCA_CAMARA = ["RIC", "INC", "REQ"]
URL_CAMARA         = "https://dadosabertos.camara.leg.br/api/v2/proposicoes"

# ==============================================================================
# FUNÇÕES AUXILIARES – SALVAR/CARREGAR XLSX
# ==============================================================================
def _forcar_https(url: str | None) -> str:
    if not url:
        return ""
    if url.startswith("http://www.camara.leg.br") or url.startswith("http://www25.senado.leg.br"):
        return url.replace("http://", "https://", 1)
    return url

def salvar_df(df: pd.DataFrame):
    if LOCAL_MODE:
        df.to_excel(LOCAL_OUTPUT_FILE, index=False)
        logger.info(f"Dados salvos localmente em {LOCAL_OUTPUT_FILE}")
    else:
        from io import BytesIO
        output = BytesIO()
        df.to_excel(output, index=False, engine='openpyxl')
        output.seek(0)
        try:
            client = storage.Client()
            bucket = client.bucket(BUCKET_NAME)
            blob = bucket.blob(GCS_OBJECT_NAME)
            blob.upload_from_file(
                output,
                content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            )
            logger.info(f"Arquivo XLSX enviado diretamente ao GCS: gs://{BUCKET_NAME}/{GCS_OBJECT_NAME}")
        except Exception as e:
            logger.error(f"Erro ao salvar DataFrame no GCS: {e}")

def carregar_df() -> pd.DataFrame:
    try:
        if LOCAL_MODE:
            if not os.path.exists(LOCAL_OUTPUT_FILE):
                logger.warning(f"Arquivo {LOCAL_OUTPUT_FILE} não encontrado.")
                return pd.DataFrame()
            df = pd.read_excel(LOCAL_OUTPUT_FILE, engine='openpyxl')
            logger.info("Arquivo carregado localmente com sucesso.")
            return df
        else:
            client = storage.Client()
            bucket = client.bucket(BUCKET_NAME)
            blob = bucket.blob(GCS_OBJECT_NAME)
            if not blob.exists():
                logger.warning(f"Objeto gs://{BUCKET_NAME}/{GCS_OBJECT_NAME} não existe.")
                return pd.DataFrame()
            from io import BytesIO
            data = BytesIO()
            blob.download_to_file(data)
            data.seek(0)
            df = pd.read_excel(data, engine='openpyxl')
            logger.info("Carregado XLSX diretamente do GCS com sucesso.")
            return df
    except Exception as e:
        logger.error(f"Erro ao carregar o arquivo: {e}")
        return pd.DataFrame()

def pegar_data_ultima_atualizacao() -> str:
    try:
        if LOCAL_MODE:
            if not os.path.exists(LOCAL_OUTPUT_FILE):
                return "Nenhum arquivo gerado ainda."
            timestamp = os.path.getmtime(LOCAL_OUTPUT_FILE)
            dt_obj = dt.fromtimestamp(timestamp)
        else:
            client = storage.Client()
            bucket = client.bucket(BUCKET_NAME)
            blob = bucket.blob(GCS_OBJECT_NAME)
            if not blob.exists():
                return "Nenhum arquivo gerado ainda."
            blob.reload()
            if not blob.updated:
                return "Arquivo sem data de atualização."
            dt_obj = blob.updated
        sp_tz = pytz.timezone("America/Sao_Paulo")
        dt_sp = dt_obj.astimezone(sp_tz)
        return dt_sp.strftime("Dados atualizados em %d/%m/%Y às %H:%M")
    except Exception as e:
        logger.error(f"Erro ao obter data de atualização: {e}")
        return "Falha ao obter data de atualização."

# ==============================================================================
# FUNÇÕES AUXILIARES GERAIS (parse, safe_request, ordenar, ...)
# ==============================================================================
def parse_data(strdata):
    formatos = [
        "%Y-%m-%dT%H:%M:%S",
        "%Y-%m-%dT%H:%M",
        "%Y-%m-%d %H:%M:%S",
        "%Y-%m-%d %H:%M",
        "%Y-%m-%d",
        "%d/%m/%Y %H:%M:%S",
        "%d/%m/%Y"
    ]
    for fmt in formatos:
        try:
            return dt.strptime(strdata, fmt)
        except:
            pass
    return None

def contem_termos(texto, termos):
    if not texto:
        return False
    texto_lower = texto.lower()
    return any(termo.lower() in texto_lower for termo in termos)

def safe_request(
    url, *, params=None, timeout=60, max_tentativas=3,
    stream=False, headers=None, **kwargs
):
    default_headers = {
        "User-Agent": (
            "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
            "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/124.0 Safari/537.36"
        ),
        "Accept": "*/*",
        "Accept-Encoding": "identity",
    }
    if headers:
        default_headers.update(headers)

    sess = requests.Session()
    retry_cfg = Retry(
        connect=max_tentativas,
        read=max_tentativas,
        total=max_tentativas,
        backoff_factor=1.5,
        status_forcelist=[500, 502, 503, 504],
        allowed_methods=["GET", "HEAD"],
        raise_on_status=False,
    )
    sess.mount("https://", HTTPAdapter(max_retries=retry_cfg))
    sess.mount("http://",  HTTPAdapter(max_retries=retry_cfg))

    for tent in range(max_tentativas):
        try:
            r = sess.get(
                url, params=params, timeout=timeout, headers=default_headers,
                stream=stream, allow_redirects=True, **kwargs
            )
            if not stream:
                try:
                    _ = r.content
                except (http.client.IncompleteRead, ChunkedEncodingError) as e:
                    raise e
            return r
        except (http.client.IncompleteRead, ChunkedEncodingError) as e:
            logger.warning(
                "safe_request - corpo incompleto em %s – tent %d/%d",
                url, tent+1, max_tentativas
            )
            time.sleep(1.5 * (tent + 1))
        except Exception as e:
            logger.warning("safe_request falhou em %s – %s", url, e)
            return None
    logger.error("safe_request esgotou tentativas em %s", url)
    return None

def ordenar_por_ultima_tramitacao(df: pd.DataFrame) -> pd.DataFrame:
    if 'DataUltimaTramitacao' in df.columns:
        df['DataUltimaTramitacao_dt'] = pd.to_datetime(
            df['DataUltimaTramitacao'], dayfirst=True, errors='coerce'
        )
        df.sort_values(by='DataUltimaTramitacao_dt', ascending=False, inplace=True)
        df.drop(columns=['DataUltimaTramitacao_dt'], inplace=True)
    return df

# =====================================================================
# === BUSCA TODOS OS TERMOS EM TODAS AS PÁGINAS DO PDF ================
# =====================================================================
def pdf_tem_termos(url_pdf: str,
                   termos: list[str] = TERMOS) -> list[str]:
    """
    Baixa o PDF (url_pdf) e devolve a lista de TERMOS encontrados
    em *todas* as páginas.  Lista vazia  ⇒  nenhum termo.
    """
    if not url_pdf:
        return []

    resp = safe_request(url_pdf, timeout=120, stream=True)
    if not (resp and resp.status_code == 200 and resp.content):
        return []

    try:
        pdf = PdfReader(io.BytesIO(resp.content))
    except Exception as e:
        logger.error(f"Erro ao abrir PDF {url_pdf}: {e}")
        return []

    termos_lower = [t.lower() for t in termos]
    encontrados: set[str] = set()

    for page in pdf.pages:                       # ← varre TODAS as páginas
        try:
            texto = (page.extract_text() or "").lower()
        except Exception:                        # página sem texto ou erro
            texto = ""

        for termo, termo_l in zip(termos, termos_lower):
            if termo_l in texto:
                encontrados.add(termo)
                # não faz early-break; precisamos continuar varrendo
                # para encontrar outros termos nas demais páginas

    return list(encontrados)

# ==============================================================================
# FUNÇÕES DE COLETA DE DADOS – CÂMARA
# ==============================================================================
def obter_local_ultima_tramitacao(tramitacao):
    sigla = tramitacao.get("siglaOrgao", "")
    uri_orgao = tramitacao.get("uriOrgao", "")
    nome_pub = ""
    if uri_orgao:
        r = safe_request(uri_orgao, timeout=30)
        if r and r.status_code == 200:
            dados_orgao = r.json().get("dados", {})
            nome_pub = dados_orgao.get("nomePublicacao", "")
    if sigla and nome_pub:
        return f"{sigla} - {nome_pub}"
    elif sigla:
        return sigla
    return ""

def obter_ultima_tramitacao_cam(det_id: int):
    """
    Retorna os dados da tramitação mais recente
    (ordenada por dataHora e, em empate, pelo maior sequencial).

    Saída:
        data_ultima (str  dd/mm/aaaa HH:MM:SS)
        desc_ultima (str  descricaoTramitacao)
        link_inteiro (str url para o teor, quando existir)
        local        (str sigla/nome do órgão)
        desp_ultima  (str despacho completo)
    """
    tram_url = f"{URL_CAMARA}/{det_id}/tramitacoes"
    tram_r   = safe_request(tram_url, timeout=30)
    tram     = tram_r.json().get("dados", []) if (tram_r and tram_r.status_code == 200) else []

    if not tram:
        return "", "", "", "", ""

    # pré-processa para ordenar
    for t in tram:
        t["dataHora_dt"]  = parse_data(t.get("dataHora", ""))
        try:
            t["sequencia_int"] = int(t.get("sequencia", 0))
        except ValueError:
            t["sequencia_int"] = 0

    tram.sort(                                       # ordem decrescente
        key=lambda x: (x["dataHora_dt"], x["sequencia_int"]),
        reverse=True
    )
    ultima = tram[0]

    data_ultima = (ultima["dataHora_dt"].strftime("%d/%m/%Y %H:%M:%S")
                   if ultima["dataHora_dt"] else "")
    desc_ultima = ultima.get("descricaoTramitacao", "") or ""
    desp_ultima = ultima.get("despacho", "") or ""
    link_inteiro = ultima.get("url", "") or ""
    local        = obter_local_ultima_tramitacao(ultima)

    return data_ultima, desc_ultima, link_inteiro, local, desp_ultima

def obter_autor_proposicao_camara(codigo_proposicao):
    logger.info(f"[DEBUG] Obtendo autor da proposição da Câmara ID {codigo_proposicao}")
    url_autores = f"https://dadosabertos.camara.leg.br/api/v2/proposicoes/{codigo_proposicao}/autores"
    try:
        response = safe_request(url_autores, timeout=30)
        if response and response.status_code == 200 and response.content:
            dados = response.json().get("dados", [])
            autor_com_ordem_1 = next((item for item in dados if item.get("ordemAssinatura")==1), None)
            autor_alvo = autor_com_ordem_1 if autor_com_ordem_1 else (dados[0] if dados else None)
            if autor_alvo:
                nome = autor_alvo.get("nome", "Autor não disponível")
                deputado_uri = autor_alvo.get("uri")
                partido = "Sem Partido"
                uf = "Sem UF"
                if deputado_uri:
                    logger.info(f"[DEBUG] Obtendo dados adicionais do deputado via URI: {deputado_uri}")
                    deputado_response = safe_request(deputado_uri, timeout=30)
                    if deputado_response and deputado_response.status_code==200 and deputado_response.content:
                        deputado_dados = deputado_response.json().get("dados", {})
                        ultimo_status = deputado_dados.get("ultimoStatus", {})
                        partido = (ultimo_status.get("siglaPartido") or "Sem Partido").strip()
                        uf = (ultimo_status.get("siglaUf") or "Sem UF").strip()
                        logger.info(f"[DEBUG] Dados do deputado obtidos: Partido={partido}, UF={uf}")
                autor_str = f"{nome} ({partido}/{uf})"
                logger.info(f"[DEBUG] Autor (ordem=1) formatado: {autor_str}")
                return autor_str
            else:
                logger.warning("[AVISO] Nenhum autor encontrado na lista.")
        else:
            logger.error(f"[ERRO] Falha ao obter autores da proposição ID {codigo_proposicao}. Status: {response.status_code if response else 'N/A'}")
    except Exception as e:
        logger.error(f"[ERRO] Exceção ao obter autores da proposição ID {codigo_proposicao}: {e}")
    return "Autor não disponível"


def obter_prazo_resposta_cam(det_id: int) -> str:
    url = f"{URL_CAMARA}/{det_id}/tramitacoes"
    r = safe_request(url, timeout=30)
    if not (r and r.status_code==200):
        return ""
    tram = r.json().get("dados", [])
    tram.sort(
        key=lambda t: parse_data(t.get("dataHora","")) or dt.min,
        reverse=True
    )
    for t in tram:
        texto = t.get("despacho") or t.get("descricaoTramitacao","")
        if "prazo para resposta" in texto.lower():
            m = re.search(r"\(de\s*([^\)]+)\)", texto, flags=re.I)
            if m:
                return m.group(1).strip()
    return ""

def _baixar_processar_camara(uri_prop: str, termos=TERMOS):
    # 1) DETALHES BÁSICOS ------------------------------------------------------
    det_r = safe_request(uri_prop, timeout=30)
    if not (det_r and det_r.status_code == 200):
        return None
    det = det_r.json().get("dados", {})

    ementa  = det.get("ementa", "") or ""
    url_pdf = _forcar_https(det.get("urlInteiroTeor"))

    if not contem_termos(ementa, termos):
        logger.debug(f"[Câmara] {det.get('siglaTipo')} {det.get('numero')}/{det.get('ano')} – sem termo, descartado")
        return None

    (data_ultima,
     desc_ultima,
     link_inteiro,
     local,
     desp_ultima) = obter_ultima_tramitacao_cam(det["id"])

    # ↓↓↓ NOVO parâmetro `det['ano']` ↓↓↓
    autor_str = obter_autor_proposicao_camara(det["id"])
    ficha_link = (f"https://www.camara.leg.br/proposicoesWeb/"
                  f"fichadetramitacao?idProposicao={det['id']}")

    termos_encontrados = [t for t in termos if t.lower() in ementa.lower()]
    onde_achou = "Ementa"

    # ------------------------------------------------------------
    # 5. Demais campos derivados
    # ------------------------------------------------------------
    dt_ap = parse_data(det.get("dataApresentacao", ""))
    data_apresentacao = (dt_ap.strftime("%d/%m/%Y %H:%M:%S")
                         if dt_ap else det.get("dataApresentacao", ""))

    casa_final = ("Congresso Nacional"
                  if "CMO - Comissão Mista de Orçamento" in local else "Câmara")

    prazo_resposta = obter_prazo_resposta_cam(det["id"])

    # ------------------------------------------------------------
    # 6. Retorno consolidado
    # ------------------------------------------------------------
    return {
        "CodigoMateria": det.get("id", ""),
        "Titulo": f"[{det['siglaTipo']} {det['numero']}/{det['ano']}]({url_pdf})",
        "Autor": autor_str,
        "DataApresentacao": data_apresentacao,
        "Ementa": ementa,
        "SituacaoAtual": desc_ultima,                # ← REQUISITO 1
        "DataUltimaTramitacao": data_ultima,         # ← REQUISITO 2
        "DescricaoUltimaTramitacao": desp_ultima,    # ← REQUISITO 3
        "LinkFicha": ficha_link,
        "LinkInteiroTeor": url_pdf,
        "TermosEncontrados": f"{', '.join(termos_encontrados)} ({onde_achou})",
        "Local": local,
        "Prazo para Resposta": prazo_resposta,
        "Casa": casa_final,
    }
# -----------------------------------------------------------------------------


def coletar_camara_datas(data_inicio, data_fim):
    resultados = []
    itens_por_pagina = 100
    pagina_atual = 1

    with ThreadPoolExecutor(MAX_WORKERS) as executor:
        logger.debug("→ coletar_camara_datas INICIO")
        while True:
            params = {
                "siglaTipo": ",".join(TIPOS_BUSCA_CAMARA),
                "itens": itens_por_pagina,
                "pagina": pagina_atual,
                "dataApresentacaoInicio": data_inicio,
                "dataApresentacaoFim": data_fim,
            }
            resp = safe_request(URL_CAMARA, params=params, timeout=30)
            if not (resp and resp.status_code==200):
                break

            proposicoes = resp.json().get("dados", [])
            if not proposicoes:
                break
            logger.debug(f"[Câmara] página {pagina_atual} – {len(proposicoes)} itens")        
            futuros = {
                executor.submit(_baixar_processar_camara, p["uri"]): p["id"]
                for p in proposicoes
            }

            for fut in as_completed(futuros):
                try:
                    dado = fut.result()
                    if dado:
                        resultados.append(dado)
                except Exception as e:
                    logger.error("Thread falhou para id %s – %s", futuros[fut], e)

            if len(proposicoes) < itens_por_pagina:
                break
            pagina_atual += 1
    logger.debug(f"[Câmara] total coletado: {len(resultados)}")
    return resultados

# ==============================================================================
# FUNÇÕES SENADO (inalteradas)
# ==============================================================================

# ---------------------------------------------------------------------
# === NOVA FUNÇÃO – PROCESSA UM ITEM DO SENADO ========================
# ---------------------------------------------------------------------
def _processar_item_senado(item: dict) -> dict | None:
    """
    Aceita o item quando ao menos 1 termo aparece na ementa OU no PDF.
    Agora o PDF é SEMPRE verificado (todas as páginas).
    """
    ementa = item.get("ementa", "") or ""
    termos_ementa = [t for t in TERMOS if t.lower() in ementa.lower()]

    url_pdf       = _forcar_https(item.get("urlDocumento"))
    termos_pdf    = pdf_tem_termos(url_pdf)

    # união dos termos encontrados
    termos_encontrados = sorted(set(termos_ementa + termos_pdf))
    if not termos_encontrados:          # nada encontrado ⇒ descarta
        return None

    # origem de onde achou
    if termos_ementa and termos_pdf:
        onde_achou = "Ementa + Inteiro Teor"
    elif termos_ementa:
        onde_achou = "Ementa"
    else:
        onde_achou = "Inteiro Teor"

    # ----------- 3) Situação mais recente ----------------------------
    detail_url   = f"http://legis.senado.gov.br/dadosabertos/processo/{item.get('id')}"
    detail_resp  = safe_request(detail_url, timeout=60)
    data_ultima_str = ""
    ultima_desc     = ""
    if detail_resp and detail_resp.status_code == 200:
        try:
            detail_data = detail_resp.json()
            autuacoes = detail_data.get("autuacoes", [])
            if isinstance(autuacoes, dict):
                autuacoes = [autuacoes]

            latest_dt = None
            for aut in autuacoes:
                situacoes = aut.get("situacoes", [])
                if isinstance(situacoes, dict):
                    situacoes = [situacoes]
                for sit in situacoes:
                    inicio = sit.get("inicio", "")
                    if inicio:
                        if " " not in inicio:
                            inicio += " 00:00:00"
                        dt_inicial = parse_data(inicio)
                        if dt_inicial and (latest_dt is None or dt_inicial > latest_dt):
                            latest_dt  = dt_inicial
                            ultima_desc = sit.get("descricao", "")
            if latest_dt:
                data_ultima_str = latest_dt.strftime("%d/%m/%Y %H:%M:%S")
        except Exception as e:
            logger.error(f"Erro ao processar detalhes {detail_url}: {e}")

    # ----------- 4) Datas e campos finais ----------------------------
    data_ap_raw = item.get("dataApresentacao", "")
    if data_ap_raw and " " not in data_ap_raw:
        data_ap_raw += " 00:00:00"
    dt_ap = parse_data(data_ap_raw)
    data_apresentacao = dt_ap.strftime("%d/%m/%Y %H:%M:%S") if dt_ap else data_ap_raw

    local       = item.get("siglaEnteIdentificador", "")
    casa_final  = ("Congresso Nacional"
                   if "CMO - Comissão Mista de Orçamento" in local else "Senado")
    url_pdf     = _forcar_https(item.get("urlDocumento"))

    return {
        "CodigoMateria": item.get("codigoMateria", ""),
        "Titulo": f"[{item.get('identificacao', '')}]({url_pdf})",
        "Autor": item.get("autoria", ""),
        "DataApresentacao": data_apresentacao,
        "Ementa": ementa,
        "SituacaoAtual": ultima_desc,
        "DataUltimaTramitacao": data_ultima_str,
        "DescricaoUltimaTramitacao": "",
        "LinkFicha": (
            f"https://www25.senado.leg.br/web/atividade/materias/-/materia/"
            f"{item.get('codigoMateria', '')}"
        ),
        "LinkInteiroTeor": url_pdf,
        "TermosEncontrados": f"{', '.join(termos_encontrados)} ({onde_achou})",
        "Local": local,
        "Prazo para Resposta": "",
        "Casa": casa_final,
    }


# =====================================================================
# FUNÇÃO SENADO – AGORA BUSCA RÁPIDA EM RQS/REQ + PDF ================
# =====================================================================
def coletar_senado_novo():
    """
    Coleta todos os 'REQ' e 'RQS' do Senado, aplica filtro de TERMOS
    primeiro na ementa e, se necessário, no PDF (urlDocumento).
    A etapa do PDF roda em paralelo para manter a execução rápida.
    """
    resultados = []
    seen       = set()
    headers    = {"Accept": "application/json"}
    siglas     = ["req", "rqs"]

    with ThreadPoolExecutor(MAX_WORKERS) as executor:
        futuros = []

        for sigla in siglas:
            url = f"http://legis.senado.gov.br/dadosabertos/processo?sigla={sigla}"
            r = safe_request(url, headers=headers, timeout=120)
            if not (r and r.status_code == 200):
                logger.error(f"Falha ao acessar {url}. Status: {r.status_code if r else 'N/A'}")
                continue

            try:
                data = r.json()          # a API já devolve uma lista
            except Exception as e:
                logger.error(f"Erro ao decodificar JSON de {url}: {e}")
                continue

            for item in data:
                key = item.get("codigoMateria") or item.get("identificacao", "")
                if key in seen:
                    continue
                seen.add(key)
                futuros.append(executor.submit(_processar_item_senado, item))

        for fut in as_completed(futuros):
            try:
                res = fut.result()
                if res:
                    resultados.append(res)
            except Exception as e:
                logger.error(f"Thread Senado falhou: {e}")

    return resultados


def coletar_senado_unificado():
    return coletar_senado_novo()

# ==============================================================================
# === ALTERADO ===  BUSCA PRINCIPAL
# ==============================================================================
def buscar_dados() -> pd.DataFrame:
    logger.info("Iniciando coleta de dados unificada (Câmara e Senado).")

    hoje = dt.now()
    inicio_periodo = (hoje - datetime.timedelta(days=90)).strftime("%Y-%m-%d")
    fim_periodo    = hoje.strftime("%Y-%m-%d")

    logger.info("Coletando dados da Câmara...")
    resultados_camara = coletar_camara_datas(inicio_periodo, fim_periodo)
    logger.info(f"Encontradas {len(resultados_camara)} proposições na Câmara.")

    logger.info("Coletando dados do Senado...")
    resultados_senado = coletar_senado_unificado()
    logger.info(f"Encontradas {len(resultados_senado)} proposições no Senado.")

    logger.info("Unificando resultados...")
    df_camara = pd.DataFrame(resultados_camara)
    df_senado = pd.DataFrame(resultados_senado)
    df_final = pd.concat([df_camara, df_senado], ignore_index=True)

    logger.info(f"Tamanho do DataFrame após concatenação: {df_final.shape}")
    if df_final.empty:
        logger.warning("Nenhum resultado após coleta unificada.")
        return pd.DataFrame()

    df_final = ordenar_por_ultima_tramitacao(df_final)
    logger.info(f"Tamanho final após ordenar: {df_final.shape}")
    return df_final

def busca_inicial():
    logger.info("Executando busca inicial (gerando XLSX direto no destino).")
    df = buscar_dados()
    if not df.empty:
        salvar_df(df)
        logger.info("Busca inicial OK e arquivo salvo.")
    else:
        logger.warning("Busca inicial não retornou dados.")

def busca_atualizacao():
    logger.info("Executando atualização (via agendador).")
    df = buscar_dados()
    if not df.empty:
        salvar_df(df)
        logger.info("Atualização concluída e arquivo XLSX enviado ao destino.")
    else:
        logger.warning("Nenhum dado retornado na atualização.")

def iniciar_agendador():
    global scheduler_global
    if scheduler_global is None:
        logger.info("Iniciando o agendador de atualizações...")
        scheduler = BackgroundScheduler()
        scheduler.add_job(busca_atualizacao, 'interval', minutes=30, id='atualizacao_periodica')
        scheduler.start()
        logger.info("Agendador iniciado com sucesso!")
        scheduler_global = scheduler
    else:
        logger.warning("Agendador já estava em execução.")

def obter_opcoes_situacao():
    df = carregar_df()
    if df.empty or "SituacaoAtual" not in df.columns:
        return []
    return [{"label": s, "value": s} for s in sorted(df["SituacaoAtual"].dropna().unique())]

# ============================================
# DASH – DEFINIÇÃO DO LAYOUT
# ============================================
from dash import Dash, html, dcc, dash_table
import dash_bootstrap_components as dbc
from dash.dependencies import Input, Output, State

ultima_mensagem_log = "Carregando dados..."

app = Dash(__name__, external_stylesheets=[dbc.themes.BOOTSTRAP], title="Reqs de Interesse", server=server)

app.layout = dbc.Container([
    html.H1("Requerimentos de Interesse do MPO - Câmara e Senado", style={"color": "#183EFF", "fontFamily": "Verdana"}),
    html.Hr(),
    html.P(id="mensagem-log", children=ultima_mensagem_log, style={
        "whiteSpace": "pre-wrap",
        "font-style": "italic",
        "color": "#6c757d",
        "text-align": "right",
        "font-size": "0.9em"
    }),
    dbc.Button("Exportar XLSX", id="btn-exportar", color="primary", className="mb-3"),
    dbc.Button("Gerar Relatório Word", id="btn-gerar-word", color="success", className="mb-3", style={"marginLeft": "10px"}),
    dcc.Download(id="download-xlsx"),
    dcc.Download(id="download-docx"),
    # Filtros
    html.Div([
        html.Label("Filtrar por Situação Atual:", style={"fontFamily": "Verdana"}),
        dcc.Dropdown(
            id="filtros-situacao",
            options=obter_opcoes_situacao(),
            value=[],
            multi=True,
            placeholder="Selecione uma ou mais situações...",
            style={"fontFamily": "Verdana"}
        )
    ], className="mb-4"),
    html.Div([
        html.Label("Filtrar por Tipo de Proposição:", style={"fontFamily": "Verdana"}),
        dcc.Checklist(
            id="filtros-tipoprop",
            options=[
                {"label": "RIC", "value": "RIC"},
                {"label": "RQS", "value": "RQS"},
                {"label": "REQ", "value": "REQ"},
                {"label": "INC", "value": "INC"},
            ],
            value=[],
            labelStyle={"display": "inline-block", "margin-right": "15px", "fontFamily": "Verdana"},
            inputStyle={"margin-right": "5px"}
        )
    ], className="mb-4"),
    html.Div([
        html.Label("Filtrar por Ano de Apresentação:", style={"fontFamily": "Verdana"}),
        dcc.Checklist(
            id="filtros-ano-apresentacao",
            options=[
                {"label": "2023", "value": "2023"},
                {"label": "2024", "value": "2024"},
                {"label": "2025", "value": "2025"}
            ],
            value=[],
            labelStyle={"display": "inline-block", "margin-right": "15px", "fontFamily": "Verdana"},
            inputStyle={"margin-right": "5px"}
        )
    ], className="mb-4"),
    html.Div([
        html.Label("Filtrar por Unidade/Autoridade:", style={"fontFamily": "Verdana"}),
        dcc.Checklist(
            id="filtros-termos",
            options=[
                {"label": " Tebet/Ministra", "value": "grupo1"},
                {"label": " IBGE", "value": "grupo2"},
                {"label": " IPEA", "value": "grupo3"},
                {"label": " SOF", "value": "grupo4"},
                {"label": " SMA", "value": "grupo5"},
                {"label": " SEPLAN", "value": "grupo6"}
            ],
            value=[],
            labelStyle={"display": "inline-block", "margin-right": "15px", "fontFamily": "Verdana"},
            inputStyle={"margin-right": "5px"}
        )
    ], className="mb-4"),
    # Novo filtro por Assunto
    html.Div([
        html.Label("Filtrar por Assunto:", style={"fontFamily": "Verdana"}),
        dcc.Checklist(
            id="filtros-assunto",
            options=[
                {"label": "LC 101/2000", "value": "LC"},
                {"label": "PPA",          "value": "PPA"},
                {"label": "LDO",          "value": "LDO"},
                {"label": "LOA",          "value": "LOA"},
            ],
            value=[],
            labelStyle={"display": "inline-block", "margin-right": "15px", "fontFamily": "Verdana"},
            inputStyle={"margin-right": "5px"}
        )
    ], className="mb-4"),
    html.Div([
        html.Label("Filtrar por Casa:", style={"fontFamily": "Verdana"}),
        dcc.Checklist(
            id="filtros-casa",
            options=[
                {"label": "Câmara", "value": "Câmara"},
                {"label": "Senado", "value": "Senado"},
                {"label": "Congresso Nacional", "value": "Congresso Nacional"}
            ],
            value=[],
            labelStyle={"display": "inline-block", "margin-right": "15px", "fontFamily": "Verdana"},
            inputStyle={"margin-right": "5px"}
        )
    ], className="mb-4"),
    html.Br(),
dcc.Interval(id="intervalo-atualizacao", interval=1800000, n_intervals=0),
dcc.Loading(
    id="loading-tabela-resultados",
    type="default",  # Você pode escolher entre 'default', 'circle', 'dot', 'cube', 'graph'
    children=[
        dash_table.DataTable(
            id="tabela-resultados",
            columns=[
                {"name": "Proposição", "id": "Titulo", "presentation": "markdown"},
                {"name": "Autor", "id": "Autor"},
                {"name": "Data Apresentação", "id": "DataApresentacao"},
                {"name": "Ementa", "id": "Ementa"},
                {"name": "Situação Atual", "id": "SituacaoAtual"},
                {"name": "Data Última Tramitação", "id": "DataUltimaTramitacao"},
                {"name": "Descrição Última Tramitação", "id": "DescricaoUltimaTramitacao"},
                {"name": "Ficha", "id": "LinkFicha", "presentation": "markdown"},
                {"name": "Termos Encontrados", "id": "TermosEncontrados"},
                {"name": "Local da Última Tramitação", "id": "Local"},
                {"name": "Prazo para Resposta", "id": "Prazo para Resposta"},
                {"name": "Casa", "id": "Casa", "hidden": True}
            ],
            data=[],
            filter_action="native",
            sort_action="native",
            sort_mode="multi",
            fixed_rows={'headers': True},
            page_action="none",
            style_table={"maxHeight": "1500px", "overflowY": "auto"},
            style_cell={
                "whiteSpace": "pre-line",
                "height": "auto",
                "textAlign": "left",
                "padding": "5px",
                "fontFamily": "Verdana",
                "fontSize": "14px",
                "color": "#000000",
                "userSelect": "text"
            },
            style_cell_conditional=[
                {"if": {"column_id": "Titulo"}, "width": "100px", "maxWidth": "100px", "whiteSpace": "normal"},
                {"if": {"column_id": "DataApresentacao"}, "width": "130px", "maxWidth": "130px", "whiteSpace": "normal"},
                {"if": {"column_id": "Ementa"}, "width": "300px", "maxWidth": "300px", "whiteSpace": "normal"},
                {"if": {"column_id": "SituacaoAtual"}, "width": "90px", "maxWidth": "90px", "whiteSpace": "normal"},
                {"if": {"column_id": "DataUltimaTramitacao"}, "width": "110px", "maxWidth": "110px", "whiteSpace": "normal"},
                {"if": {"column_id": "DescricaoUltimaTramitacao"}, "width": "350px", "maxWidth": "350px", "whiteSpace": "normal"},
                {"if": {"column_id": "LinkFicha"}, "width": "75px", "maxWidth": "75px", "whiteSpace": "normal"},
                {"if": {"column_id": "LinkInteiroTeor"}, "width": "75px", "maxWidth": "75px", "whiteSpace": "normal"},
                {"if": {"column_id": "TermosEncontrados"}, "width": "120px", "maxWidth": "120px", "whiteSpace": "normal"},
                {"if": {"column_id": "Autor"}, "width": "160px", "maxWidth": "160px", "whiteSpace": "normal"},
                {"if": {"column_id": "Local"}, "width": "100px", "maxWidth": "100px", "whiteSpace": "normal"},
                {"if": {"column_id": "Prazo para Resposta"}, "width": "100px", "maxWidth": "100px", "whiteSpace": "normal"},
                {"if": {"column_id": "Casa"}, "width": "70px", "maxWidth": "70px", "whiteSpace": "normal"}
            ],
            style_header={"backgroundColor": "#183EFF", "fontWeight": "bold", "color": "#FFFFFF", "fontFamily": "Verdana"},
            css=[
                {
                    "selector": ".dash-table-container .dash-cell.column-Titulo",
                    "rule":     "pointer-events: auto !important;"
                },
                {
                    "selector": ".dash-table-container .dash-cell.column-Titulo a",
                    "rule":     "pointer-events: auto !important; text-decoration: underline;"
                },
            ],
            style_data_conditional=[
                # linhas pares com fundo cinza
                {"if": {"row_index": "even"}, "backgroundColor": "#F2F2F2"},

                # libera o clique *nas células* da coluna Titulo
                {
                    "if": {"column_id": "Titulo"},
                    "pointerEvents": "auto",     # <- camelCase
                    "textDecoration": "underline",
                },
            ],
            row_selectable="multi",
            markdown_options={"link_target": "_blank"},
        )
    ]
),
html.Br(),
    html.Div([
        dbc.Button("Gerar Mensagem WhatsApp", id="gerar-mensagem-btn", color="success", className="mb-3"),
        dcc.Clipboard(
            target_id="mensagem-whatsapp",
            title="Copiar Mensagem",
            style={"cursor": "pointer", "fontSize": 20, "margin-left": "10px"}
        )
    ], style={"display": "flex", "alignItems": "center"}),
    dcc.Textarea(id="mensagem-whatsapp", style={"width": "100%", "height": "150px", "margin-top": "10px"},
                 placeholder="A mensagem gerada aparecerá aqui...")
], fluid=True, style={"backgroundColor": "#FFFFFF", "color": "#000000", "padding": "20px"})

# ============================================
# CALLBACKS
# ============================================
@app.callback(
    Output("mensagem-log", "children"),
    Input("tabela-resultados", "data")
)
def atualizar_mensagem_exibida(_):
    """
    Em vez de exibir "próxima atualização", exibimos a data de atualização do arquivo.
    """
    return pegar_data_ultima_atualizacao()

@app.callback(
    Output("tabela-resultados", "data"),
    [
        Input("intervalo-atualizacao", "n_intervals"),
        Input("filtros-termos", "value"),
        Input("filtros-situacao", "value"),
        Input("filtros-tipoprop", "value"),
        Input("filtros-ano-apresentacao", "value"),
        Input("filtros-casa", "value"),
        Input("filtros-assunto", "value"),
    ]
)
def atualizar_tabelas(n_intervals, filtros_termos, filtros_situacao,
                      filtros_tipos, filtros_anos, filtros_casa, filtros_assunto):
    """
    Lê o XLSX diretamente (do GCS ou localmente) e aplica os filtros.
    """
    df_latest = carregar_df()
    if df_latest.empty:
        return []
    
    if "LinkInteiroTeor" in df_latest.columns and "Titulo" in df_latest.columns:
        def rebuild(row):
            # extrai o texto que está entre colchetes [...]
            m = re.match(r"\[([^\]]+)\]", row["Titulo"] or "")
            label = m.group(1) if m else row["Titulo"] or ""
            url   = row.get("LinkInteiroTeor", "") or ""
            return f"[{label}]({url})" if url else label

        df_latest["Titulo"] = df_latest.apply(rebuild, axis=1)

    # Ajuste de links (Ficha, InteiroTeor)
    if "LinkFicha" in df_latest.columns:
        df_latest["LinkFicha"] = df_latest["LinkFicha"].apply(
            lambda x: f"[Ver Ficha]({x})" if pd.notnull(x) and x else ""
        )

    df_filtrado = df_latest.copy()
    # --- FILTRO POR ASSUNTO ---
    if filtros_assunto:
        mask = pd.Series(False, index=df_filtrado.index)
        if "LC"  in filtros_assunto:
            mask |= df_filtrado["TermosEncontrados"].str.contains(
                "Lei Complementar nº 101", case=False, na=False)
        if "PPA" in filtros_assunto:
            mask |= df_filtrado["TermosEncontrados"].str.contains(
                "Plano Plurianual", case=False, na=False)
        if "LDO" in filtros_assunto:
            mask |= (
                df_filtrado["TermosEncontrados"].str.contains(
                    "Lei de Diretrizes Orçamentárias", case=False, na=False
                ) |
                df_filtrado["TermosEncontrados"].str.contains(
                    "Projeto de Lei de Diretrizes Orçamentárias", case=False, na=False
                )
            )
        if "LOA" in filtros_assunto:
            mask |= (
                df_filtrado["TermosEncontrados"].str.contains(
                    "Lei Orçamentária Anual", case=False, na=False
                ) |
                df_filtrado["TermosEncontrados"].str.contains(
                    "Projeto de Lei Orçamentária Anual", case=False, na=False
                )
            )
        df_filtrado = df_filtrado[mask]
    # -------------------------------
    grupo1_terms = ["Tebet", "Ministério do Planejamento", "Ministério do Planejamento e Orçamento",
                    "Ministra do Planejamento", "Ministra de Estado do Planejamento e Orçamento", "(MPO)"]
    grupo2_terms = ["IBGE", "Instituto Brasileiro de Geografia e Estatística", "Pochmann"]
    grupo3_terms = ["Luciana Mendes Santos Servo", "IPEA", "Instituto de Pesquisa Econômica Aplicada"]
    grupo4_terms = ["Clayton Luiz Montes", "Secretaria de Orçamento Federal", "SOF"]
    grupo5_terms = ["Sergio Pinheiro Firpo", "Secretaria de Monitoramento e Avaliação de Políticas Públicas e Assuntos Econômicos"]
    grupo6_terms = ["Virginia de Angelis", "Secretaria Nacional de Planejamento", "SEPLAN"]
    grupo7_terms = ["Lei Complementar nº 101"]

    termos_filtragem = []
    if filtros_termos is None:
        filtros_termos = []
    if "grupo1" in filtros_termos:
        termos_filtragem.extend(grupo1_terms)
    if "grupo2" in filtros_termos:
        termos_filtragem.extend(grupo2_terms)
    if "grupo3" in filtros_termos:
        termos_filtragem.extend(grupo3_terms)
    if "grupo4" in filtros_termos:
        termos_filtragem.extend(grupo4_terms)
    if "grupo5" in filtros_termos:
        termos_filtragem.extend(grupo5_terms)
    if "grupo6" in filtros_termos:
        termos_filtragem.extend(grupo6_terms)
    if "grupo7" in filtros_termos:
        termos_filtragem.extend(grupo7_terms)

    if not df_filtrado.empty:
        if termos_filtragem:
            df_filtrado = df_filtrado[
                df_filtrado["TermosEncontrados"].apply(
                    lambda x: any(t.lower() in x.lower() for t in termos_filtragem)
                    if isinstance(x, str) else False
                )
            ]
        if filtros_situacao:
            df_filtrado = df_filtrado[df_filtrado["SituacaoAtual"].isin(filtros_situacao)]
        if filtros_tipos:
                siglas = df_filtrado["Titulo"].str.extract(r"\[?(RIC|RQS|REQ|INC)\b",
                                                expand=False)
                df_filtrado = df_filtrado[siglas.isin(filtros_tipos)]

        if filtros_anos:
            df_filtrado["DataApresentacao_dt"] = pd.to_datetime(
                df_filtrado["DataApresentacao"],
                dayfirst=True,
                errors="coerce"
            )
            anos_selecionados = [int(a) for a in filtros_anos if a.isdigit()]
            df_filtrado = df_filtrado[df_filtrado["DataApresentacao_dt"].dt.year.isin(anos_selecionados)]
            df_filtrado.drop(columns=["DataApresentacao_dt"], inplace=True, errors="ignore")
        if filtros_casa and "Casa" in df_filtrado.columns:
            df_filtrado = df_filtrado[df_filtrado["Casa"].isin(filtros_casa)]
        if not df_filtrado.empty and "DataUltimaTramitacao" in df_filtrado.columns:
            df_filtrado["DataUltimaTramitacao"] = pd.to_datetime(
                df_filtrado["DataUltimaTramitacao"],
                errors="coerce",
                dayfirst=True
            )
            df_filtrado.sort_values(by="DataUltimaTramitacao", ascending=False, inplace=True)
        return df_filtrado.to_dict("records")
    else:
        return []

@app.callback(
    Output("filtros-situacao", "options"),
    Input("intervalo-atualizacao", "n_intervals"),
)
def atualizar_opcoes_situacao(n_intervals):
    """
    Atualiza as opções de Situação Atual com base nos valores
    presentes na coluna 'SituacaoAtual' do arquivo XLSX.
    """
    df = carregar_df()
    if df.empty or "SituacaoAtual" not in df.columns:
        return []
    opcs = sorted(df["SituacaoAtual"].dropna().unique())
    return [{"label": s, "value": s} for s in opcs]

@app.callback(
    Output("download-xlsx", "data"),
    Input("btn-exportar", "n_clicks"),
    State("tabela-resultados", "data"),
    prevent_initial_call=True
)
def exportar_xlsx(n_clicks, current_data):
    df_export = pd.DataFrame(current_data)
    from io import BytesIO
    import xlsxwriter
    output = BytesIO()
    with pd.ExcelWriter(output, engine="xlsxwriter") as writer:
        df_export.to_excel(writer, index=False)
    output.seek(0)
    return dash.dcc.send_bytes(output.getvalue(), "resultados_exportados.xlsx")

@app.callback(
    Output("mensagem-whatsapp", "value"),
    Input("gerar-mensagem-btn", "n_clicks"),
    State("tabela-resultados", "selected_rows"),
    State("tabela-resultados", "data")
)
def gerar_mensagem_whatsapp(n_clicks, selected_rows, data):
    if n_clicks is None or n_clicks <= 0 or not selected_rows:
        return ""
    linhas_msg = []
    for i in selected_rows:
        linha = data[i]
        titulo = linha.get("Titulo", "")
        ementa = linha.get("Ementa", "")
        autor = linha.get("Autor", "")
        data_ultima = linha.get("DataUltimaTramitacao", "")
        local = linha.get("Local", "")
        link_inteiro = linha.get("LinkInteiroTeor", "")
        link_ficha = linha.get("LinkFicha", "")
        situacao_atual = linha.get("SituacaoAtual", "")
        descricao = linha.get("DescricaoUltimaTramitacao", "")
        if data_ultima:
            dt_parsed = parse_data(data_ultima)
            if dt_parsed:
                data_ultima = dt_parsed.strftime("%d/%m/%Y %H:%M:%S")
        msg = (
            f"📌 *Proposição*: {titulo} - {autor}\n"
            f"📝 *Ementa*: {ementa}\n"
            f"🗒️ *Situação Atual*: {situacao_atual}\n"
            f"📃 *Descrição*: {descricao}\n"
            f"📍 *Local*: {local}\n"
            f"⏰ *Última Tramitação*: {data_ultima}\n"
            f"🔗 *Inteiro Teor*: {link_inteiro}\n"
            f"🔎 *Ficha*: {link_ficha}\n"
            f"----------------------------------------\n"
        )
        linhas_msg.append(msg)
    return "📢 *Atualização Legislativa* 📢\n\n" + "".join(linhas_msg)

@app.callback(
    Output("download-docx", "data"),
    Input("btn-gerar-word", "n_clicks"),
    State("tabela-resultados", "data"),
    prevent_initial_call=True
)
def gerar_relatorio_word(n_clicks, current_data):
    if n_clicks is None or n_clicks == 0:
        return None
    df_export = pd.DataFrame(current_data)
    if df_export.empty:
        return None
    from io import BytesIO
    from docx import Document
    from docx.shared import Pt
    from docx.oxml.ns import qn
    document = Document()
    styles = document.styles
    normal_style = styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
    titulo = document.add_heading("Relatório de Proposições Filtradas", level=1)
    titulo.style.font.name = 'Calibri'
    titulo.style.font.size = Pt(18)
    document.add_paragraph("")
    for index, row in df_export.iterrows():
        p_prop = document.add_paragraph()
        run_prop = p_prop.add_run(f"Proposição: {row.get('Titulo', '')}")
        run_prop.bold = True
        run_prop.font.size = Pt(14)
        run_prop.font.name = 'Calibri'
        run_prop._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')

        ementa_par = document.add_paragraph()
        ementa_run = ementa_par.add_run(f"Ementa: {row.get('Ementa', '')}")
        ementa_run.font.size = Pt(12)
        ementa_run.font.name = 'Calibri'
        ementa_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')

        situacao_par = document.add_paragraph()
        situacao_run = situacao_par.add_run(f"Situação Atual: {row.get('SituacaoAtual', '')}")
        situacao_run.font.size = Pt(12)
        situacao_run.font.name = 'Calibri'
        situacao_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')

        local_par = document.add_paragraph()
        local_run = local_par.add_run(f"Local: {row.get('Local', '')}")
        local_run.font.size = Pt(12)
        local_run.font.name = 'Calibri'
        local_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')

        ultima_par = document.add_paragraph()
        ultima_run = ultima_par.add_run(f"Última Tramitação: {row.get('DataUltimaTramitacao', '')}")
        ultima_run.font.size = Pt(12)
        ultima_run.font.name = 'Calibri'
        ultima_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')

        document.add_paragraph("")

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return dash.dcc.send_bytes(buffer.getvalue(), "Relatorio_Proposicoes.docx")

# ============================================
# NOVO ENDPOINT PARA ATUALIZAÇÃO VIA CLOUD SCHEDULER
# ============================================
@server.route('/atualizar', methods=['POST'])
def atualizar_manual():
    try:
        # Chamada síncrona: executa busca_atualizacao() e retorna após sua conclusão
        busca_atualizacao()
        logger.info("Atualização concluída de forma síncrona.")
        return "Atualização concluída", 200
    except Exception as e:
        logger.error(f"Erro na atualização: {e}")
        return f"Erro: {e}", 500

# ============================================
# INICIALIZAÇÃO
# ============================================
if __name__ == "__main__":
    logger.info("Iniciando aplicação...")

    # Dispara a busca inicial somente se o arquivo ainda não existir.
    if LOCAL_MODE:
        if not os.path.exists(LOCAL_OUTPUT_FILE):
            logger.info("Arquivo local não encontrado. Disparando busca inicial de forma síncrona.")
            busca_inicial()
        else:
            logger.info("Arquivo local já existe. Não disparar busca inicial.")
    else:
        try:
            client = storage.Client()
            bucket = client.bucket(BUCKET_NAME)
            blob = bucket.blob(GCS_OBJECT_NAME)
            if not blob.exists():
                logger.info("Arquivo no GCS não encontrado. Disparando busca inicial de forma síncrona.")
                busca_inicial()
            else:
                logger.info("Arquivo no GCS já existe. Não disparar busca inicial.")
        except Exception as e:
            logger.error(f"Erro ao verificar existência do arquivo no GCS: {e}")
            # Em caso de erro, você pode optar por disparar a busca inicial ou não.
            # Aqui, optamos por não dispará-la.
    
    # Se for utilizar o agendador local para testes, descomente a linha abaixo
    #iniciar_agendador()
    
    port = int(os.getenv("PORT", "8080"))
    app.run(host="0.0.0.0", port=port, debug=False)
